"""M01 — gerador DOCX de iterações do questionário customizado.

Substitui o PDF editável: DOCX permite imagens, diagramas, tabelas e
formatação nativa dos editores (Word/Google Docs/LibreOffice). O marker
canônico pra auto-linkagem vai nos metadados `core_properties.keywords`
do DOCX, e sobrevive a edições em qualquer editor compatível com o
padrão OpenXML.
"""
from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

VIOLET = RGBColor(0x63, 0x66, 0xF1)
SLATE_DARK = RGBColor(0x1F, 0x29, 0x37)
SLATE_MID = RGBColor(0x6B, 0x72, 0x80)


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Pinta o fundo de uma célula de tabela (workaround — python-docx não expõe)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


class DocxQuestionnaireGenerator:
    """Gerador DOCX das iterações do questionário M01."""

    def generate_docx(
        self,
        project_name: str,
        questions: list[dict[str, Any]],
        iteration: int,
        iteration_id: str | None = None,
    ) -> bytes:
        """
        Gera DOCX com as questões da iteração.

        Args:
            project_name: nome do projeto.
            questions: lista de questões (id, type, text, context, pillar, options?).
            iteration: número sequencial da iteração.
            iteration_id: UUID da iteração (CustomQuestionnaireIteration.id).
                Gravado em `core_properties.keywords` como marker canônico
                `gca_iteration_id=<uuid>` — usado pela Ingestão pra
                auto-linkar a resposta à iteração correta.

        Returns:
            bytes do arquivo .docx
        """
        doc = Document()

        # Metadados canônicos
        core = doc.core_properties
        core.title = f"Questões Abertas — Iteração {iteration}"
        core.author = "GCA — Gerenciador Central de Arquiteturas"
        core.subject = "GCA M01 — resposta iterativa do questionário customizado"
        if iteration_id:
            # Marker pro detector na Ingestão (auto-linkagem).
            core.keywords = f"gca_iteration_id={iteration_id}"

        # Margens A4 razoáveis
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

        # ── Cabeçalho ──
        p = doc.add_paragraph()
        run = p.add_run("GCA")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = VIOLET
        sub = p.add_run("  Gerenciador Central de Arquiteturas")
        sub.font.size = Pt(9)
        sub.font.color.rgb = SLATE_MID

        title = doc.add_paragraph()
        t_run = title.add_run("Questões Abertas para Melhoria")
        t_run.bold = True
        t_run.font.size = Pt(15)
        t_run.font.color.rgb = SLATE_DARK

        sub_p = doc.add_paragraph()
        s_run = sub_p.add_run(f"{project_name} — Iteração {iteration}")
        s_run.bold = True
        s_run.font.size = Pt(12)
        s_run.font.color.rgb = VIOLET

        # Instruções
        doc.add_paragraph()
        instr_title = doc.add_paragraph()
        it_run = instr_title.add_run("INSTRUÇÕES")
        it_run.bold = True
        it_run.font.size = Pt(10)
        it_run.font.color.rgb = SLATE_DARK

        instrucoes = [
            "1. Responda cada questão abaixo no campo correspondente, em português-BR, de forma objetiva e factual.",
            "2. Você pode inserir imagens, diagramas, tabelas ou qualquer conteúdo relevante — o sistema analisa tudo.",
            "3. Salve o arquivo e faça upload pela aba Ingestão de Documentos. O sistema vincula automaticamente.",
            "4. Não remova o identificador técnico dos metadados do arquivo — ele é usado pra linkar esta resposta à iteração.",
        ]
        for line in instrucoes:
            p = doc.add_paragraph(line)
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = SLATE_MID

        doc.add_paragraph()

        # ── Questões ──
        q_title = doc.add_paragraph()
        qt_run = q_title.add_run("QUESTÕES")
        qt_run.bold = True
        qt_run.font.size = Pt(11)
        qt_run.font.color.rgb = SLATE_DARK

        for idx, q in enumerate(questions, start=1):
            self._render_question(doc, q, idx)

        # ── Rodapé ──
        doc.add_paragraph()
        foot = doc.add_paragraph()
        f_run = foot.add_run(
            "Depois de responder, salve o arquivo e faça upload pela aba Ingestão de Documentos do projeto."
        )
        f_run.italic = True
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = SLATE_MID
        foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    def _render_question(self, doc: Document, q: dict[str, Any], idx: int) -> None:
        qid = q.get("id") or f"Q{idx}"
        text = q.get("text") or q.get("question") or ""
        context = q.get("context") or ""
        pillar = q.get("pillar") or q.get("target_pillar") or ""
        qtype = q.get("type") or "text"
        options = q.get("options") or []

        pillar_code = "—"
        if pillar and pillar.startswith("P"):
            pillar_code = pillar.split("_", 1)[0]

        # Cabeçalho da questão: Qx [Py]  texto
        h = doc.add_paragraph()
        id_run = h.add_run(f"{qid}  [{pillar_code}]  ")
        id_run.bold = True
        id_run.font.size = Pt(11)
        id_run.font.color.rgb = VIOLET
        t_run = h.add_run(text)
        t_run.bold = True
        t_run.font.size = Pt(11)
        t_run.font.color.rgb = SLATE_DARK

        # Contexto (o que a resposta destrava)
        if context:
            ctx = doc.add_paragraph(context)
            for r in ctx.runs:
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = SLATE_MID

        # Opções (pra type=choice) — lista numerada
        if qtype == "choice" and options:
            doc.add_paragraph("Opções (escolha uma ou escreva livremente no campo abaixo):").runs[0].font.size = Pt(9)
            for i, opt in enumerate(options, start=1):
                li = doc.add_paragraph(f"  {i}. {opt}")
                for r in li.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = SLATE_DARK

        # Caixa de resposta — tabela 1×1 com borda + shading suave.
        # O usuário escreve/cola conteúdo dentro (inclusive imagens/tabelas).
        rotulo = doc.add_paragraph()
        rot_run = rotulo.add_run("Resposta:")
        rot_run.bold = True
        rot_run.font.size = Pt(9)
        rot_run.font.color.rgb = SLATE_DARK

        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.rows[0].cells[0]
        _set_cell_shading(cell, "F9FAFB")
        # Célula vazia — sem placeholder, sem formatação. Usuário digita livre.

        # Espaço depois da caixa
        doc.add_paragraph()


# Singleton global
docx_generator = DocxQuestionnaireGenerator()
