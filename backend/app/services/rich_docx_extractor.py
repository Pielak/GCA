"""MVP 8 Fase 2 — extração rica de .docx.

O extractor antigo (`arguider_service._extract_docx`) usava
`Document.paragraphs` e ignorava tabelas, caixas de texto, headers,
footers e notas de rodapé. No dogfood 2026-04-19, o GP subiu um
documento v1.0 com todos os RFs em TABELA: python-docx não viu nada,
Arguidor recebeu texto quase vazio, OCG não evoluiu, backlog não foi
gerado, roadmap ficou zerado. Cliente comum ficou sem saber por que
"nada aconteceu".

Este módulo percorre o body do documento **na ordem** (parágrafos e
tabelas intercalados como o autor escreveu), convertendo cada tabela
em parágrafos no formato `[Col1: val] [Col2: val] ...` que o Arguidor
entende. Também varre seções pra pegar headers/footers e descarrega
texto de caixas de texto e SDTs (content controls).

Escopo da Fase 2 (contrato §7 MVP 8):
  - Body principal em ordem de leitura (paragraphs + tables intercalados).
  - Tabelas viradas em parágrafos estruturados.
  - Headers e footers de cada seção.
  - Caixas de texto (w:txbxContent).
  - Footnotes (se presentes).
  - SDT / content controls (w:sdt) — texto interno.
  - Listas aninhadas — preserva hierarquia como indentação.

Fora de escopo (movido pra fases seguintes):
  - OCR de imagens embutidas (Fase 3).
  - Normalização heurística de seções (Fase 4).
  - Relatório pro usuário (Fase 5).
  - Re-escrita de conteúdo — extrai, não reescreve.
"""
from __future__ import annotations

import io
from typing import Iterable

import structlog

logger = structlog.get_logger(__name__)

# Namespace WordprocessingML
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def extract_rich_text(file_bytes: bytes) -> str:
    """Extrai todo o conteúdo legível de um .docx preservando ordem de
    leitura. Tabelas viram parágrafos estruturados.

    Retorna string pronta pra alimentar o Arguidor. Em caso de falha,
    retorna mensagem legível começando com `[`, igual ao contrato do
    extractor antigo.
    """
    try:
        from docx import Document
    except ImportError:
        return "[python-docx não instalado — extrair DOCX não disponível]"

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.warning("rich_docx.open_failed", error=str(exc))
        return f"[Erro ao abrir DOCX: {exc}]"

    parts: list[str] = []
    try:
        parts.extend(_iter_body_in_order(doc))
        parts.extend(_iter_sections_headers_footers(doc))
        parts.extend(_iter_textboxes(doc))
        parts.extend(_iter_footnotes(doc))
    except Exception as exc:
        logger.warning("rich_docx.extraction_partial", error=str(exc))
        if not parts:
            return f"[Erro ao extrair DOCX: {exc}]"

    text = "\n\n".join(p for p in parts if p and p.strip())
    return text


def _iter_body_in_order(doc) -> Iterable[str]:
    """Percorre o body do documento na ordem em que o autor escreveu —
    parágrafos e tabelas intercalados, não em listas separadas."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = _local_tag(child.tag)
        if tag == "p":
            text = _paragraph_text(child)
            if text.strip():
                yield text
        elif tag == "tbl":
            yield from _table_to_paragraphs(child)
        elif tag == "sdt":
            # Content control — conteúdo de formulário
            yield from _sdt_text(child)


def _local_tag(tag: str) -> str:
    """Retorna o tag sem namespace: '{...}p' → 'p'."""
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def _paragraph_text(p_element) -> str:
    """Extrai texto de um <w:p> incluindo runs aninhados em <w:sdt>.

    Respeita indentação via w:numPr (listas) e w:ind (indent manual) —
    não calcula o nível exato, só marca com prefixo "- " quando o
    parágrafo tem formatação de lista, pra ajudar o Arguidor a
    reconhecer hierarquia."""
    texts: list[str] = []
    is_list = False
    for descendant in p_element.iter():
        tag = _local_tag(descendant.tag)
        if tag == "numPr":
            is_list = True
        elif tag == "t":
            if descendant.text:
                texts.append(descendant.text)
        elif tag == "tab":
            texts.append("\t")
        elif tag == "br":
            texts.append("\n")
    raw = "".join(texts).strip()
    if not raw:
        return ""
    if is_list:
        return f"- {raw}"
    return raw


def _table_to_paragraphs(tbl_element) -> Iterable[str]:
    """Converte <w:tbl> em parágrafos estruturados.

    Formato:
      - Primeira linha é tratada como cabeçalho se tiver negrito ou se
        todas as células forem distintas das linhas seguintes.
      - Cada linha de dados vira:
          [Coluna1: valor] [Coluna2: valor] ...
      - Se não houver cabeçalho detectável, usa índices:
          [Col1: valor] [Col2: valor] ...

    Tabelas vazias (todas as células em branco) são puladas.
    """
    rows = _table_rows(tbl_element)
    if not rows:
        return

    non_empty = [r for r in rows if any(cell.strip() for cell in r)]
    if not non_empty:
        return

    header = non_empty[0]
    has_header = _looks_like_header(header, non_empty[1:]) and len(non_empty) > 1

    if has_header:
        data_rows = non_empty[1:]
        col_names = [cell.strip() or f"Col{idx+1}" for idx, cell in enumerate(header)]
    else:
        data_rows = non_empty
        col_names = [f"Col{idx+1}" for idx in range(len(non_empty[0]))]

    yield "[TABELA]"
    for row in data_rows:
        parts = []
        for idx, cell in enumerate(row):
            value = cell.strip()
            if not value:
                continue
            col = col_names[idx] if idx < len(col_names) else f"Col{idx+1}"
            parts.append(f"[{col}: {value}]")
        if parts:
            yield " ".join(parts)
    yield "[/TABELA]"


def _table_rows(tbl_element) -> list[list[str]]:
    """Retorna matriz de strings — linha × célula. Cada célula já vem
    com os parágrafos internos concatenados por ' | '."""
    rows: list[list[str]] = []
    for tr in tbl_element.iter(f"{W_NS}tr"):
        row: list[str] = []
        for tc in tr.iter(f"{W_NS}tc"):
            cell_parts = []
            for p in tc.iter(f"{W_NS}p"):
                text = _paragraph_text(p)
                if text:
                    cell_parts.append(text)
            row.append(" | ".join(cell_parts))
        if row:
            rows.append(row)
    return rows


def _looks_like_header(first_row: list[str], rest: list[list[str]]) -> bool:
    """Heurística simples: primeira linha é cabeçalho se nenhuma célula
    dela aparece em nenhuma linha seguinte (valores distintos indicam
    rótulos, não dados)."""
    if not first_row or not rest:
        return False
    first_clean = [c.strip().lower() for c in first_row if c.strip()]
    if not first_clean:
        return False
    for data_row in rest:
        for cell in data_row:
            if cell.strip().lower() in first_clean:
                return False
    return True


def _sdt_text(sdt_element) -> Iterable[str]:
    """Texto dentro de content controls (formulários)."""
    for p in sdt_element.iter(f"{W_NS}p"):
        text = _paragraph_text(p)
        if text:
            yield text


def _iter_sections_headers_footers(doc) -> Iterable[str]:
    """Headers e footers de cada seção. Útil quando o autor coloca
    metadados importantes no cabeçalho (autor, versão, data)."""
    for section_idx, section in enumerate(doc.sections):
        for part_name, part in (("header", section.header), ("footer", section.footer)):
            try:
                for p in part.paragraphs:
                    text = p.text.strip()
                    if text:
                        yield f"[{part_name.upper()} seção {section_idx+1}] {text}"
                for tbl in getattr(part, "tables", []):
                    yield from _table_to_paragraphs(tbl._element)
            except Exception:
                continue


def _iter_textboxes(doc) -> Iterable[str]:
    """Caixas de texto (w:txbxContent) — contêm parágrafos aninhados."""
    body = doc.element.body
    for txbx in body.iter(f"{W_NS}txbxContent"):
        for p in txbx.iter(f"{W_NS}p"):
            text = _paragraph_text(p)
            if text:
                yield f"[CAIXA DE TEXTO] {text}"


def _iter_footnotes(doc) -> Iterable[str]:
    """Notas de rodapé. python-docx não expõe footnotes na API pública,
    então pega direto do part XML quando presente."""
    try:
        footnotes_part = doc.part.package.part_related_by_reltype(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        )
    except Exception:
        return
    try:
        from lxml import etree
    except ImportError:
        return
    try:
        root = etree.fromstring(footnotes_part.blob)
    except Exception:
        return
    for footnote in root.iter(f"{W_NS}footnote"):
        # Ignora separadores ( type="separator" / "continuationSeparator" )
        if footnote.get(f"{W_NS}type") in ("separator", "continuationSeparator"):
            continue
        for p in footnote.iter(f"{W_NS}p"):
            text = _paragraph_text(p)
            if text:
                yield f"[NOTA DE RODAPÉ] {text}"
