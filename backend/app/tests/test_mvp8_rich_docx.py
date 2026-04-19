"""MVP 8 Fase 2 — testes do extractor rico de .docx.

O dogfood 2026-04-19 mostrou que documentos com RFs em tabela viravam
texto vazio para o Arguidor. Este teste constrói .docx programaticamente
em memória com os padrões problemáticos (tabela de RFs, caixa de texto,
header/footer, lista aninhada) e valida que o extractor captura tudo.

Sem dependência de fixtures em disco — `python-docx` cria o .docx in
memory, o extractor lê os bytes, os asserts checam o texto final.
"""
import io

import pytest

from app.services.rich_docx_extractor import extract_rich_text


def _build_docx_with_table() -> bytes:
    """Doc com um parágrafo de abertura + tabela 3x3 de RFs + parágrafo
    de fechamento. Reproduz o padrão do v1.0 da Automação Jurídica."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Introdução ao sistema.")

    table = doc.add_table(rows=4, cols=3)
    header = table.rows[0].cells
    header[0].text = "Código"
    header[1].text = "Nome"
    header[2].text = "Descrição"

    table.rows[1].cells[0].text = "RF-001"
    table.rows[1].cells[1].text = "Login"
    table.rows[1].cells[2].text = "Permitir login com email e senha"

    table.rows[2].cells[0].text = "RF-002"
    table.rows[2].cells[1].text = "Cadastro"
    table.rows[2].cells[2].text = "Permitir cadastro de usuário"

    table.rows[3].cells[0].text = "RF-003"
    table.rows[3].cells[1].text = "Relatórios"
    table.rows[3].cells[2].text = "Gerar relatórios em PDF"

    doc.add_paragraph("Conclusão dos requisitos funcionais.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_paragraphs_only() -> bytes:
    """Doc sem tabelas — garante que o extractor novo não regride no
    caso simples (mesmo resultado do extractor antigo)."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Primeiro parágrafo.")
    doc.add_paragraph("Segundo parágrafo.")
    doc.add_paragraph("Terceiro parágrafo com acentuação: ação, ânsia.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_with_header_footer() -> bytes:
    """Doc com metadados em header/footer — comum em templates
    corporativos (versão, autor, data)."""
    from docx import Document
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Documento de Requisitos v2.0"
    doc.sections[0].footer.paragraphs[0].text = "Confidencial — Automação Jurídica"
    doc.add_paragraph("Conteúdo do documento.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_with_nested_list() -> bytes:
    """Doc com lista hierárquica — Arguidor precisa reconhecer itens."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Entregáveis do projeto:")
    doc.add_paragraph("Módulo de autenticação", style="List Bullet")
    doc.add_paragraph("Módulo de relatórios", style="List Bullet")
    doc.add_paragraph("Módulo de administração", style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================================
# Casos que motivaram a Fase 2
# ============================================================================

def test_tabela_de_rfs_nao_desaparece():
    """O problema original: RFs em tabela ficavam invisíveis. O novo
    extractor tem que emitir cada célula de forma legível."""
    text = extract_rich_text(_build_docx_with_table())

    # Conteúdo antes e depois da tabela preservado
    assert "Introdução ao sistema" in text
    assert "Conclusão dos requisitos funcionais" in text

    # Cada RF aparece
    assert "RF-001" in text
    assert "RF-002" in text
    assert "RF-003" in text
    assert "Login" in text
    assert "Cadastro" in text
    assert "Relatórios" in text

    # Formato estruturado [Coluna: valor] foi aplicado
    assert "[Código:" in text or "[Col1:" in text
    assert "[Nome:" in text or "[Col2:" in text


def test_tabela_marca_inicio_e_fim():
    """O Arguidor se beneficia de delimitadores pra saber que bloco é
    dado tabular vs. texto corrente."""
    text = extract_rich_text(_build_docx_with_table())
    assert "[TABELA]" in text
    assert "[/TABELA]" in text


def test_ordem_paragrafos_antes_depois_da_tabela_preservada():
    """Se o autor escreveu 'introdução → tabela → conclusão', a ordem
    tem que ser mantida — senão contexto quebra."""
    text = extract_rich_text(_build_docx_with_table())
    idx_intro = text.find("Introdução")
    idx_tabela = text.find("[TABELA]")
    idx_conclusao = text.find("Conclusão")
    assert 0 <= idx_intro < idx_tabela < idx_conclusao


def test_sem_tabelas_nao_regride():
    """Fase 2 não pode quebrar docs simples — só enriquece."""
    text = extract_rich_text(_build_docx_paragraphs_only())
    assert "Primeiro parágrafo" in text
    assert "Segundo parágrafo" in text
    assert "Terceiro parágrafo com acentuação: ação, ânsia" in text


def test_header_e_footer_extraidos():
    """Versão, autor, data viram metadados no cabeçalho/rodapé —
    precisam chegar ao Arguidor."""
    text = extract_rich_text(_build_docx_with_header_footer())
    assert "Documento de Requisitos v2.0" in text
    assert "Confidencial — Automação Jurídica" in text
    assert "Conteúdo do documento" in text


def test_listas_preservam_itens():
    """Itens de lista não podem sumir. O prefixo '- ' só é aplicado
    quando o parágrafo tem <w:numPr> no XML (listas reais do Word);
    o importante aqui é que nenhum item desaparece."""
    text = extract_rich_text(_build_docx_with_nested_list())
    assert "Entregáveis do projeto:" in text
    assert "Módulo de autenticação" in text
    assert "Módulo de relatórios" in text
    assert "Módulo de administração" in text


def test_lista_com_numPr_ganha_prefixo():
    """Quando o parágrafo tem <w:numPr> (lista real do Word), o
    extractor prefixa com '- ' pro Arguidor reconhecer hierarquia.
    Testado construindo XML direto, já que python-docx não anexa
    numbering.xml automaticamente em docs criados in-memory."""
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()
    p = doc.add_paragraph("Item com numeração")
    pPr = p._p.get_or_add_pPr()
    numPr = etree.SubElement(pPr, qn("w:numPr"))
    etree.SubElement(numPr, qn("w:ilvl")).set(qn("w:val"), "0")
    etree.SubElement(numPr, qn("w:numId")).set(qn("w:val"), "1")

    buf = io.BytesIO()
    doc.save(buf)
    text = extract_rich_text(buf.getvalue())
    assert "- Item com numeração" in text


def test_bytes_invalidos_nao_quebram():
    """Lixo binário tem que retornar string com prefixo '[' — contrato
    do extractor antigo que o arguider_service espera."""
    result = extract_rich_text(b"\x00\x01\x02 not a real docx")
    assert isinstance(result, str)
    assert result.startswith("[")


def test_arquivo_vazio_nao_quebra():
    """Bytes vazios não podem explodir o pipeline."""
    result = extract_rich_text(b"")
    assert isinstance(result, str)
    assert result.startswith("[")


def test_arguider_service_usa_extractor_rico():
    """Ligação end-to-end: arguider_service._extract_docx delega ao
    novo extractor. Se alguém quebrar a ligação, este teste pega."""
    from app.services.arguider_service import DocumentExtractor
    extractor = DocumentExtractor()
    text = extractor._extract_docx(_build_docx_with_table())
    assert "RF-001" in text
    assert "[TABELA]" in text


def test_valores_unicode_preservados():
    """Acentos, ç, aspas tipográficas — sem re-codificação errada."""
    from docx import Document
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Atributo"
    t.rows[0].cells[1].text = "Valor"
    t.rows[1].cells[0].text = "Descrição"
    t.rows[1].cells[1].text = "Sistema de gestão — inclusão e edição"
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_rich_text(buf.getvalue())
    assert "Sistema de gestão — inclusão e edição" in text
    assert "Descrição" in text
