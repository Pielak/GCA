"""MVP 8 Fase 5 — Relatório de extração.

Valida que o report:
  - Roda em .docx, .pdf, .md/.txt.
  - Conta parágrafos, tabelas, caixas de texto, headers/footers.
  - Reporta camadas de PDF usadas (acroform/text).
  - Detecta RFs / RNFs / módulos por heurística simples.
  - Lida com arquivo vazio / tipo desconhecido sem explodir.
  - Preserva warnings vindos dos extractors.
"""
import io

import pytest

from app.services.extraction_report_service import build_extraction_report


# ============================================================================
# Helpers — constroem fixtures in-memory
# ============================================================================

def _docx_with_table_and_rfs() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Sistema de Gestão Jurídica — documento de requisitos.")
    doc.add_paragraph("Lista de Requisitos Funcionais:")

    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Código"
    table.rows[0].cells[1].text = "Descrição"
    table.rows[1].cells[0].text = "RF-001"
    table.rows[1].cells[1].text = "Login"
    table.rows[2].cells[0].text = "RF-002"
    table.rows[2].cells[1].text = "Cadastro"
    table.rows[3].cells[0].text = "RNF-001"
    table.rows[3].cells[1].text = "Resposta < 2s"

    doc.add_paragraph("Módulo de autenticação deve prover SSO.")
    doc.add_paragraph("Módulo de relatórios precisa de exportação PDF.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_simple_text() -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(72, 800, "RF-010: Sistema de login")
    c.drawString(72, 780, "RF-011: Dashboard")
    c.drawString(72, 760, "Módulo de autenticação")
    c.showPage()
    c.save()
    return buf.getvalue()


# ============================================================================
# .docx
# ============================================================================

def test_docx_report_conta_tabela_e_paragrafos():
    report = build_extraction_report(_docx_with_table_and_rfs(), "docx")
    assert report["ok"] is True
    assert report["chars"] > 0
    assert report["tables_detected"] == 1
    assert report["paragraphs"] > 0


def test_docx_report_detecta_rfs_e_rnfs():
    report = build_extraction_report(_docx_with_table_and_rfs(), "docx")
    assert "RF-001" in report["requirements_functional"]
    assert "RF-002" in report["requirements_functional"]
    assert "RNF-001" in report["requirements_non_functional"]


def test_docx_report_detecta_modulos():
    report = build_extraction_report(_docx_with_table_and_rfs(), "docx")
    hints = " | ".join(report["module_hints"]).lower()
    assert "módulo de autenticação" in hints or "modulo de autenticação" in hints
    assert len(report["module_hints"]) >= 1


def test_rfs_sem_duplicatas_e_limitados():
    """Se o doc repete 'RF-001' várias vezes, a lista mostra o único."""
    from docx import Document
    doc = Document()
    for _ in range(20):
        doc.add_paragraph("RF-001 aparece muitas vezes.")
    buf = io.BytesIO()
    doc.save(buf)

    report = build_extraction_report(buf.getvalue(), "docx", max_preview_items=10)
    assert report["requirements_functional"] == ["RF-001"]


# ============================================================================
# .pdf
# ============================================================================

def test_pdf_report_reporta_camada_text():
    report = build_extraction_report(_pdf_simple_text(), "pdf")
    assert report["ok"] is True
    assert "text" in report["pdf_layers"]
    assert report["pdf_pages_with_text"] >= 1
    # Detecta RFs do conteúdo
    assert "RF-010" in report["requirements_functional"]
    assert "RF-011" in report["requirements_functional"]


def test_pdf_vazio_reporta_warnings():
    report = build_extraction_report(b"", "pdf")
    assert report["ok"] is False
    assert report["warnings"]
    assert "vazio" in report["warnings"][0].lower()


def test_pdf_bytes_invalidos_nao_quebra():
    report = build_extraction_report(b"not a pdf", "pdf")
    # Pode reportar warning mas não lança exceção
    assert isinstance(report, dict)
    assert report["ok"] is False or report["chars"] == 0


# ============================================================================
# Outros tipos / casos defensivos
# ============================================================================

def test_markdown_plain_text():
    content = "# Documento\n\nRF-042: algo\n\nMódulo de pagamentos integrado.\n"
    report = build_extraction_report(content.encode("utf-8"), "md")
    assert report["ok"] is True
    assert "RF-042" in report["requirements_functional"]
    assert any("pagamentos" in m.lower() for m in report["module_hints"])


def test_tipo_desconhecido_reporta_warning():
    report = build_extraction_report(b"some bytes", "exotic-format-xyz")
    warnings_joined = " ".join(report["warnings"])
    assert "não tem extractor rico" in warnings_joined or "Nenhum texto" in warnings_joined


def test_text_sample_preview():
    """Sample é curto — 500 chars — pra UI render sem scroll."""
    long_text = ("x" * 5000).encode("utf-8")
    report = build_extraction_report(long_text, "txt")
    assert len(report["text_sample"]) <= 500


def test_report_shape_estavel():
    """Contrato com frontend: chaves sempre presentes, mesmo em caso
    de falha total. Evita bug de `undefined` no React."""
    report = build_extraction_report(b"", "pdf")
    expected_keys = {
        "ok", "file_type", "chars", "paragraphs", "tables_detected",
        "text_boxes", "headers_footers", "pdf_layers",
        "pdf_pages_with_text", "acroform_fields",
        "requirements_functional", "requirements_non_functional",
        "module_hints", "warnings", "text_sample",
    }
    assert set(report.keys()) >= expected_keys
