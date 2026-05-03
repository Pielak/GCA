"""Tests for chunker implementations."""
import pytest
from pathlib import Path
import tempfile
from docx import Document as DocxDocument
from docx.shared import Pt

from app.services.chunkers.docx_chunker import DocxChunker
from app.services.chunkers.markdown_chunker import MarkdownChunker
from app.services.chunkers.pdf_chunker import PdfChunker


@pytest.fixture
def sample_docx():
    """Create a sample DOCX file for testing."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = DocxDocument()
        doc.add_heading("Main Section", level=1)
        doc.add_paragraph("This is content in the main section.")
        doc.add_paragraph("More content here.")
        doc.add_heading("Subsection", level=2)
        doc.add_paragraph("Content in subsection.")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"

        doc.save(f.name)
        return f.name


@pytest.fixture
def sample_markdown():
    """Create a sample Markdown file for testing."""
    with tempfile.NamedTemporaryFile(suffix=".md", delete=False, mode="w") as f:
        f.write("# Main Title\n\n")
        f.write("This is content under the main title.\n\n")
        f.write("More content.\n\n")
        f.write("## Subsection\n\n")
        f.write("Subsection content.\n")
        f.flush()
        return f.name


def test_docx_chunker_basic(sample_docx):
    """Test DOCX chunker produces chunks."""
    chunker = DocxChunker()
    chunks = chunker.chunk(sample_docx)

    assert len(chunks) > 0
    assert all(c.id for c in chunks)
    assert all(c.heading_path for c in chunks)
    assert all(c.chunk_type in ["section", "table", "list", "code"] for c in chunks)
    assert any(c.chunk_type == "table" for c in chunks)


def test_docx_chunker_token_count(sample_docx):
    """Test token estimation."""
    chunker = DocxChunker()
    chunks = chunker.chunk(sample_docx)

    for chunk in chunks:
        assert chunk.token_count > 0 or len(chunk.text) == 0


def test_markdown_chunker_basic(sample_markdown):
    """Test Markdown chunker produces chunks."""
    chunker = MarkdownChunker()
    chunks = chunker.chunk(sample_markdown)

    assert len(chunks) > 0
    assert all(c.id for c in chunks)


def test_pdf_chunker_exists():
    """Test PDF chunker can be instantiated (no actual PDF test without pypdf)."""
    chunker = PdfChunker()
    assert chunker is not None


def test_chunker_estimate_tokens():
    """Test token estimation is consistent."""
    text = "This is a simple test sentence with some words in it."
    estimated = DocxChunker.estimate_tokens(text)

    # Rough check: ~10 words * 1.4 = ~14 tokens
    assert 10 <= estimated <= 20


def test_chunker_first_sentence():
    """Test first sentence extraction."""
    text = "This is the first sentence. This is the second sentence. More text."
    first = DocxChunker.first_sentence(text)

    assert "first sentence" in first
    assert "second" not in first
