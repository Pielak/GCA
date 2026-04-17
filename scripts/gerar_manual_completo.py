#!/usr/bin/env python3
"""
Gera o Manual Completo do GCA (Admin + GP) em .docx — 'Manual_GCA_Completo.docx'.

Consume:
    /home/luiz/GCA/screenshots/manifest.json
    /home/luiz/GCA/screenshots/*.png  (28 capturas)

Estrutura do manual:
    Parte I   — Introdução ao GCA
    Parte II  — Tour Visual (todas as telas com prints + descrição)
    Parte III — Tutorial Passo a Passo
    Parte IV  — O que o GCA faz / NÃO faz
    Parte V   — Perguntas e Respostas
    Parte VI  — Glossário e Suporte
"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ──────────────────────────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────────────────────────
SCREENSHOTS = Path("/home/luiz/GCA/screenshots")
MANIFEST = SCREENSHOTS / "manifest.json"
OUTPUT = Path("/home/luiz/GCA/Manual_GCA_Completo.docx")

# ──────────────────────────────────────────────────────────────────
# Cores e tipografia
# ──────────────────────────────────────────────────────────────────
VIOLET = RGBColor(0x7C, 0x3A, 0xED)
SLATE_DARK = RGBColor(0x1E, 0x29, 0x3B)
SLATE_MID = RGBColor(0x47, 0x55, 0x69)
EMERALD = RGBColor(0x05, 0x96, 0x69)
AMBER = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)
CYAN = RGBColor(0x06, 0x95, 0xA9)

IMG_WIDTH_CM = 15.5  # largura útil A4 (margens 2.5cm)


# ──────────────────────────────────────────────────────────────────
# Descrições por tela — escritas com base nas capturas reais.
# Cada item: (caption, body, role, dica?)
# ──────────────────────────────────────────────────────────────────
DESCRICOES = {
    # ─── Públicas ─────────────────────────────────────
    "01_publica_login": {
        "caption": "Tela de login unificada",
        "role": "Pública",
        "body": (
            "Ponto único de entrada no GCA. À esquerda, painel institucional com "
            "particles animados, identidade da plataforma e cards das quatro "
            "capacidades centrais (7 Pilares, IA Assistida, Multi-tenant e "
            "CodeGen). À direita, o formulário com três campos: PROJETO (combo), "
            "EMAIL e SENHA. O combo lista todos os projetos ativos da instância — "
            "o último projeto acessado fica no topo (lembrado pelo navegador). "
            "Se o usuário entra sem selecionar projeto, só admins conseguem "
            "passar (vão para a área administrativa). Selecionando um projeto, "
            "o sistema valida que a pessoa é membro daquele projeto antes de "
            "deixar entrar."
        ),
        "dica": (
            "O link 'Solicitar novo projeto' (rodapé) leva ao wizard público de "
            "submissão. Útil para pessoas que ainda não têm conta no GCA."
        ),
    },
    "02_publica_solicitar_projeto_passo1": {
        "caption": "Solicitar novo projeto — Passo 1 (Básico)",
        "role": "Pública",
        "body": (
            "Primeiro passo do wizard de solicitação. Captura: nome do "
            "solicitante, email (será o canal oficial de comunicação), nome "
            "do projeto, tipo de entregável (combo com 9 opções incluindo "
            "'Outro') e descrição obrigatória de no mínimo 30 caracteres. "
            "O stepper no topo deixa claro que existe um segundo passo. "
            "Ao escolher 'Outro', um campo extra aparece para o solicitante "
            "descrever livremente o tipo (ex: 'Browser Extension', 'CLI Tool')."
        ),
        "dica": (
            "Quem solicita um projeto vira automaticamente o GP quando o admin "
            "aprovar. A descrição alimenta o seed inicial do OCG — quanto mais "
            "rica, melhor o pipeline arranca."
        ),
    },
    "03_publica_solicitar_projeto_passo2": {
        "caption": "Solicitar novo projeto — Passo 2 (Requisitos)",
        "role": "Pública",
        "body": (
            "Segundo passo: perguntas obrigatórias específicas do tipo de "
            "entregável escolhido + duas perguntas comuns (stakeholders e "
            "critérios de sucesso). Para 'Novo sistema' as perguntas cobrem "
            "público-alvo, volume de usuários, autenticação e hospedagem. "
            "Cada tipo (mobile_app, integration, modernization, etl, etc) "
            "tem seu próprio bloco de 3-5 perguntas. As respostas são "
            "armazenadas como JSON estruturado e servem de contexto inicial "
            "para o Arguidor quando o projeto for aprovado."
        ),
        "dica": (
            "O botão 'Voltar' permite ajustar o passo 1 sem perder respostas. "
            "Validação visual mostra perguntas faltando antes do envio."
        ),
    },
    "04_publica_reset_password": {
        "caption": "Redefinição de senha",
        "role": "Pública",
        "body": (
            "Tela acessada pelo link 'Esqueci minha senha' do login OU pelo "
            "link recebido por email após solicitação de redefinição. Tokens "
            "de redefinição são válidos por 24 horas e expiram após uso. "
            "Validação de força da senha aplicada (mínimo 8 caracteres, "
            "letras maiúsculas/minúsculas, número e símbolo)."
        ),
    },
    # ─── Admin ─────────────────────────────────────
    "10_admin_dashboard": {
        "caption": "Administração GCA — Dashboard",
        "role": "Admin",
        "body": (
            "Tela inicial do administrador. Quatro cards de métricas no topo: "
            "Total de Projetos, Projetos Ativos, Pendentes de Aprovação, "
            "Usuários. Abaixo, à esquerda, lista de 'Projetos do Sistema' com "
            "status; à direita, 'Auditoria Global' mostrando os últimos eventos "
            "(BACKLOG_REGENERATED, OCG_UPDATED, etc) com timestamps. A barra "
            "lateral esquerda lista as áreas administrativas: Dashboard Global, "
            "Configurações, Auditoria Global e Provedores IA."
        ),
        "dica": (
            "Os cards são clicáveis: 'Pendentes de Aprovação' leva direto à "
            "lista filtrada de solicitações aguardando ação."
        ),
    },
    "11_admin_gestao_projetos": {
        "caption": "Gestão de Projetos",
        "role": "Admin",
        "body": (
            "Lista todos os projetos da instância (pendentes, aprovados e "
            "rejeitados). Colunas: PROJETO, TIPO, STATUS, SOLICITANTE / GP, "
            "PENDÊNCIAS, AÇÕES. O coluna 'SOLICITANTE / GP' deixa claro que "
            "quem solicita é sempre o futuro GP. As ações por linha incluem: "
            "'Detalhes da solicitação' (ícone de documento, sempre visível), "
            "'Mensagem ao solicitante' (lápis), 'Aprovar' (check verde, "
            "só pendentes), 'Rejeitar com razão' (sino âmbar, só pendentes — "
            "envia email ao solicitante) e 'Excluir sem notificar' (lixeira "
            "vermelha)."
        ),
        "dica": (
            "Antes de aprovar, sempre clique no ícone de documento para abrir "
            "o modal de detalhes. Você verá descrição completa + perguntas e "
            "respostas do wizard, podendo avaliar a aderência da solicitação."
        ),
    },
    "12_admin_gestao_projetos_detalhes_modal": {
        "caption": "Modal 'Detalhes da solicitação'",
        "role": "Admin",
        "body": (
            "Aberto pelo ícone de documento na lista. Mostra tudo o que o "
            "solicitante enviou: nome e email do solicitante, tipo de "
            "entregável (com label custom se for 'Outro'), descrição completa, "
            "respostas do wizard formatadas como pares pergunta/resposta. "
            "Perguntas não respondidas aparecem em amarelo. Para projetos "
            "rejeitados, mostra também o motivo da rejeição. Datas e status "
            "ficam no rodapé."
        ),
        "dica": (
            "Use este modal como base para decidir entre aprovar, rejeitar com "
            "razão (envia email automático) ou pedir esclarecimentos via "
            "'Mensagem ao solicitante' (mantém o status pendente)."
        ),
    },
    "13_admin_gestao_usuarios": {
        "caption": "Gestão de Usuários",
        "role": "Admin",
        "body": (
            "Lista todos os usuários da instância com discriminação clara de "
            "papéis. Colunas: USUÁRIO, PERFIL, PROJETOS / PAPEL, STATUS, "
            "ÚLTIMO ACESSO, CADASTRADO EM. Admin recebe badge especial "
            "'Admin (sistema)' com ícone de escudo — não atua em projetos por "
            "regra de negócio. Não-admins mostram pills coloridas por papel "
            "{papel · projeto} com cores distintas (gp=violeta, tech_lead=ciano, "
            "dev=verde, qa=âmbar, compliance=laranja, viewer=cinza). Ações "
            "por linha: ativar/desativar (ícone de raio) e excluir (lixeira)."
        ),
        "dica": (
            "Excluir usuários ativos é bloqueado: o sistema avisa quais projetos "
            "impedem e pede transferência prévia de papel. Uma medida de "
            "segurança contra órfãos no banco."
        ),
    },
    "14_admin_auditoria_global": {
        "caption": "Auditoria Global",
        "role": "Admin",
        "body": (
            "Log completo e cronológico de todas as ações sensíveis na "
            "instância: logins, aprovações, rejeições, geração de código, "
            "edição de OCG, exclusões. Cada entrada registra: timestamp, "
            "ator (quem), tipo de evento, recurso afetado, correlação ID e "
            "detalhes em JSON. Filtros disponíveis por tipo de evento, "
            "ator e período."
        ),
        "dica": (
            "A auditoria é imutável e idempotente. É a fonte da verdade quando "
            "houver investigação ou compliance check. Exportável para CSV."
        ),
    },
    # ─── Projeto ─────────────────────────────────────
    "20_projeto_lista_projetos": {
        "caption": "Lista de projetos do GP",
        "role": "GP",
        "body": (
            "Tela inicial após login com projeto selecionado. Mostra todos os "
            "projetos dos quais o GP é membro, com nome, slug, status (ativo, "
            "provisionando, arquivado) e ação rápida para entrar. Em uma "
            "instância com poucos projetos, geralmente só o ativo aparece."
        ),
    },
    "21_projeto_dashboard": {
        "caption": "Dashboard do projeto",
        "role": "GP / Tech Lead / Dev / QA",
        "body": (
            "Visão geral do projeto. KPIs no topo: Score OCG, Aderência "
            "(Compliance), Stacks recomendadas, Pilares avaliados. Cards "
            "centrais: 'Saúde do Contexto OCG' (gauge de Funcionalidade, "
            "Configuração, Modelagem, Implementação) e 'Consumo de IA do "
            "Projeto' (custo em USD + tokens). Cards inferiores: Gatekeeper "
            "(radar dos 7 pilares + score), Stack Recomendada (linguagens e "
            "frameworks decididos pelo Arguidor) e Equipe (membros + papéis)."
        ),
        "dica": (
            "Esta tela é o pulso do projeto. Verifique semanalmente: se o "
            "score do OCG cair, é sinal de que entrou ingestão conflitante e "
            "vale revisar o Arguidor."
        ),
    },
    "22_projeto_team": {
        "caption": "Equipe do projeto",
        "role": "GP",
        "body": (
            "Gestão dos membros do projeto. Lista quem é GP, Tech Lead, Dev, "
            "QA, Compliance, Stakeholder. Permite convidar novos membros por "
            "email com papel atribuído (link válido por 5 dias úteis), "
            "alterar papel, desativar, e visualizar quando a pessoa aceitou o "
            "convite. Substituição de GP é registrada em audit."
        ),
        "dica": (
            "Múltiplos papéis no mesmo usuário são permitidos. Restrições mais "
            "fortes prevalecem (ex: Dev+QA continua sem aprovar suas próprias "
            "entregas)."
        ),
    },
    "23_projeto_questionnaire": {
        "caption": "Questionário inicial",
        "role": "GP / Stakeholder",
        "body": (
            "Visualização read-only do questionário de 49 perguntas que o "
            "solicitante (ou time inicial) preencheu. As respostas alimentam "
            "diretamente o seed do OCG e são auditadas. Dividido em seções: "
            "Negócio, Compliance, Escopo, NFR, Arquitetura, Dados, Segurança, "
            "Operacional. Cada pergunta com tooltip explicativo e opção 'N/A' "
            "documentada."
        ),
        "dica": (
            "Se um stakeholder questionar uma decisão de stack, esta tela é a "
            "primeira parada — mostra exatamente o que foi declarado no início."
        ),
    },
    "24_projeto_repository": {
        "caption": "Repositório do projeto",
        "role": "GP",
        "body": (
            "Configuração do repositório Git oficial do projeto (obrigatório). "
            "Suporta GitHub, GitLab e Bitbucket. Captura: provider, URL do "
            "repo, Personal Access Token (com permissão read+write) e branch "
            "padrão (default 'main'). Botão 'Verificar conexão' valida e "
            "salva. Status de conexão exibido em tempo real."
        ),
        "dica": (
            "Use PAT com escopo MÍNIMO necessário (apenas o repo deste "
            "projeto). O token será usado para commit automático do CodeGen — "
            "trate como credencial sensível."
        ),
    },
    "25_projeto_external_repos": {
        "caption": "Repositórios externos",
        "role": "GP / Tech Lead",
        "body": (
            "Apontamento para repositórios de referência ou legado que o "
            "projeto pode usar como contexto. Cada repo externo é analisado "
            "automaticamente: o Arguidor lê o código, identifica padrões, "
            "stack e dependências, e gera um documento sintético que entra "
            "na ingestão alimentando o OCG. Útil para projetos de "
            "modernização ou que precisam manter compatibilidade com algo "
            "existente."
        ),
        "dica": (
            "Um repo externo precisa ser aprovado pelo GP antes de virar "
            "documento ingerido — etapa de governança contra ingestão "
            "indesejada."
        ),
    },
    "26_projeto_ingestion": {
        "caption": "Ingestão de Documentos",
        "role": "GP / Tech Lead / Compliance",
        "body": (
            "Upload e gestão de documentos que viram contexto para o OCG. "
            "Aceita PDF, DOCX e Markdown. Drag-and-drop ou botão 'Selecionar "
            "Arquivos'. A tabela mostra cada documento com tipo, tamanho, "
            "status (pending → analyzing → processed) e categorização. "
            "Documentos vindos de repos externos aparecem com o prefixo "
            "'external_*'. Deduplicação via SHA-256: subir o mesmo arquivo "
            "duas vezes é ignorado."
        ),
        "dica": (
            "Documentos ricos em requisitos, regulamentações e decisões "
            "passadas alimentam um OCG forte. Documentos genéricos contraem o "
            "OCG (sinalizam baixa confiança aos agentes)."
        ),
    },
    "27_projeto_gatekeeper": {
        "caption": "Gatekeeper — 7 Pilares",
        "role": "Compliance / GP",
        "body": (
            "Avaliação dos 7 pilares com gauge global central, radar visual "
            "à direita e lista detalhada de cada pilar abaixo. Cada pilar "
            "tem score numérico, status (aprovado, atenção, bloqueante) e "
            "lista de itens identificados. Itens com severidade BLOCKER ou "
            "CRITICAL impedem o pipeline de avançar. O botão 'Gerar Código' "
            "só fica habilitado quando todos os bloqueios estão resolvidos."
        ),
        "dica": (
            "Score < 70% em P2 (Compliance) ou P7 (Segurança) sempre bloqueia "
            "— intencional, regra de negócio. Para outros pilares, o Admin "
            "pode parametrizar thresholds em Admin → Configurações."
        ),
    },
    "28_projeto_ocg": {
        "caption": "OCG — Objeto de Contexto Global",
        "role": "GP / Tech Lead / Compliance",
        "body": (
            "Visualização completa do OCG em sua versão atual. Mostra: scores "
            "dos pilares, recomendações de stack, decisões arquiteturais, "
            "compliance mapeado, estratégia de testes, deliverables planejados. "
            "Histórico de versões disponível com possibilidade de rollback. "
            "Cada mudança no OCG (delta) registra origem (document_ingestion, "
            "manual_edit, pillar_agent, propagation, rollback) e quem fez."
        ),
        "dica": (
            "O OCG não é editado diretamente em texto. Mudanças vêm de "
            "ingestão (documentos), Arguidor (recálculo automático) ou "
            "edições parametrizadas via UI. Para reverter algo, use o "
            "rollback de versão."
        ),
    },
    "29_projeto_arguider": {
        "caption": "Arguidor",
        "role": "Tech Lead",
        "body": (
            "Módulo que ajusta arquitetura, stack e padrões com base no OCG e "
            "nos repos externos analisados. Mostra recomendações de stack "
            "(linguagens, frameworks, banco de dados, infraestrutura) com "
            "justificativa para cada escolha, dependências sugeridas com "
            "versões, e padrões de design recomendados. Aceitar/rejeitar "
            "recomendações alimenta o OCG."
        ),
        "dica": (
            "O Arguidor reroda automaticamente quando entra novo documento "
            "ou repo externo. Suas decisões são versionadas — você pode "
            "comparar 'antes vs depois' de uma mudança."
        ),
    },
    "30_projeto_backlog": {
        "caption": "Backlog Vivo",
        "role": "GP / Tech Lead",
        "body": (
            "Lista de itens derivados automaticamente do OCG. Cada item "
            "mostra: título, descrição, pilar relacionado (Negócio, "
            "Compliance, etc), prioridade (P0-P3), tags (Compliance, "
            "Segurança, Performance) e status (pending, blocked, in_progress, "
            "done). Filtros por pilar e status. Botões 'Gerar Backlog "
            "Inteligente' e 'Regenerar do OCG' atualizam a lista quando "
            "houve mudança no contexto."
        ),
        "dica": (
            "Itens marcados 'blocked' geralmente dependem de ingestão "
            "adicional ou aprovação de outro item. O badge 'Compliance' marca "
            "itens auditáveis para LGPD ou regulamentação setorial."
        ),
    },
    "31_projeto_roadmap": {
        "caption": "Roadmap",
        "role": "GP",
        "body": (
            "Sequenciamento dos itens do backlog em sprints/marcos. "
            "Visualização tipo timeline com dependências, datas estimadas, "
            "responsáveis sugeridos. Drag-and-drop para reordenar. Cada item "
            "do roadmap aciona o CodeGen quando estiver pronto."
        ),
        "dica": (
            "O roadmap é uma proposta — cabe ao GP ajustar sprint, "
            "responsável e prazo. As estimativas são geradas com base em "
            "complexidade declarada no OCG."
        ),
    },
    "32_projeto_codegen": {
        "caption": "CodeGen — Gerador de Código",
        "role": "Dev / Tech Lead",
        "body": (
            "Tela principal de geração de código. À esquerda, painel com "
            "ações ('Gerar Código', 'Avisos', 'Recompiar Reset Cobertura', "
            "'CD Pendente'). À direita, árvore de arquivos do repositório "
            "(estrutura completa do projeto). No centro, editor Monaco para "
            "visualização e edição. Quando você seleciona um item do roadmap "
            "e clica 'Gerar Código', o sistema invoca o LLM com o contexto "
            "OCG e produz os arquivos correspondentes. Validação automática "
            "(Pyflakes para Python, esprima para JS/TS) marca erros e "
            "warnings inline."
        ),
        "dica": (
            "Use 'Regerar arquivo' (não 'Regerar projeto') para corrigir um "
            "arquivo pontual sem perder edições no resto. Após aprovar, o "
            "GCA commita os arquivos diretamente no Git do projeto."
        ),
    },
    "33_projeto_qa_readiness": {
        "caption": "QA Readiness — Planos de Teste",
        "role": "QA",
        "body": (
            "Geração e revisão de planos de teste a partir do OCG. Cada item "
            "do backlog gera casos de teste sugeridos (unitário, integração, "
            "E2E) com cenários, dados de massa e critérios de aceite. O QA "
            "revisa, edita e aprova os planos. Estatísticas no topo: cobertura "
            "estimada, casos pendentes, casos aprovados."
        ),
        "dica": (
            "O QA é o único papel que pode editar planos de teste. Devs podem "
            "executar mas não alterar — separação de responsabilidades para "
            "evitar 'auto-validação'."
        ),
    },
    "34_projeto_tester_review": {
        "caption": "Revisão do Tester",
        "role": "QA / Tester",
        "body": (
            "Execução dos planos de teste e registro de evidências. Cada "
            "execução grava: data, executor, status (passou/falhou), tempo, "
            "logs e artefatos (screenshots, JSONL de output). 6 abas "
            "organizam: Pendentes, Em Execução, Passou, Falhou, Bloqueados, "
            "Histórico. Aprovações destravam itens correspondentes no roadmap."
        ),
        "dica": (
            "Falhas registradas viram automaticamente itens do backlog com "
            "tag 'bug', preservando rastreabilidade entre teste e correção."
        ),
    },
    "35_projeto_docs": {
        "caption": "Documentação Viva",
        "role": "Todos",
        "body": (
            "Visualização da documentação do projeto, regenerada "
            "automaticamente a cada commit. Árvore lateral mostra todos os "
            "arquivos .md em /docs do repositório. Visualização read-only — "
            "para editar, use seu IDE preferido e faça commit (a doc se "
            "atualiza sozinha). Cada documento mostra autor, data e link "
            "para o commit que originou a versão atual."
        ),
        "dica": (
            "Compartilhe esta tela com stakeholders que precisam acompanhar "
            "sem mexer no código. É a versão sempre atual da documentação, "
            "sem precisar baixar repo."
        ),
    },
    "36_projeto_readiness": {
        "caption": "Definition of Done — Readiness",
        "role": "GP / Compliance",
        "body": (
            "Painel central da entrega. Mostra todos os deliverables esperados "
            "(código, OCG, ADRs, SBOM, plano de testes, diagrama de "
            "arquitetura, compliance doc, dockerfile, CI pipeline, etc) com "
            "status: Verificado (verde), Presente (azul), Manual (amarelo) ou "
            "Faltando (vermelho). Filtros por pilar e por status. Indicador "
            "de readiness percentual no topo. Quando todos os deliverables "
            "críticos estão 'verified', o botão 'Gerar Release Bundle' fica "
            "habilitado."
        ),
        "dica": (
            "Deliverables 'manual_only' precisam ser anexados manualmente "
            "(business case, decisões registradas em PDF, etc). O sistema "
            "não tenta gerar — apenas reserva o slot e cobra a presença."
        ),
    },
    "37_projeto_settings": {
        "caption": "Configurações do projeto",
        "role": "GP",
        "body": (
            "Parametrização específica do projeto: provedor de IA preferido "
            "(pode sobrescrever o default da instância), modelos para cada "
            "agente (Analyzer, P1-P7, Consolidator, Arguidor, CodeGen), "
            "thresholds dos pilares (se a instância permitir override), "
            "configurações de notificação por evento, e parâmetros de "
            "geração (idioma do código, padrão de docstring, formato de ADR)."
        ),
        "dica": (
            "Mexa aqui só quando souber o que está fazendo. Os defaults da "
            "instância já são otimizados pelo Admin. Mudança em modelo de "
            "agente reflete na próxima execução, não retroativamente."
        ),
    },
    "38_projeto_audit": {
        "caption": "Auditoria do projeto",
        "role": "GP / Compliance",
        "body": (
            "Log local do projeto (separado da auditoria global do sistema). "
            "Registra todas as ações dentro do projeto: ingestão de documento, "
            "atualização do OCG (com correlation_id), geração de código, "
            "aprovações de QA, mudanças de papel, edições de configuração. "
            "Filtros por usuário, evento e período. Exportável para CSV — "
            "útil para auditoria externa ou compliance check."
        ),
        "dica": (
            "Compliance pode pedir esta tela como evidência em auditorias "
            "regulatórias. O log é imutável e cada ação tem ator identificado."
        ),
    },
}


# ──────────────────────────────────────────────────────────────────
# Helpers de formatação
# ──────────────────────────────────────────────────────────────────

def set_default_font(doc: Document, family: str = "Calibri", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)
    style.font.color.rgb = SLATE_DARK


def add_h1(doc: Document, text: str, color: RGBColor = VIOLET) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = color


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = VIOLET


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = SLATE_DARK


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_callout(doc: Document, label: str, text: str, color: RGBColor = AMBER) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    label_run = p.add_run(f"[{label}] ")
    label_run.bold = True
    label_run.font.color.rgb = color
    body_run = p.add_run(text)
    body_run.font.color.rgb = SLATE_DARK


def add_qa(doc: Document, question: str, answer: str) -> None:
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(8)
    qrun = pq.add_run(f"P: {question}")
    qrun.bold = True
    qrun.font.color.rgb = VIOLET
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    pa.paragraph_format.space_after = Pt(10)
    arun = pa.add_run(f"R: {answer}")
    arun.font.color.rgb = SLATE_DARK


def add_table_two_cols(doc: Document, header: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid Accent 4"
    hdr = table.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    for i, (k, v) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v


def add_screenshot_section(doc: Document, shot: dict, desc: dict) -> None:
    """Insere um bloco de tela: título, badge de papel, screenshot, descrição, dica."""
    # Título da tela
    add_h3(doc, f"{shot['n']:02d}. {desc.get('caption', shot['slug'])}")

    # Badge de papel
    badge = doc.add_paragraph()
    badge.paragraph_format.space_after = Pt(4)
    badge_run = badge.add_run(f"Quem usa: {desc.get('role', '—')}")
    badge_run.italic = True
    badge_run.font.size = Pt(9)
    badge_run.font.color.rgb = CYAN

    # URL real
    if shot.get("url"):
        url_p = doc.add_paragraph()
        url_p.paragraph_format.space_after = Pt(8)
        url_run = url_p.add_run(f"URL: {shot['url']}")
        url_run.font.size = Pt(9)
        url_run.font.color.rgb = SLATE_MID
        url_run.italic = True

    # Screenshot
    img_path = SCREENSHOTS / shot["file"]
    if img_path.exists():
        try:
            doc.add_picture(str(img_path), width=Cm(IMG_WIDTH_CM))
            # Centraliza imagem
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            add_callout(doc, "ERRO", f"Não foi possível inserir imagem {shot['file']}: {e}", RED)
    else:
        add_callout(doc, "AUSENTE", f"Imagem não encontrada: {shot['file']}", RED)

    # Descrição
    add_p(doc, desc.get("body", "Descrição pendente."))

    # Dica (se houver)
    if desc.get("dica"):
        add_callout(doc, "Dica", desc["dica"], EMERALD)

    # Espaço antes da próxima
    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────────
# Construção do documento
# ──────────────────────────────────────────────────────────────────

def add_capa(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("\n\n\nGCA")
    t.bold = True
    t.font.size = Pt(56)
    t.font.color.rgb = VIOLET

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Gerenciador Central de Arquiteturas")
    s.italic = True
    s.font.size = Pt(20)
    s.font.color.rgb = SLATE_MID

    doc.add_paragraph()
    doc.add_paragraph()

    pitch = doc.add_paragraph()
    pitch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = pitch.add_run("Manual Completo do Usuário")
    p_run.bold = True
    p_run.font.size = Pt(22)
    p_run.font.color.rgb = SLATE_DARK

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = sub2.add_run("Tour visual + tutorial + Q&A — para Admin e Gerente de Projeto")
    s2.font.size = Pt(13)
    s2.font.color.rgb = SLATE_MID

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(f"Versão do documento: {datetime.now().strftime('%Y-%m-%d')}")
    m.font.size = Pt(10)
    m.font.color.rgb = SLATE_MID

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = foot.add_run("28 telas capturadas em ambiente de produção (instância FinanceHub Pro)")
    f.font.size = Pt(9)
    f.font.color.rgb = SLATE_MID
    f.italic = True

    doc.add_page_break()


def add_sumario(doc: Document) -> None:
    add_h1(doc, "Sumário")
    add_p(doc, "")
    sumario = [
        ("Parte I", "Introdução ao GCA — visão, conceitos fundamentais"),
        ("Parte II", "Tour Visual — todas as 28 telas com descrição"),
        ("  Cap. 1", "Acesso ao sistema (4 telas públicas)"),
        ("  Cap. 2", "Área administrativa (5 telas — Admin)"),
        ("  Cap. 3", "Área do projeto (19 telas — GP/Tech Lead/Dev/QA)"),
        ("Parte III", "Tutorial Passo a Passo — 10 jornadas comuns"),
        ("Parte IV", "O que o GCA FAZ / NÃO FAZ"),
        ("Parte V", "Perguntas e Respostas (FAQ)"),
        ("Parte VI", "Glossário e Suporte"),
    ]
    for k, v in sumario:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_k = p.add_run(f"{k}  ")
        run_k.bold = True
        run_k.font.color.rgb = VIOLET
        p.add_run(v)
    doc.add_page_break()


def add_parte_intro(doc: Document) -> None:
    add_h1(doc, "Parte I — Introdução ao GCA")

    add_h2(doc, "1.1 O que é o GCA")
    add_p(doc,
        "O GCA (Gerenciador Central de Arquiteturas) é uma plataforma de "
        "governança e geração assistida de software com IA. Ele organiza o "
        "ciclo de vida completo de um projeto de TI — da solicitação inicial "
        "até a entrega de código, testes e documentação — garantindo que "
        "decisões técnicas, de negócio e de compliance sejam registradas, "
        "avaliadas e auditáveis.")
    add_p(doc,
        "Em uma frase: o GCA recebe a descrição de um projeto, conduz uma "
        "análise rigorosa em sete dimensões (os Pilares), produz um documento "
        "vivo de contexto (o OCG), gera código sob governança e mantém a "
        "documentação atualizada a cada mudança.")

    add_h2(doc, "1.2 Para quem é?")
    add_bullet(doc, "Administradores de TI que precisam orquestrar múltiplos projetos com rastreabilidade.")
    add_bullet(doc, "Gerentes de Projeto (GP) que querem governar entregas sem perder velocidade.")
    add_bullet(doc, "Times de desenvolvimento que aceitam apoio de IA mas com aprovação humana e auditoria.")
    add_bullet(doc, "Times de QA, Compliance e Stakeholders que precisam visibilidade sem editar código.")

    add_h2(doc, "1.3 Filosofia central")
    add_p(doc,
        "Tudo o que o GCA produz precisa ser explicável, reversível e "
        "aprovável por um humano. A IA acelera, mas nunca decide sozinha. "
        "Cada arquivo gerado, cada decisão de stack, cada item do backlog "
        "tem origem rastreável e responsável designado.")

    add_h2(doc, "1.4 Conceitos fundamentais")

    add_h3(doc, "OCG — Objeto de Contexto Global")
    add_p(doc,
        "O OCG é o documento vivo do seu projeto. Começa pequeno (com as "
        "respostas do questionário) e cresce a cada nova ingestão. É a fonte "
        "única de verdade que alimenta todos os outros módulos. Ingestão "
        "boa expande o OCG; ingestão ruim contrai.")

    add_h3(doc, "Os 7 Pilares")
    add_p(doc, "Toda decisão técnica passa por avaliação em sete dimensões:")
    add_table_two_cols(doc, ("Pilar", "Descrição"), [
        ("P1 — Negócio", "Modelo de negócio, ROI, KPIs, monetização."),
        ("P2 — Compliance", "LGPD, regulatório setorial (BACEN, ANS), políticas."),
        ("P3 — Escopo", "Funcional vs não-funcional, MVP, fora de escopo."),
        ("P4 — NFR", "Performance, escalabilidade, disponibilidade, SLA."),
        ("P5 — Arquitetura", "Stack, padrões, dependências, infraestrutura."),
        ("P6 — Dados", "Modelagem, ETL, governança, retenção."),
        ("P7 — Segurança", "Autenticação, criptografia, auditoria, OWASP."),
    ])
    add_callout(doc, "Bloqueante",
                "Score < 70% em P2 (Compliance) ou P7 (Segurança) trava o pipeline. "
                "Outros thresholds são parametrizáveis pelo Admin.", RED)

    add_h3(doc, "Pipeline de 8 agentes IA")
    add_p(doc,
        "Para gerar o OCG, o GCA aciona oito agentes especializados em "
        "paralelo: Analyzer (decompõe), 7 especialistas (um por pilar) e "
        "Consolidator (junta tudo). Cada agente declara confiança e gaps. "
        "Provedores suportados: Anthropic, OpenAI, Gemini, DeepSeek, Grok, Qwen.")

    add_h3(doc, "RBAC — 7 papéis")
    add_table_two_cols(doc, ("Papel", "O que faz / não faz"), [
        ("Admin", "Configura sistema. NÃO atua dentro de projetos."),
        ("GP", "Conduz o projeto, aprova entregas. NÃO escreve código."),
        ("Tech Lead", "Define arquitetura, revisa CodeGen."),
        ("Dev", "Edita código, abre PRs. NÃO aprova suas próprias entregas."),
        ("QA", "Aprova qualidade. NUNCA edita código."),
        ("Compliance", "Audita P2 e P7."),
        ("Stakeholder", "Visualiza progresso. Sem permissão de edição."),
    ])

    add_h3(doc, "Pipeline de 10 fases por projeto")
    fases = [
        "1. Questionário inicial — captura escopo, restrições, KPIs.",
        "2. Repositório do Projeto (obrigatório).",
        "3. Repos Externos — código de referência/legado.",
        "4. Ingestão — documentos viram contexto.",
        "5. Gatekeeper — avalia 7 pilares e libera/trava o pipeline.",
        "6. OCG — gerado e atualizado reativamente.",
        "7. Arguidor — ajusta arquitetura, stack, padrões.",
        "8. Backlog → Roadmap → CodeGen — geração de código sob aprovação.",
        "9. Testes — QA Readiness + Tester Review.",
        "10. Documentação Viva + Release Bundle.",
    ]
    for f in fases:
        add_bullet(doc, f)

    doc.add_page_break()


def add_parte_tour_visual(doc: Document, shots: list[dict]) -> None:
    add_h1(doc, "Parte II — Tour Visual")
    add_p(doc,
        "Esta parte percorre todas as telas do GCA com captura real, "
        "descrição de uso e dica prática. As telas estão agrupadas em "
        "três capítulos: Acesso ao sistema (públicas), Área administrativa "
        "e Área do projeto.")
    add_p(doc, "")

    # Cap 1 — públicas
    add_h2(doc, "Capítulo 1 — Acesso ao sistema")
    add_p(doc,
        "Telas que qualquer pessoa acessa antes de autenticar. Cobrem login, "
        "solicitação de novo projeto e recuperação de senha.")
    for shot in [s for s in shots if s["category"] == "publica" and s["ok"]]:
        slug = shot["file"].replace(".png", "")
        desc = DESCRICOES.get(slug, {})
        add_screenshot_section(doc, shot, desc)
    doc.add_page_break()

    # Cap 2 — admin
    add_h2(doc, "Capítulo 2 — Área administrativa (Admin)")
    add_p(doc,
        "Telas exclusivas do Admin. Por regra de negócio, o Admin nunca "
        "atua dentro de projetos — sua função é configurar a instância, "
        "aprovar solicitações, gerenciar usuários e provedores de IA, e "
        "auditar todo o sistema.")
    for shot in [s for s in shots if s["category"] == "admin" and s["ok"]]:
        slug = shot["file"].replace(".png", "")
        desc = DESCRICOES.get(slug, {})
        add_screenshot_section(doc, shot, desc)
    doc.add_page_break()

    # Cap 3 — projeto
    add_h2(doc, "Capítulo 3 — Área do projeto (GP / Tech Lead / Dev / QA)")
    add_p(doc,
        "Telas dentro de um projeto. As permissões variam por papel — "
        "todas as telas são visíveis para o GP; outros papéis veem "
        "subconjuntos relevantes (Dev vê CodeGen, QA vê QA Readiness "
        "e Tester Review, etc).")
    for shot in [s for s in shots if s["category"] == "projeto" and s["ok"]]:
        slug = shot["file"].replace(".png", "")
        desc = DESCRICOES.get(slug, {})
        add_screenshot_section(doc, shot, desc)
    doc.add_page_break()


def add_parte_tutorial(doc: Document) -> None:
    add_h1(doc, "Parte III — Tutorial Passo a Passo")
    add_p(doc,
        "Dez jornadas comuns desenhadas como sequências de passos.")
    add_p(doc, "")

    add_h2(doc, "3.1 Solicitar um novo projeto")
    add_numbered(doc, "Acesse a página de login do GCA.")
    add_numbered(doc, "Clique em 'Solicitar novo projeto' (rodapé do card de login).")
    add_numbered(doc, "Preencha o passo 1: nome, email, nome do projeto, tipo, descrição (mín. 30 chars).")
    add_numbered(doc, "Se nenhum tipo padrão couber, escolha 'Outro' e descreva.")
    add_numbered(doc, "No passo 2, responda perguntas obrigatórias do tipo escolhido.")
    add_numbered(doc, "Envie. O administrador é notificado in-app e por email.")
    add_numbered(doc, "Você receberá email de aprovação ou recusa no endereço cadastrado.")
    add_callout(doc, "Atenção",
        "Quem solicita vira automaticamente o GP quando aprovado. Use seu "
        "email principal — é o canal oficial.")

    add_h2(doc, "3.2 Login (Admin OU membro de projeto)")
    add_bullet(doc, "Sem projeto selecionado: só admins entram (área administrativa).")
    add_bullet(doc, "Com projeto no combo: valida que você é membro e aplica permissões do papel.")
    add_p(doc, "Passos:")
    add_numbered(doc, "Selecione projeto no combo (último acessado fica no topo).")
    add_numbered(doc, "Digite email e senha.")
    add_numbered(doc, "Clique em Entrar.")
    add_callout(doc, "Erros comuns",
        "'Email ou senha inválidos' = credenciais erradas. 'Você não é membro' = "
        "GP precisa adicionar você ao time. 'Selecione seu projeto' = "
        "não-admin tentando entrar sem projeto.")

    add_h2(doc, "3.3 Aprovar solicitação (Admin)")
    add_numbered(doc, "Em Admin → Gestão de Projetos, localize linha 'Pendente'.")
    add_numbered(doc, "Clique no ícone de documento para abrir 'Detalhes da solicitação'.")
    add_numbered(doc, "Avalie aderência (descrição, perguntas e respostas do wizard).")
    add_numbered(doc, "Para aprovar: ícone verde de check. Tenant é provisionado, GP vira membro, 2 emails enviados.")
    add_numbered(doc, "Para rejeitar: ícone âmbar de aviso. Escreva motivo (mín. 10 chars) — vai por email.")
    add_numbered(doc, "Para spam/duplicata: lixeira vermelha (exclui sem notificar).")

    add_h2(doc, "3.4 Configurar repositório (GP)")
    add_p(doc, "Sem repositório, o pipeline não avança. Suporta GitHub, GitLab e Bitbucket.")
    add_numbered(doc, "Entre no projeto e vá em 'Repositório'.")
    add_numbered(doc, "Escolha provider e cole URL do repo.")
    add_numbered(doc, "Cole PAT com permissão de leitura e escrita.")
    add_numbered(doc, "Clique 'Verificar conexão'. Repo é vinculado e branch padrão detectada.")

    add_h2(doc, "3.5 Ingerir documentos para o OCG")
    add_numbered(doc, "Vá em Ingestão.")
    add_numbered(doc, "Faça upload de PDFs/DOCX/MD com requisitos, regulamentações, atas.")
    add_numbered(doc, "Sistema indexa, calcula SHA-256 (deduplicação) e analisa.")
    add_numbered(doc, "Acompanhe status: pending → analyzing → processed.")
    add_numbered(doc, "Documentos processados alimentam o OCG automaticamente.")
    add_callout(doc, "Repos externos",
        "Em 'Repos Externos' aponte para repositórios existentes. O Arguidor "
        "analisa o código e gera documento sintético que entra na ingestão.")

    add_h2(doc, "3.6 Acompanhar Gatekeeper")
    add_numbered(doc, "Vá em Gatekeeper.")
    add_numbered(doc, "Veja score de cada um dos 7 pilares + radar visual.")
    add_numbered(doc, "Itens BLOCKER ou CRITICAL precisam ser resolvidos antes do pipeline avançar.")
    add_numbered(doc, "Para cada item: marcar resolvido, anexar evidência ou registrar justificativa.")

    add_h2(doc, "3.7 Gerar código (CodeGen)")
    add_numbered(doc, "Após OCG verde, vá em Backlog → Roadmap.")
    add_numbered(doc, "Selecione item do roadmap e dispare CodeGen.")
    add_numbered(doc, "Aguarde geração (barra de progresso visível).")
    add_numbered(doc, "Revise arquivos no editor Monaco — diagnostics destacam erros.")
    add_numbered(doc, "Para regerar arquivo único, use 'Regerar arquivo' (não 'Regerar projeto').")
    add_numbered(doc, "Ao aprovar, GCA commita os arquivos no Git do projeto.")
    add_callout(doc, "Docstrings obrigatórias",
        "Política GCA: todo módulo/classe/função em código gerado precisa ter "
        "docstring. Arquivos sem ficam em status 'todo' automaticamente.")

    add_h2(doc, "3.8 Planejar e executar testes")
    add_numbered(doc, "QA Readiness: GCA gera planos a partir do OCG.")
    add_numbered(doc, "QA revisa, aprova ou edita o plano (devs não editam).")
    add_numbered(doc, "Tester executa testes. GCA registra evidências (logs, artefatos).")
    add_numbered(doc, "QA aprova ou rejeita execução. Aprovações destravam roadmap.")

    add_h2(doc, "3.9 Documentação Viva")
    add_numbered(doc, "Vá em Documentação Viva.")
    add_numbered(doc, "Árvore mostra todos .md em /docs do repo.")
    add_numbered(doc, "Clique para visualizar (read-only — para editar use IDE e commit).")
    add_numbered(doc, "Documentação se atualiza sozinha a cada commit.")

    add_h2(doc, "3.10 Empacotar entrega (Release Bundle)")
    add_numbered(doc, "Em Readiness, quando deliverables estiverem 'verified', clique 'Gerar Release Bundle'.")
    add_numbered(doc, "GCA empacota: código, OCG, docs, ADRs, SBOM, plano de testes, métricas.")
    add_numbered(doc, "Você recebe .zip versionado com MANIFEST.json + RELEASE_NOTES.md + SHA-256.")
    add_numbered(doc, "Esse bundle é o entregável oficial para o cliente / produção.")

    doc.add_page_break()


def add_parte_faz_nao_faz(doc: Document) -> None:
    add_h1(doc, "Parte IV — O que o GCA FAZ / NÃO FAZ")

    add_h2(doc, "4.1 O que o GCA FAZ")

    add_h3(doc, "Governança de projeto")
    add_bullet(doc, "Recebe e aprova solicitações com wizard tipado e perguntas obrigatórias.")
    add_bullet(doc, "Provisiona schema PostgreSQL isolado por projeto (multi-tenancy real).")
    add_bullet(doc, "RBAC granular com 7 papéis e auditoria de toda ação.")
    add_bullet(doc, "Substituição de GP, transferência de papéis, bloqueio de exclusão de usuário ativo.")

    add_h3(doc, "Análise e contexto (OCG)")
    add_bullet(doc, "Constrói e atualiza OCG reativamente conforme você adiciona subsídios.")
    add_bullet(doc, "Aplica deltas (mudanças incrementais) ao invés de reescrever — rápido e auditável.")
    add_bullet(doc, "Histórico completo de versões com rollback.")
    add_bullet(doc, "Avalia 7 pilares e bloqueia pipeline quando Segurança/Compliance < 70%.")

    add_h3(doc, "Geração assistida")
    add_bullet(doc, "Gera código (scaffold ou arquivo único) com IA, validado e com docstrings.")
    add_bullet(doc, "Commita arquivos aprovados direto no Git do projeto.")
    add_bullet(doc, "Gera 9+ artefatos: ADR, OpenAPI, Dockerfile, CI pipeline, observability, SBOM, plano de testes, diagrama, compliance doc.")
    add_bullet(doc, "Detecta gaps e cobra novos subsídios quando agentes têm baixa confiança.")

    add_h3(doc, "Integração e infraestrutura")
    add_bullet(doc, "Integra com GitHub, GitLab e Bitbucket via PAT.")
    add_bullet(doc, "Suporta seis provedores de IA: Anthropic, OpenAI, Gemini, DeepSeek, Grok, Qwen.")
    add_bullet(doc, "Parametriza pesos de pilares, thresholds e modelos por instância.")
    add_bullet(doc, "Notificações in-app + email para todos os eventos relevantes.")

    add_h3(doc, "Qualidade e entrega")
    add_bullet(doc, "Pipeline de QA: planos gerados, revisão humana, execução com evidências.")
    add_bullet(doc, "Validador tier-1: Pyflakes (Python), esprima (JS/TS), stdlib (JSON/YAML/TOML).")
    add_bullet(doc, "Editor Monaco com diagnostics em tempo real.")
    add_bullet(doc, "Release Bundle: .zip versionado com manifest e changelog.")
    add_bullet(doc, "Documentação viva regenerada a cada commit.")

    add_h2(doc, "4.2 O que o GCA NÃO FAZ")

    add_h3(doc, "Limites por design")
    add_bullet(doc, "Não substitui decisão humana. IA propõe, humano aprova.")
    add_bullet(doc, "Não escreve código sem repositório Git configurado (bloqueio explícito).")
    add_bullet(doc, "Não compartilha dados entre projetos (schema isolado, RLS).")
    add_bullet(doc, "Admin nunca atua dentro de projetos. QA nunca edita código. Dev nunca aprova suas próprias entregas.")
    add_bullet(doc, "Não exclui usuários que são GP de projetos ativos sem transferência prévia.")
    add_bullet(doc, "Não aceita aprovação de projetos com Segurança/Compliance < 70% sem justificativa.")

    add_h3(doc, "Não é um IDE")
    add_bullet(doc, "Edição de código no GCA é leve (Monaco read-mostly).")
    add_bullet(doc, "Não roda debugger nem testes na infra dele — apenas registra que foram executados.")
    add_bullet(doc, "Não substitui pipelines CI/CD existentes — gera o YAML mas execução é externa.")

    add_h3(doc, "Não é um SaaS multi-cliente")
    add_bullet(doc, "GCA é produto instalável. Cada cliente tem sua própria instância.")
    add_bullet(doc, "Não há tenant compartilhado, billing centralizado nem marketplace ainda.")

    add_h3(doc, "Gaps conhecidos (no roadmap)")
    add_bullet(doc, "Instalador único (binário/installer GUI) — hoje só Docker Compose.")
    add_bullet(doc, "Auto-upgrade com preview de changes — hoje upgrade manual.")
    add_bullet(doc, "Backup/restore com export/import de configuração.")
    add_bullet(doc, "Hardening de produção (rate limit, WAF, certificados).")
    add_bullet(doc, "PAT do Git criptografado com Fernet.")
    add_bullet(doc, "Marketplace de templates de projeto.")

    add_h3(doc, "Não é mágico")
    add_bullet(doc, "Descrição vaga = OCG raso = Gatekeeper cobra mais ingestão.")
    add_bullet(doc, "Sem documentos/regulamentações/contratos = pilares Negócio e Compliance amarelos.")
    add_bullet(doc, "Aprovar geração sem revisar = retrabalho depois.")
    add_bullet(doc, "Ignorar bloqueios do Gatekeeper = pipeline quebra mais adiante.")

    doc.add_page_break()


def add_parte_faq(doc: Document) -> None:
    add_h1(doc, "Parte V — Perguntas e Respostas")

    add_h2(doc, "Sobre o produto")
    add_qa(doc, "Por que GCA e não outro framework de IA?",
        "O GCA não é só um wrapper de LLM. Orquestra papéis humanos, governa "
        "decisões com 7 pilares, mantém auditoria completa. A IA é meio; o fim é "
        "entrega rastreável.")
    add_qa(doc, "O GCA precisa de internet?",
        "Sim — para chamar provedores de IA e Git remoto. Nada do código do projeto "
        "sai da sua instância; saem prompts construídos com contexto necessário "
        "(você controla o provedor).")
    add_qa(doc, "Posso usar modelos locais (Ollama, LM Studio)?",
        "Hoje suportamos Anthropic, OpenAI, Gemini, DeepSeek, Grok, Qwen. Modelos "
        "locais via API compatível com OpenAI funcionam (configurar endpoint custom). "
        "Suporte oficial está no roadmap.")
    add_qa(doc, "É open source?",
        "A instância atual é dogfood do criador. Política de licenciamento da v1.0 "
        "será definida antes do release público.")

    add_h2(doc, "Sobre conta e acesso")
    add_qa(doc, "Esqueci minha senha. O que fazer?",
        "Tela de login → 'Esqueci minha senha'. Recebe email com link válido por 24h.")
    add_qa(doc, "Posso ter múltiplos papéis em um mesmo projeto?",
        "Sim. Permissões somam, restrições mais fortes prevalecem (Dev+QA ainda não "
        "aprova sua própria entrega).")
    add_qa(doc, "Como adiciono pessoas ao meu projeto?",
        "Como GP: aba 'Equipe' (dentro do projeto). Convide por email com papel. "
        "Link de aceite válido por 5 dias úteis.")

    add_h2(doc, "Sobre o pipeline")
    add_qa(doc, "Quanto tempo leva para o OCG ficar pronto?",
        "Depende de volume de ingestão e provedor. Com Claude Sonnet 4.6 e ingestão "
        "moderada (5–10 docs), 30s a 3 min. Atualizações reativas (delta): 5–15s.")
    add_qa(doc, "O GCA pode regenerar o OCG do zero?",
        "Sim, em casos excepcionais (Admin → Configurações). Cuidado: descarta deltas "
        "manuais. Use só após mudança estrutural (novo escopo, regulamentação).")
    add_qa(doc, "O que acontece se eu rejeitar uma sugestão da IA?",
        "Decisão fica registrada com motivo. GCA aprende do contexto — sugestões "
        "futuras consideram seus rejeições passados (via OCG).")
    add_qa(doc, "O CodeGen sobrescreve meus arquivos editados?",
        "Depende: 'Regerar projeto' (scaffold) sobrescreve tudo — só na primeira "
        "geração. 'Regerar arquivo' regera um arquivo apenas — preserva o resto. "
        "Sempre revise diff no Monaco antes de commitar.")

    add_h2(doc, "Sobre dados e segurança")
    add_qa(doc, "Onde ficam meus documentos ingeridos?",
        "Volume de storage da sua instância (backend/storage/ingested/<project_id>/). "
        "Não saem do servidor. Apenas o necessário ao prompt vai ao provedor escolhido.")
    add_qa(doc, "Senhas em texto plano?",
        "Não. Bcrypt antes de gravar. Tokens de redefinição expiram em 24h. "
        "JWT assinado com RS256.")
    add_qa(doc, "E os PATs do Git?",
        "Hoje em texto plano no campo 'pat_encrypted' (nome herdado). Roadmap: "
        "Fernet (criptografia simétrica). Use PATs com escopo mínimo até lá.")
    add_qa(doc, "Tem auditoria?",
        "Sim. Toda ação relevante gera entrada em audit_log_global ou audit local "
        "do tenant. Visível em Admin → Auditoria Global e em cada projeto.")

    add_h2(doc, "Sobre desenvolvimento e código gerado")
    add_qa(doc, "Em quais linguagens o CodeGen funciona?",
        "Suportadas: Python, JS/TS, Java, Go. Validador cobre Python (Pyflakes), "
        "JS/TS (esprima), JSON/YAML/TOML (stdlib). Outras linguagens são geradas "
        "sem validação automática.")
    add_qa(doc, "Posso usar meu próprio template?",
        "Hoje scaffold é determinado por OCG e Arguidor. Templates customizados "
        "estão no roadmap (marketplace futuro). Workaround: faça scaffold e adapte; "
        "GCA respeita edições nos arquivos individuais.")
    add_qa(doc, "Por que docstrings obrigatórias?",
        "Política do GCA: código sem documentação inline é dívida. Validador rebaixa "
        "arquivos sem docstring para 'todo'. Isenções: configs, __init__.py "
        "triviais, markdown.")

    add_h2(doc, "Sobre a entrega final")
    add_qa(doc, "O que está dentro de um Release Bundle?",
        "Arquivo .zip versionado: código (snapshot do repo), OCG completo, doc viva "
        "(markdown), ADRs, SBOM CycloneDX 1.5, plano de testes, diagrama de "
        "arquitetura, MANIFEST.json com deliverables, RELEASE_NOTES.md e SHA-256 "
        "para integridade.")
    add_qa(doc, "Posso entregar o Release Bundle direto pro cliente?",
        "Sim — esse é o propósito. Bundle é auto-contido e auditável. Cliente recebe "
        "tudo que precisa para entender, executar, manter e auditar a entrega.")

    doc.add_page_break()


def add_parte_glossario(doc: Document) -> None:
    add_h1(doc, "Parte VI — Glossário e Suporte")

    add_h2(doc, "6.1 Glossário")
    add_table_two_cols(doc, ("Termo", "Definição"), [
        ("OCG", "Objeto de Contexto Global. Documento vivo único por projeto, alimentado pela ingestão e usado por todos os módulos."),
        ("Pilares", "As 7 dimensões avaliadas pelo Gatekeeper: Negócio, Compliance, Escopo, NFR, Arquitetura, Dados, Segurança."),
        ("Gatekeeper", "Módulo que avalia e bloqueia ou libera o pipeline com base no score dos pilares."),
        ("Arguidor", "Módulo que ajusta arquitetura, stack e padrões com base no OCG e em repos externos."),
        ("CodeGen", "Geração de código assistida por IA, com validação automática e commit no Git."),
        ("RBAC", "Role-Based Access Control. Sistema de papéis e permissões."),
        ("Tenant", "Schema PostgreSQL isolado por projeto. Garante separação de dados."),
        ("PAT", "Personal Access Token. Token usado para autenticar com GitHub/GitLab/Bitbucket."),
        ("SBOM", "Software Bill of Materials. Lista de dependências e versões. Padrão CycloneDX 1.5."),
        ("ADR", "Architecture Decision Record. Documento curto registrando uma decisão arquitetural. Padrão MADR."),
        ("Definition of Done", "Conjunto de deliverables que precisam estar 'verified' para o projeto ser considerado pronto."),
        ("Release Bundle", "Empacotamento final em .zip com código + OCG + docs + ADRs + SBOM + manifest + changelog."),
        ("Delta OCG", "Mudança incremental aplicada ao OCG (ao invés de reescrever todo o documento). Mais rápido e auditável."),
        ("Dogfood", "Uso do próprio produto pela equipe que o desenvolve para validá-lo na prática."),
    ])

    add_h2(doc, "6.2 Suporte")
    add_h3(doc, "Onde pedir ajuda")
    add_bullet(doc, "Dentro do GCA: tooltips de ajuda em ícones de '?' nas telas.")
    add_bullet(doc, "Documentação viva: cada projeto mantém em /docs material atualizado.")
    add_bullet(doc, "Suporte ao admin da sua instância: ele pode acessar logs, auditoria e parametrizações.")

    add_h3(doc, "Reportar problema")
    add_p(doc,
        "Ao reportar um erro, sempre inclua: (1) tela onde ocorreu, (2) ação "
        "executada, (3) mensagem completa de erro, (4) horário aproximado. "
        "Permite ao admin localizar o evento na auditoria rapidamente.")

    add_h3(doc, "Versão do GCA")
    add_p(doc,
        "Versão da instância aparece no rodapé da tela de login (ex: 'v0.8.0'). "
        "Use sempre que reportar bugs ou comparar comportamentos. Em release "
        "futuro, haverá tela 'Sobre' com detecção de update disponível.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final = p.add_run("— Fim do Manual Completo —")
    final.italic = True
    final.font.color.rgb = SLATE_MID


# ──────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────

def build_document(shots: list[dict]) -> Document:
    doc = Document()
    set_default_font(doc)

    # Margens A4
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_capa(doc)
    add_sumario(doc)
    add_parte_intro(doc)
    add_parte_tour_visual(doc, shots)
    add_parte_tutorial(doc)
    add_parte_faz_nao_faz(doc)
    add_parte_faq(doc)
    add_parte_glossario(doc)

    return doc


def main() -> None:
    if not MANIFEST.exists():
        raise SystemExit(f"Manifest não encontrado: {MANIFEST}. Rode antes: bash scripts/capturar_telas_gca.sh")

    manifest_data = json.loads(MANIFEST.read_text())
    shots = manifest_data["shots"]
    print(f"Lidos {len(shots)} shots do manifest ({manifest_data['ok']} ok, {manifest_data['failed']} falhas)")

    doc = build_document(shots)
    doc.save(str(OUTPUT))

    size_kb = OUTPUT.stat().st_size / 1024
    print(f"\n✓ Manual gerado: {OUTPUT}")
    print(f"  Tamanho: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
