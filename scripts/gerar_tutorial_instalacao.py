#!/usr/bin/env python3
"""
Gera GCA_Tutorial_Instalacao_v1.docx — tutorial passo a passo de instalação
em Windows e Ubuntu, com wireframes dos 10 passos e screenshots capturados
da aplicação em funcionamento.

Autor: Luiz Carlos Pielak
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

WIREFRAME_DIR = Path("/home/luiz/GCA/docs/wireframes")
SCREENSHOT_DIR = Path("/home/luiz/GCA/screenshots_v3")
OUT_PATH = Path("/home/luiz/GCA/docs/GCA_Tutorial_Instalacao_v1.docx")

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
SLATE_DARK = RGBColor(0x1E, 0x29, 0x3B)
SLATE_MEDIUM = RGBColor(0x64, 0x74, 0x8B)
EMERALD = RGBColor(0x05, 0x96, 0x69)
AMBER = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = VIOLET
        run.font.size = Pt(20)


def h2(doc, text: str):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = SLATE_DARK
        run.font.size = Pt(15)


def h3(doc, text: str):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = SLATE_MEDIUM
        run.font.size = Pt(12)


def para(doc, text: str, *, size: int = 11, bold: bool = False,
         color: RGBColor = None, justify: bool = True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color


def bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


def numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


def code(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = SLATE_DARK


def callout(doc, title: str, text: str, color: RGBColor = AMBER):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(title + " ")
    r1.bold = True
    r1.font.color.rgb = color
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)


def image(doc, path: Path, caption: str = "", width_in: float = 6.0):
    if not path.exists():
        para(doc, f"[imagem ausente: {path.name}]", color=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = SLATE_MEDIUM


def table(doc, headers: list[str], rows: list[list[str]], widths: list[float] = None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 4"
    t.autofit = False
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "6D28D9")
        if widths and i < len(widths):
            cell.width = Cm(widths[i])
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if widths and ci < len(widths):
                cell.width = Cm(widths[ci])


def page_break(doc):
    doc.add_page_break()


# ─── Documento ────────────────────────────────────────────────────────────

def build(doc: Document):
    core = doc.core_properties
    core.title = "GCA — Tutorial de Instalação"
    core.author = "Luiz Carlos Pielak"
    core.subject = "Passo a passo de instalação do GCA em Windows e Ubuntu"
    core.keywords = "GCA, instalação, Windows, Ubuntu, Docker, Inno Setup"
    core.last_modified_by = "Luiz Carlos Pielak"
    core.created = datetime.now()

    # Capa
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("GCA")
    tr.font.size = Pt(60)
    tr.bold = True
    tr.font.color.rgb = VIOLET

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Gestão de Codificação Assistida")
    sr.font.size = Pt(20)
    sr.font.color.rgb = SLATE_DARK

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = desc.add_run("Tutorial de Instalação")
    dr.font.size = Pt(22)
    dr.italic = True
    dr.font.color.rgb = SLATE_MEDIUM

    for _ in range(6):
        doc.add_paragraph()

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sub2.add_run("Para usuários técnicos nível pleno")
    sr2.font.size = Pt(14)
    sr2.font.color.rgb = SLATE_MEDIUM

    for _ in range(3):
        doc.add_paragraph()

    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = auth.add_run("Autor: Luiz Carlos Pielak")
    ar.font.size = Pt(13)
    ar.bold = True

    dateline = doc.add_paragraph()
    dateline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt = dateline.add_run(f"Versão 1.0 — {datetime.now().strftime('%d/%m/%Y')}")
    dt.font.size = Pt(11)
    dt.font.color.rgb = SLATE_MEDIUM

    page_break(doc)

    # ─── 1. Introdução ────────────────────────────────────────────────
    h1(doc, "1. Introdução")
    para(doc,
         "Este tutorial conduz o instalador através das etapas necessárias para colocar o GCA "
         "(Gestão de Codificação Assistida) em produção numa máquina com Windows 10/11 ou Ubuntu "
         "22.04+. A instalação é inteiramente containerizada (Docker) e o processo é assistido por "
         "um assistente gráfico em Windows (Inno Setup) e por um script interativo em Ubuntu "
         "(install.sh).")
    para(doc,
         "Todo o processo leva entre 6 e 15 minutos, dependendo da velocidade da Internet (imagens "
         "Docker totalizam aproximadamente 900 MB) e da máquina. Ao final, o GCA ficará disponível "
         "no navegador e o primeiro Administrador já estará cadastrado.")

    h2(doc, "1.1. Quem deve seguir este documento")
    bullet(doc, "Administradores de TI responsáveis por instalar o GCA no cliente.")
    bullet(doc, "Desenvolvedores responsáveis por validar um ambiente de homologação.")
    bullet(doc, "Analistas de suporte que precisam reinstalar a partir de backup.")
    para(doc, "Pré-requisito mínimo: domínio de linha de comando (CMD / PowerShell no Windows, bash no Linux) e noções básicas de Docker.")

    h2(doc, "1.2. O que você vai precisar")
    bullet(doc, "Um arquivo de licença válido (formato texto, chave GCA-PROD-*).")
    bullet(doc, "Acesso de Administrador (Windows) ou sudo (Ubuntu).")
    bullet(doc, "Credenciais do registry privado do GCA (usuário e token, fornecidos no e-mail de contratação).")
    bullet(doc, "Conexão estável com a Internet durante a instalação.")
    bullet(doc, "Endereço (domínio ou IP) e porta que vão expor o GCA.")
    bullet(doc, "E-mail + senha que serão usados como primeiro Administrador.")
    bullet(doc, "Provedor de IA escolhido + chave de API (Anthropic, OpenAI, DeepSeek, Google ou Ollama local).")

    h2(doc, "1.3. Arquitetura instalada")
    para(doc,
         "Ao final da instalação, a máquina vai estar rodando seis containers Docker orquestrados "
         "por Docker Compose:")
    table(doc,
          ["Container", "Função", "Porta"],
          [
              ["gca-postgres", "Banco de dados PostgreSQL 15 com volume persistente gca-postgres-data.", "5432 (interna)"],
              ["gca-backend", "API FastAPI em Python 3.11 com 4 workers Uvicorn.", "8000"],
              ["gca-frontend", "Build estático React servido por vite preview.", "5173"],
              ["gca-redis", "Cache opcional para pipelines de ingestão.", "6379 (interna)"],
              ["gca-ollama", "Provedor de IA local (opcional — só sobe se configurado).", "11434 (interna)"],
              ["gca-n8n", "Orquestrador de workflows (opcional — histórico).", "5678 (interna)"],
          ],
          widths=[3.5, 9, 3])

    page_break(doc)

    # ─── 2. Pré-requisitos detalhados ─────────────────────────────────
    h1(doc, "2. Pré-requisitos de máquina")

    h2(doc, "2.1. Hardware recomendado")
    table(doc,
          ["Recurso", "Mínimo", "Recomendado"],
          [
              ["Processador", "4 núcleos 64-bit (Intel/AMD com AVX2)", "8 núcleos 64-bit"],
              ["Memória RAM", "8 GB", "16 GB ou mais"],
              ["Espaço em disco", "30 GB livres em SSD", "100 GB em SSD NVMe"],
              ["Rede", "10 Mbps para download inicial", "100 Mbps simétrico"],
              ["Sistema", "Windows 10/11 (build 19045+) ou Ubuntu 22.04+", "Linux nativo de preferência"],
          ],
          widths=[3.5, 6, 6])
    callout(doc, "Atenção:",
            "O GCA pode funcionar em máquinas menores, mas os backups diários e a geração do OCG "
            "podem ficar lentos em configurações abaixo do mínimo.", AMBER)

    h2(doc, "2.2. Docker")
    h3(doc, "Windows")
    numbered(doc, "Baixe Docker Desktop em https://www.docker.com/products/docker-desktop")
    numbered(doc, "Durante a instalação, marque \"Use WSL 2 instead of Hyper-V\".")
    numbered(doc, "Após instalar, reinicie a máquina.")
    numbered(doc, "Abra o Docker Desktop e confirme que o ícone da baleia ficou verde (\"Docker Desktop is running\").")
    numbered(doc, "Em Settings → Resources, reserve ao menos 4 GB de RAM para o Docker.")

    h3(doc, "Ubuntu")
    para(doc, "Execute em um terminal com sudo:")
    code(doc, "sudo apt update && sudo apt install -y ca-certificates curl gnupg")
    code(doc, "sudo install -m 0755 -d /etc/apt/keyrings")
    code(doc, "curl -fsSL https://download.docker.com/linux/ubuntu/gpg | sudo gpg --dearmor -o /etc/apt/keyrings/docker.gpg")
    code(doc, "echo \"deb [arch=$(dpkg --print-architecture) signed-by=/etc/apt/keyrings/docker.gpg] https://download.docker.com/linux/ubuntu $(. /etc/os-release && echo $VERSION_CODENAME) stable\" | sudo tee /etc/apt/sources.list.d/docker.list")
    code(doc, "sudo apt update && sudo apt install -y docker-ce docker-ce-cli containerd.io docker-buildx-plugin docker-compose-plugin")
    code(doc, "sudo usermod -aG docker $USER && newgrp docker")

    page_break(doc)

    # ─── 3. Instalação Windows ────────────────────────────────────────
    h1(doc, "3. Instalação em Windows")
    para(doc,
         "O instalador do GCA para Windows é gerado com Inno Setup e distribuído como um único "
         "arquivo GCA-Setup-1.0.exe. O fluxo é assistido em 11 telas (10 passos interativos + tela "
         "de conclusão). As imagens abaixo são capturas do próprio instalador em funcionamento.")

    # ─── Passo 1: Boas-vindas ─────────────────────────────────────────
    h2(doc, "Passo 1 — Boas-vindas")
    para(doc,
         "Dê duplo clique em GCA-Setup-1.0.exe. Uma janela UAC do Windows vai solicitar permissões "
         "administrativas — confirme. Em seguida, a tela de boas-vindas do assistente aparece "
         "listando as 10 etapas que virão. Clique em \"Avançar\" para prosseguir.")
    image(doc, WIREFRAME_DIR / "01_boas_vindas.png",
          "Figura 1 — Tela de boas-vindas do instalador GCA.")

    # ─── Passo 2: EULA ────────────────────────────────────────────────
    h2(doc, "Passo 2 — Aceite do contrato de licença (EULA)")
    para(doc,
         "Leia atentamente os termos de uso. O contrato estabelece cinco cláusulas centrais: "
         "proibição de engenharia reversa, responsabilidades do cliente (backups, chave mestra, "
         "credenciais), não compartilhamento de dados entre instâncias, mecanismo de atualização "
         "via release versionada e prestação de suporte pelo contrato comercial separado.")
    para(doc,
         "Após ler, marque \"Li e aceito os termos do contrato de licença\" e clique em "
         "\"Avançar\".")
    image(doc, WIREFRAME_DIR / "02_eula.png",
          "Figura 2 — Tela de aceite do contrato de licença.")
    callout(doc, "Importante:",
            "O aceite do EULA é obrigatório. Não é possível prosseguir sem marcar a caixa.", AMBER)

    # ─── Passo 3: Chave de ativação ───────────────────────────────────
    h2(doc, "Passo 3 — Chave de ativação")
    para(doc,
         "Cole a chave que você recebeu por e-mail no campo único da tela. A chave tem o formato "
         "GCA-PROD-XXXXX-XXXXX-XXXXX-XXXXX (24 caracteres em 5 grupos separados por hífen). "
         "O instalador valida a chave contra o registry privado do GCA e exibe:")
    bullet(doc, "A data de expiração da licença.")
    bullet(doc, "O número máximo de projetos simultâneos contratados.")
    bullet(doc, "O número máximo de usuários ativos contratados.")
    bullet(doc, "O nível de suporte contratado.")
    para(doc,
         "Se a chave for inválida ou estiver expirada, o instalador exibe erro em destaque e não "
         "permite avançar. Nesses casos, consulte o e-mail de contratação ou entre em contato com "
         "o fornecedor.")
    image(doc, WIREFRAME_DIR / "03_chave_ativacao.png",
          "Figura 3 — Validação da chave de ativação.")

    # ─── Passo 4: Pré-requisitos ──────────────────────────────────────
    h2(doc, "Passo 4 — Verificação de pré-requisitos")
    para(doc,
         "O instalador verifica automaticamente seis pré-requisitos da máquina. Cada item aparece "
         "com um ícone verde (OK) ou vermelho (falha):")
    bullet(doc, "Sistema operacional — Windows 10/11 64-bit build 19045 ou superior.")
    bullet(doc, "Memória RAM — no mínimo 8 GB.")
    bullet(doc, "Espaço em disco — no mínimo 30 GB livres na partição escolhida.")
    bullet(doc, "Docker Desktop — versão 4.30 ou superior, com WSL 2 ativado.")
    bullet(doc, "Acesso à Internet — para baixar as imagens do GCA Registry.")
    bullet(doc, "Permissão de Administrador — necessária para registrar serviços.")
    para(doc,
         "Se algum item falhar, corrija o problema e clique em \"Verificar novamente\". Só é "
         "possível avançar quando os seis itens estão verdes.")
    image(doc, WIREFRAME_DIR / "04_prerequisitos.png",
          "Figura 4 — Verificação automatizada dos seis pré-requisitos.")

    # ─── Passo 5: Pasta de instalação ─────────────────────────────────
    h2(doc, "Passo 5 — Pasta de instalação")
    para(doc,
         "Escolha onde os arquivos de configuração do GCA ficarão. O default é "
         "C:\\Program Files\\GCA. A pasta aloca ~25 GB — o volume maior de dados (banco, uploads, "
         "backups) fica em volumes Docker nomeados, separados da pasta acima.")
    para(doc,
         "Os três volumes nomeados criados automaticamente são:")
    bullet(doc, "gca-postgres-data — banco de dados PostgreSQL.")
    bullet(doc, "gca-uploads-storage — uploads e anexos de tickets.")
    bullet(doc, "gca-backups — backups por projeto (DT-063, até 10 retidos por projeto).")
    image(doc, WIREFRAME_DIR / "05_pasta_instalacao.png",
          "Figura 5 — Escolha de pasta de instalação.")

    # ─── Passo 6: Porta e domínio ─────────────────────────────────────
    h2(doc, "Passo 6 — Porta e domínio")
    para(doc,
         "Configure os endereços pelos quais o GCA será acessado:")
    table(doc,
          ["Campo", "Padrão", "Descrição"],
          [
              ["Domínio", "(vazio)", "Opcional. Informe um domínio se você tem proxy reverso com HTTPS (nginx, Caddy, Cloudflare Tunnel)."],
              ["Porta do frontend", "5173", "Porta HTTP do Vite preview. Altere se já tiver algum serviço rodando nela."],
              ["Porta da API", "8000", "Porta HTTP do FastAPI. Altere se conflitar."],
              ["Porta PostgreSQL", "5432", "Porta interna, normalmente não exposta ao host."],
          ],
          widths=[3.5, 2.5, 9])
    callout(doc, "Recomendado:",
            "Em produção, use proxy reverso com HTTPS válido (Let's Encrypt via Caddy, por "
            "exemplo). Expor Vite preview diretamente na Internet só é aceitável em ambiente de "
            "teste interno.", AMBER)
    image(doc, WIREFRAME_DIR / "06_rede_porta.png",
          "Figura 6 — Configuração de rede (domínio, portas).")

    # ─── Passo 7: Admin inicial ───────────────────────────────────────
    h2(doc, "Passo 7 — Criação do primeiro Administrador")
    para(doc,
         "Nesta etapa você cria o usuário com is_admin=true que terá acesso à camada "
         "administrativa da instância. Esse administrador pode, depois, criar outros admins via "
         "a UI /admin/users.")
    bullet(doc, "Nome completo — usado em notificações e no cabeçalho da UI.")
    bullet(doc, "E-mail — usado como login. Precisa ser único na instância.")
    bullet(doc, "Senha — mínimo 10 caracteres, 1 maiúscula, 1 número e 1 caractere especial.")
    bullet(doc, "Confirmação da senha — redigitação para evitar erro de digitação.")
    para(doc,
         "A senha é gravada como hash bcrypt com salt único. O texto claro não é armazenado em "
         "nenhum log, arquivo ou banco.")
    image(doc, WIREFRAME_DIR / "07_admin_inicial.png",
          "Figura 7 — Cadastro do primeiro Administrador.")

    # ─── Passo 8: Provedor IA ─────────────────────────────────────────
    h2(doc, "Passo 8 — Provedor de IA padrão")
    para(doc,
         "Escolha qual provedor de IA será o padrão da instância. Essa escolha pode ser alterada "
         "depois por instância (em /admin/settings) e também por projeto (em "
         "/projects/{id}/settings).")
    para(doc,
         "Os cinco provedores suportados são:")
    table(doc,
          ["Provedor", "Criticidade", "Custo aprox. por OCG", "Observações"],
          [
              ["Anthropic (Claude)", "Alta", "US$ 0,008", "Recomendado para OCG e decisões críticas."],
              ["OpenAI (GPT-4)", "Alta", "US$ 0,012", "Qualidade premium, um pouco mais caro."],
              ["Google (Gemini)", "Média-Alta", "US$ 0,004", "Bom para textos longos."],
              ["DeepSeek", "Baixa-Média", "US$ 0,0005", "Adequado para tarefas simples. Não recomendado para OCG consolidado."],
              ["Ollama (local)", "Baixa-Média", "Gratuito", "Requer GPU (mínimo 8 GB VRAM) para qualidade decente."],
          ],
          widths=[3.5, 2.5, 3, 6])
    para(doc,
         "Cole a chave de API do provedor escolhido. A chave é criptografada com Fernet (chave "
         "derivada da GCA_MASTER_KEY da instância) antes de ser salva.")
    image(doc, WIREFRAME_DIR / "08_provedor_ia.png",
          "Figura 8 — Configuração do provedor de IA padrão.")

    # ─── Passo 9: Resumo ──────────────────────────────────────────────
    h2(doc, "Passo 9 — Resumo da instalação")
    para(doc,
         "Antes de aplicar a instalação, o assistente mostra um resumo com todas as configurações "
         "coletadas nas etapas anteriores. Reveja com atenção:")
    bullet(doc, "Pasta de instalação.")
    bullet(doc, "Domínio configurado.")
    bullet(doc, "Portas do frontend e da API.")
    bullet(doc, "Administrador inicial (nome + e-mail).")
    bullet(doc, "Provedor de IA.")
    bullet(doc, "Imagens Docker que serão baixadas.")
    bullet(doc, "Volumes Docker que serão criados.")
    bullet(doc, "Tempo estimado de instalação.")
    para(doc,
         "Se algum item estiver errado, clique em \"Voltar\" até o passo correspondente. Se tudo "
         "estiver correto, clique em \"Instalar\" para iniciar o processo.")
    image(doc, WIREFRAME_DIR / "09_resumo.png",
          "Figura 9 — Resumo com todas as configurações antes de aplicar.")

    # ─── Passo 10: Instalando ─────────────────────────────────────────
    h2(doc, "Passo 10 — Instalação em curso")
    para(doc,
         "O instalador executa 13 sub-etapas sequencialmente. A barra de progresso mostra o "
         "percentual geral e o log abaixo mostra cada sub-etapa com status (pendente, em "
         "andamento, concluída). Não feche esta janela — você pode minimizar.")
    bullet(doc, "Validação da chave de ativação contra o registry privado.")
    bullet(doc, "Verificação do Docker Desktop.")
    bullet(doc, "Criação dos volumes Docker nomeados.")
    bullet(doc, "Geração do arquivo .env com as configurações coletadas.")
    bullet(doc, "Gravação do Admin inicial criptografado.")
    bullet(doc, "Gravação do provedor de IA com api_key criptografada.")
    bullet(doc, "Download das imagens Docker do registry privado autenticado.")
    bullet(doc, "docker compose up -d para subir os containers.")
    bullet(doc, "Aguardar health check do backend (timeout 60s).")
    bullet(doc, "Aplicar migrations do banco.")
    bullet(doc, "Bootstrap do primeiro Admin no banco.")
    image(doc, WIREFRAME_DIR / "10_instalando.png",
          "Figura 10 — Progresso da instalação com log detalhado.")
    callout(doc, "Se alguma etapa falhar:",
            "O instalador faz rollback automático e exibe o erro em destaque. Consulte o capítulo "
            "5 (Troubleshooting) para os erros mais comuns e correções.", RED)

    # ─── Passo 11: Conclusão ──────────────────────────────────────────
    h2(doc, "Passo 11 — Conclusão")
    para(doc,
         "Instalação concluída. A tela final mostra os próximos passos e o acesso imediato ao "
         "GCA. Se a caixa \"Abrir o GCA agora no navegador\" estiver marcada, o navegador padrão "
         "abre automaticamente na URL configurada.")
    bullet(doc, "Acesse o GCA pelo domínio configurado (ou http://localhost:5173 se não configurou domínio).")
    bullet(doc, "Faça login com o e-mail e a senha definidos no Passo 7.")
    bullet(doc, "Aprove a primeira solicitação de projeto em /admin/projects.")
    image(doc, WIREFRAME_DIR / "11_conclusao.png",
          "Figura 11 — Tela de conclusão com instruções de próximos passos.")

    page_break(doc)

    # ─── 4. Instalação Ubuntu ─────────────────────────────────────────
    h1(doc, "4. Instalação em Ubuntu")
    para(doc,
         "A instalação em Ubuntu 22.04+ usa um script interativo install.sh que cobre as mesmas "
         "10 etapas do instalador Windows, mas em modo texto. Alternativamente, o GCA pode ser "
         "instalado via pacote .deb para distribuição homogênea em frota.")

    h2(doc, "4.1. Instalação via install.sh")
    para(doc, "Baixe o pacote do registry privado e execute:")
    code(doc, "wget https://registry.gca-produto.com/downloads/install.sh")
    code(doc, "chmod +x install.sh")
    code(doc, "sudo ./install.sh")
    para(doc,
         "O script apresenta as 10 etapas interativamente no terminal. Cada pergunta tem default "
         "entre colchetes — basta apertar Enter para aceitar. Exemplos:")
    code(doc, "$ Informe a chave de ativação: GCA-PROD-ABCD1-EFGH2-IJKL3-MNOPQ")
    code(doc, "$ Pasta de instalação [/opt/gca]:")
    code(doc, "$ Domínio [localhost]: gca.empresa.com.br")
    code(doc, "$ Porta frontend [5173]:")
    code(doc, "$ Porta API [8000]:")
    code(doc, "$ E-mail do Admin inicial: admin@empresa.com.br")
    code(doc, "$ Senha do Admin inicial: ****")
    code(doc, "$ Confirmar senha: ****")
    code(doc, "$ Provedor de IA (anthropic/openai/gemini/deepseek/ollama) [anthropic]:")
    code(doc, "$ Chave de API do provedor: sk-ant-*******************")
    code(doc, "$ Confirmar instalação? [s/N]: s")
    para(doc, "Ao final, o script mostra a URL e o status de cada container.")

    h2(doc, "4.2. Instalação via pacote .deb")
    para(doc, "Para distribuição em frota (Ansible, Puppet, Chef):")
    code(doc, "wget https://registry.gca-produto.com/downloads/gca_1.0.0_amd64.deb")
    code(doc, "sudo dpkg -i gca_1.0.0_amd64.deb")
    code(doc, "sudo apt-get install -f   # resolve dependências, se alguma faltar")
    para(doc,
         "O pacote instala os arquivos em /opt/gca, registra o serviço systemd gca.service e "
         "dispara install.sh em modo não-interativo lendo variáveis de ambiente (GCA_LICENSE, "
         "GCA_ADMIN_EMAIL, GCA_ADMIN_PASSWORD, GCA_DOMAIN etc).")
    para(doc, "Após instalar, inicie o serviço:")
    code(doc, "sudo systemctl enable --now gca.service")
    code(doc, "sudo systemctl status gca.service")

    h2(doc, "4.3. Pós-instalação Ubuntu")
    bullet(doc, "Conferir que os containers estão de pé: docker ps.")
    bullet(doc, "Conferir o log do backend: docker logs gca-backend --tail 30.")
    bullet(doc, "Acessar a URL configurada pelo navegador.")
    bullet(doc, "Fazer login com as credenciais de Admin definidas.")

    page_break(doc)

    # ─── 5. Pós-instalação — telas reais ──────────────────────────────
    h1(doc, "5. Primeiros passos na aplicação")
    para(doc,
         "Esta seção mostra as telas reais da aplicação GCA em funcionamento, capturadas com os "
         "usuários Administrador e GP. Use-as como referência visual após concluir a instalação.")

    h2(doc, "5.1. Login Administrador")
    para(doc, "Acesse a URL configurada. A tela de login apresenta campos de e-mail e senha e redireciona "
              "para /admin após autenticação bem-sucedida.")
    image(doc, SCREENSHOT_DIR / "01_publica_login_admin.png", "Figura 12 — Tela de login do Administrador.", width_in=5.9)

    h2(doc, "5.2. Dashboard Global (primeira tela após login Admin)")
    para(doc, "O Dashboard Global agrega contadores de projetos, tickets e releases recentes da instância.")
    image(doc, SCREENSHOT_DIR / "10_admin_dashboard_global.png", "Figura 13 — Dashboard Global do Admin.", width_in=5.9)

    h2(doc, "5.3. Gestão de Projetos")
    para(doc, "Lista de todas as solicitações e projetos existentes na instância, com badges de "
              "lifecycle (Ativo / Pausado / Desativado / Excluído-órfão) e ações correspondentes.")
    image(doc, SCREENSHOT_DIR / "11_admin_gestao_projetos.png", "Figura 14 — Gestão de Projetos.", width_in=5.9)

    h2(doc, "5.4. Gestão de Usuários")
    para(doc, "Lista todos os usuários, permite promover a Administrador via o ícone Shield e convidar "
              "novos administradores pelo botão no topo direito.")
    image(doc, SCREENSHOT_DIR / "13_admin_gestao_usuarios.png", "Figura 15 — Gestão de Usuários.", width_in=5.9)

    h2(doc, "5.5. Modal \"Convidar Administrador\"")
    para(doc, "Preenchido com nome completo e e-mail. Se SMTP não estiver configurado, a senha "
              "temporária é exibida na tela para comunicação manual.")
    image(doc, SCREENSHOT_DIR / "14_admin_convidar_admin_modal.png", "Figura 16 — Modal de convite de Admin.", width_in=5.9)

    h2(doc, "5.6. Métricas (global + por projeto)")
    para(doc, "Página de métricas consolida totais globais (chamadas, tokens, custo) e faz breakdown "
              "por projeto em tabela ordenada por custo decrescente.")
    image(doc, SCREENSHOT_DIR / "16_admin_metricas.png", "Figura 17 — Métricas com breakdown por projeto.", width_in=5.9)

    h2(doc, "5.7. Equipe Sustentação")
    para(doc, "Administra a flag is_support. Admins herdam Sustentação automaticamente; usuários "
              "comuns podem ser promovidos via busca e botão de adição.")
    image(doc, SCREENSHOT_DIR / "19_admin_equipe_sustentacao.png", "Figura 18 — Equipe Sustentação.", width_in=5.9)

    h2(doc, "5.8. Releases (admin)")
    para(doc, "Lista de releases aplicadas e pendentes. Releases destrutivas ficam marcadas em "
              "âmbar e mostram botão \"Aplicar com snapshot\".")
    image(doc, SCREENSHOT_DIR / "20_admin_releases.png", "Figura 19 — Releases (visão admin).", width_in=5.9)

    h2(doc, "5.9. Detalhe de release com log de aplicação")
    image(doc, SCREENSHOT_DIR / "21_admin_release_detail.png", "Figura 20 — Detalhe da release + itens + log.", width_in=5.9)

    h2(doc, "5.10. Incidents (admin agregado)")
    image(doc, SCREENSHOT_DIR / "18_admin_incidents.png", "Figura 21 — Tickets escalados a Admin/Sustentação.", width_in=5.9)

    h2(doc, "5.11. Backups agregados")
    image(doc, SCREENSHOT_DIR / "17_admin_backups.png", "Figura 22 — Backups por projeto (visão admin).", width_in=5.9)

    h2(doc, "5.12. Auditoria Global")
    image(doc, SCREENSHOT_DIR / "15_admin_auditoria_global.png", "Figura 23 — Auditoria global com hash chain.", width_in=5.9)

    page_break(doc)

    # ─── 6. Dentro de um projeto ──────────────────────────────────────
    h1(doc, "6. Dentro de um projeto — fluxo do GP")
    para(doc,
         "A área de projeto concentra o fluxo diário do Gerente de Projeto (GP). Após a aprovação "
         "do projeto pelo Administrador, o GP acessa a lista de projetos (visão segmentada — sem "
         "os menus de camada administrativa) e entra no seu projeto. A sidebar se reconfigura "
         "automaticamente mostrando todos os sub-itens do projeto, que são o fluxo operacional "
         "completo do GP, Dev, Tester e QA.")
    para(doc,
         "As telas abaixo foram capturadas com o usuário Fernando, GP real do projeto Automação "
         "Jurídica Assistida, para demonstrar o layout correto da sidebar de projeto (em vez da "
         "sidebar administrativa). Observe a diferença: o menu \"ADMINISTRAÇÃO\" foi substituído "
         "por \"MEUS PROJETOS\" com o projeto ativo expandido e seus 18 sub-itens.")

    h2(doc, "6.0. Lista de projetos (visão do GP)")
    para(doc, "Entrada da sessão do GP: lista apenas os projetos onde ele é membro aceito. "
              "Cliques levam à rota /projects/{id}.")
    image(doc, SCREENSHOT_DIR / "30_projeto_gp_lista_projetos.png",
          "Figura — Lista de projetos visível ao GP.", width_in=5.9)

    project_shots = [
        ("6.1. Dashboard do projeto", "31_projeto_gp_dashboard.png",
         "Visão consolidada do projeto: readiness do OCG, últimos deltas, backlog priorizado, pendências. Note na sidebar: sub-itens específicos do projeto (Dashboard, Equipe, OCG, Repositórios, Ingestão, Gatekeeper, Arguidor, CodeGen, QA, Revisão, Backlog, Roadmap, Docs, Readiness, Settings, Audit, Backups, Incidentes, Métricas) — visualização do GP."),
        ("6.2. Equipe", "32_projeto_gp_team.png",
         "Gestão da equipe do projeto — convites, papéis, multi-papel."),
        ("6.3. OCG — Objeto Canônico de Governança", "33_projeto_gp_ocg.png",
         "Fonte única de verdade do projeto. Exibe os sete pilares (Business, Architecture, Stack, Testing, Compliance, Risk, Deliverables) com status e score."),
        ("6.4. Repositórios externos", "34_projeto_gp_external_repos.png",
         "Lista de repositórios Git vinculados ao projeto (integração via PAT criptografada)."),
        ("6.5. Ingestão", "35_projeto_gp_ingestion.png",
         "Uploads de documentos para análise. Detector de PII aplica quarentena quando necessário."),
        ("6.6. Gatekeeper", "36_projeto_gp_gatekeeper.png",
         "Avaliação por sete pilares. Bloqueia avanço de fase quando thresholds não são atingidos."),
        ("6.7. Arguidor", "37_projeto_gp_arguider.png",
         "Análise LLM dos documentos ingeridos. Gera findings que alimentam o backlog."),
        ("6.8. Geração de Código (CodeGen)", "38_projeto_gp_codegen.png",
         "Gera estrutura inicial de projeto com base no OCG + scaffolders determinísticos. Desde a emenda RBAC 2026-04-19, o GP pode operar CodeGen sem depender exclusivamente de Dev."),
        ("6.9. QA Readiness", "39_projeto_gp_qa_readiness.png",
         "Estado de prontidão de QA com cobertura por pilar."),
        ("6.10. Revisão de Testes", "40_projeto_gp_tester_review.png",
         "Workflow de aprovação de testes executados (Tester cria, QA aprova)."),
        ("6.11. Backlog", "41_projeto_gp_backlog.png",
         "Itens derivados do OCG + findings do Arguidor, priorizados."),
        ("6.12. Roadmap", "42_projeto_gp_roadmap.png",
         "Visão temporal de entregas planejadas."),
        ("6.13. Documentação Viva", "43_projeto_gp_docs.png",
         "Documentação gerada automaticamente a partir do OCG e atualizada a cada mudança."),
        ("6.14. Definition of Done", "44_projeto_gp_readiness.png",
         "Checklist de critérios de entrega final — Release Bundle."),
        ("6.15. Configurações", "45_projeto_gp_settings.png",
         "Três tabs: Questionário (PDF), Repositório (Git + PAT), Provedor de IA do projeto."),
        ("6.16. Pipeline Audit", "46_projeto_gp_audit.png",
         "Logs de auditoria específicos do projeto, filtrados por projeto."),
        ("6.17. Backups do projeto", "47_projeto_gp_backups.png",
         "Lista de backups do projeto, download, botão 'Backup agora', rollback."),
        ("6.18. Incidentes", "48_projeto_gp_incidents.png",
         "Tickets abertos pelo time do projeto. Filtros de status e prioridade."),
        ("6.19. Modal de abertura de ticket", "49_projeto_gp_abrir_ticket_modal.png",
         "Formulário: título, descrição, categoria, prioridade, seção (autopreenchida), fluxo (obrigatório), anexos."),
        ("6.20. Métricas do Projeto", "50_projeto_gp_metrics.png",
         "Uso de IA e eventos de auditoria compartimentalizados — visão do GP e demais membros."),
    ]

    for title, filename, desc in project_shots:
        h2(doc, title)
        para(doc, desc)
        image(doc, SCREENSHOT_DIR / filename, f"Figura — {title.split('. ', 1)[-1]}.", width_in=5.9)

    page_break(doc)

    # ─── 7. Operação depois de instalado ──────────────────────────────
    h1(doc, "7. Operação após a instalação")

    h2(doc, "7.1. Scripts que já vêm no pacote")
    table(doc,
          ["Script", "Função"],
          [
              ["/opt/gca/scripts/upgrade.sh", "Upgrade idempotente (DT-062) — 9 etapas com backup pré, fetch, build, migrations, recreate, health."],
              ["/opt/gca/scripts/backup.sh", "Backup completo pré-upgrade (postgres dump + volumes)."],
              ["/opt/gca/scripts/restore.sh", "Restore validando SHA256, dupla confirmação."],
              ["/opt/gca/scripts/health-check.sh", "Health endpoints do stack."],
          ],
          widths=[6, 11])

    h2(doc, "7.2. Rotina de manutenção mensal")
    numbered(doc, "Verificar espaço dos volumes Docker (docker system df).")
    numbered(doc, "Conferir contagem de backups por projeto (máx 10 retidos por projeto).")
    numbered(doc, "Revisar /admin/metrics — custo de IA acumulado, eventos de audit, uso por projeto.")
    numbered(doc, "Aplicar releases pendentes em /admin/releases (se houver).")
    numbered(doc, "Atualizar chave de ativação quando próximo da expiração.")

    h2(doc, "7.3. Atualização do GCA")
    para(doc, "Quando uma nova versão é liberada:")
    code(doc, "cd /opt/gca && sudo ./scripts/upgrade.sh")
    para(doc,
         "O script roda: backup → fetch → exit se já atualizado → git pull --ff-only → build → "
         "alembic upgrade → docker compose up -d --force-recreate → healthcheck 60s → smoke. "
         "Se qualquer passo falhar, cita o comando exato de restore.sh para reverter.")

    h2(doc, "7.4. Backup manual adicional")
    para(doc, "Além do backup automático às 12:00, o Admin ou GP pode disparar backup imediato:")
    bullet(doc, "UI: /projects/{id}/backups → botão 'Backup agora'.")
    bullet(doc, "UI admin: /admin/backups → selecionar projeto → 'Backup imediato a pedido do GP'.")
    bullet(doc, "CLI: sudo /opt/gca/scripts/backup.sh.")

    page_break(doc)

    # ─── 8. Troubleshooting ───────────────────────────────────────────
    h1(doc, "8. Troubleshooting")

    h2(doc, "8.1. Erros comuns de instalação")
    table(doc,
          ["Sintoma", "Causa provável", "Correção"],
          [
              ["\"Docker Desktop não responde\"", "Docker parado ou WSL 2 desativado", "Reiniciar Docker Desktop; habilitar WSL 2 em Settings."],
              ["\"Chave de ativação inválida\"", "Chave expirada ou registry inacessível", "Verificar e-mail original; confirmar conexão com registry.gca-produto.com."],
              ["\"Porta 5173 já em uso\"", "Outro serviço ocupando a porta", "Alterar porta no Passo 6 ou parar o serviço conflitante."],
              ["\"Health check falhou após 60s\"", "Backend não subiu — ver log", "docker logs gca-backend --tail 100 e reportar ao suporte se erro opaco."],
              ["\"Migration failed\"", "Migração parcial aplicada antes", "Restore do último backup antes de retentar upgrade."],
          ],
          widths=[5.5, 5, 6.5])

    h2(doc, "8.2. Comandos de diagnóstico")
    code(doc, "docker ps                                      # status dos containers")
    code(doc, "docker logs gca-backend --tail 100             # log do backend")
    code(doc, "docker logs gca-frontend --tail 50             # log do frontend")
    code(doc, "docker exec gca-postgres pg_isready -U gca     # ping DB")
    code(doc, "curl http://localhost:8000/api/v1/metrics/health   # health público")
    code(doc, "curl http://localhost:8000/docs                # OpenAPI")

    h2(doc, "8.3. Onde buscar ajuda")
    bullet(doc, "Contrato canônico GCA_CANONICAL_CONTRACT.md — regras soberanas do produto.")
    bullet(doc, "Documento de Requisitos v1.0 — detalhes técnicos completos.")
    bullet(doc, "Suporte comercial: ver contrato comercial assinado no momento da contratação.")
    bullet(doc, "Auditoria: /admin/audit dentro da instância para investigar eventos históricos.")

    page_break(doc)

    # ─── 9. Glossário ─────────────────────────────────────────────────
    h1(doc, "9. Glossário")
    para(doc, "Termos em ordem alfabética referentes à instalação e operação do GCA.", size=10, color=SLATE_MEDIUM)

    glossary = [
        ("Admin", "Usuário da instância com is_admin=true. Configura a instância; não atua operacionalmente em projetos."),
        ("Ansible", "Ferramenta de automação de infraestrutura. Pode ser usada para instalar o .deb em frota."),
        ("Anthropic", "Provedor de IA (Claude). Recomendado para tarefas de alta criticidade no GCA."),
        ("APScheduler", "Biblioteca Python de agendamento. Dispara o backup diário às 12:00."),
        ("Backup", "Cópia consistente dos dados do projeto. Gerenciado por DT-063. Até 10 retidos por projeto."),
        ("BCC", "BytecodeCompiler do PyArmor (variante gratuita) usada no empacotamento."),
        ("Bcrypt", "Algoritmo de hash de senhas utilizado no GCA com salt único por usuário."),
        ("Caddy", "Servidor web com HTTPS automático via Let's Encrypt. Recomendado como proxy reverso em produção."),
        ("Catch-up", "Rotina do scheduler que dispara backup se o último foi há mais de 24h (pós inicialização)."),
        ("Chave de ativação", "Token GCA-PROD-XXXXX-XXXXX-XXXXX-XXXXX fornecido na contratação. Determina validade e limites da licença."),
        ("Container", "Unidade de deployment isolada via Docker. O GCA usa seis containers."),
        ("Cython", "Ferramenta que compila Python para C e depois para binário. Usada na proteção de código."),
        ("DT-062", "Dívida técnica do MVP 5 — script upgrade.sh idempotente."),
        ("DT-063", "Dívida técnica — backup compartimentalizado por projeto com scheduler diário."),
        ("Docker", "Runtime de containers usado pelo GCA. Versão mínima 24."),
        ("Docker Compose", "Orquestrador de multi-container definido em docker-compose.yml."),
        ("Docker Desktop", "Aplicação Docker para Windows/Mac. Requer WSL 2 ativo em Windows."),
        ("EULA", "End User License Agreement — contrato de licença aceito no Passo 2 da instalação."),
        ("Fernet", "Algoritmo de criptografia simétrica usado pelo GCA para credenciais persistidas."),
        ("GCA", "Gestão de Codificação Assistida / Gerenciador Central de Arquiteturas — nome do produto."),
        ("GCA_MASTER_KEY", "Chave mestra da instância. Gerada durante a instalação; usada pelo Fernet."),
        ("GP", "Gerente de Projeto. Soberano do projeto. Tem acesso a todas as funcionalidades dentro do projeto."),
        ("Health check", "Endpoint /api/v1/metrics/health que retorna 200 quando o backend está operacional."),
        ("Inno Setup", "Ferramenta open-source para gerar instaladores Windows .exe. Usada no GCA."),
        ("install.sh", "Script interativo de instalação em Ubuntu."),
        ("Migration", "Script SQL aplicado ao banco durante upgrade. Nunca destrutivo por default (MVP 7)."),
        ("Multi-stage", "Técnica Docker de separar build e runtime. Usada nas imagens de produção."),
        ("Ollama", "Provedor de IA local (self-hosted). Opcional. Requer hardware com GPU."),
        ("OCG", "Objeto Canônico de Governança. Fonte única de verdade do projeto."),
        ("PostgreSQL", "Banco de dados relacional usado pelo GCA (versão 15)."),
        ("PyArmor", "Ferramenta de obfuscação de Python. Variante BCC (gratuita) é a usada no GCA."),
        ("Release", "Entrega versionada do GCA. Declarada em backend/releases/*.yaml. Pode ser destrutiva."),
        ("Registry", "Repositório de imagens Docker. O GCA usa registry privado autenticado."),
        ("Rollback", "Restauração de dados via snapshot pré-release (DT-063). Por-projeto."),
        ("Scheduler", "APScheduler — dispara backups diários às 12:00."),
        ("SHA-256", "Algoritmo de hash usado para verificar integridade de backups e manifestos de release."),
        ("Smoke test", "Verificação rápida pós-upgrade de que o sistema está operacional."),
        ("SMTP", "Protocolo de envio de e-mail. O GCA usa SMTP compartimentalizado por projeto (DT-016)."),
        ("Snapshot", "Backup prévio automático antes de aplicação de release destrutiva."),
        ("Sustentação", "Conjunto de usuários com is_support=true que recebem tickets escalados a Admin."),
        ("Systemd", "Gerenciador de serviços do Linux. O GCA registra o serviço gca.service."),
        ("Ticket", "Registro de incidente aberto por usuário. Ver Documento de Requisitos §3.5."),
        ("Ubuntu", "Distribuição Linux suportada nativamente (22.04+)."),
        ("Upgrade", "Processo de atualização do GCA via scripts/upgrade.sh."),
        ("Uvicorn", "Servidor ASGI que executa o FastAPI do GCA."),
        ("Vite", "Ferramenta de build para frontend React. O GCA usa vite preview em produção."),
        ("Volume", "Área de armazenamento persistente do Docker. Três volumes nomeados: gca-postgres-data, gca-uploads-storage, gca-backups."),
        ("Windows", "Sistema operacional suportado via instalador Inno Setup (.exe)."),
        ("WSL 2", "Windows Subsystem for Linux versão 2. Obrigatório para Docker Desktop em Windows."),
        ("YAML", "Formato de declaração de releases em backend/releases/*.yaml."),
    ]

    for term, definition in sorted(glossary, key=lambda x: x[0].lower()):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        t_run = p.add_run(f"{term}. ")
        t_run.bold = True
        t_run.font.size = Pt(11)
        t_run.font.color.rgb = VIOLET
        d_run = p.add_run(definition)
        d_run.font.size = Pt(11)

    doc.add_paragraph()
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = final.add_run("— Fim do tutorial —")
    fr.italic = True
    fr.font.color.rgb = SLATE_MEDIUM
    fr.font.size = Pt(10)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    build(doc)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"✓ Tutorial gerado: {OUT_PATH}")
    print(f"  Tamanho: {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
