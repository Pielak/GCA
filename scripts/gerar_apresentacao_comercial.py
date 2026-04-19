#!/usr/bin/env python3
"""
Gera "GCA - Gestão de Codificação Assistida.docx" — documento de
apresentação comercial do produto, com as 40 screenshots e os 10
diagramas Mermaid renderizados embedados.

Baseado em docs/GCA_Apresentacao_Comercial.md.

Autor: Luiz Carlos Pielak
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

DIAGRAM_DIR = Path("/home/luiz/GCA/docs/diagrams")
SCREENSHOT_DIR = Path("/home/luiz/GCA/screenshots_v3")
OUT_PATH = Path("/home/luiz/GCA/docs/GCA - Gestão de Codificação Assistida.docx")

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
SLATE_DARK = RGBColor(0x1E, 0x29, 0x3B)
SLATE_MEDIUM = RGBColor(0x64, 0x74, 0x8B)
EMERALD = RGBColor(0x05, 0x96, 0x69)
AMBER = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = VIOLET
        run.font.size = Pt(22)


def h2(doc, text: str):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = SLATE_DARK
        run.font.size = Pt(16)


def h3(doc, text: str):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = SLATE_MEDIUM
        run.font.size = Pt(13)


def para(doc, text: str, *, size: int = 11, bold: bool = False,
         color: RGBColor = None, justify: bool = True, italic: bool = False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


def bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


def numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


def code(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = SLATE_DARK


def image(doc, path: Path, caption: str = "", width_in: float = 6.0):
    if not path.exists():
        para(doc, f"[imagem ausente: {path.name}]", color=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = SLATE_MEDIUM


def table(doc, headers: list[str], rows: list[list[str]], widths: list[float] = None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 4"
    t.autofit = False
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "6D28D9")
        if widths and i < len(widths):
            cell.width = Cm(widths[i])
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if widths and ci < len(widths):
                cell.width = Cm(widths[ci])


def page_break(doc):
    doc.add_page_break()


def callout(doc, label: str, text: str, color: RGBColor = AMBER):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.color.rgb = color
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)


def quote(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = SLATE_DARK


# ─── Documento ────────────────────────────────────────────────────────────

def build(doc: Document):
    core = doc.core_properties
    core.title = "GCA — Gestão de Codificação Assistida"
    core.author = "Luiz Carlos Pielak"
    core.subject = "Apresentação comercial do produto GCA"
    core.keywords = "GCA, apresentação, produto, governança, IA, desenvolvimento"
    core.last_modified_by = "Luiz Carlos Pielak"
    core.created = datetime.now()

    # ═══ Capa ═══
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("GCA")
    tr.font.size = Pt(80)
    tr.bold = True
    tr.font.color.rgb = VIOLET

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Gestão de Codificação Assistida")
    sr.font.size = Pt(26)
    sr.font.color.rgb = SLATE_DARK

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sub2.add_run("Gerenciador Central de Arquiteturas")
    sr2.font.size = Pt(16)
    sr2.font.color.rgb = SLATE_MEDIUM
    sr2.italic = True

    for _ in range(4):
        doc.add_paragraph()

    pres = doc.add_paragraph()
    pres.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = pres.add_run("Apresentação Comercial do Produto")
    pr.font.size = Pt(18)
    pr.italic = True
    pr.font.color.rgb = SLATE_DARK

    for _ in range(6):
        doc.add_paragraph()

    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = auth.add_run("Autor: Luiz Carlos Pielak")
    ar.font.size = Pt(14)
    ar.bold = True

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dt.add_run(f"Versão 1.0 — {datetime.now().strftime('%d/%m/%Y')}")
    dr.font.size = Pt(11)
    dr.font.color.rgb = SLATE_MEDIUM

    stat = doc.add_paragraph()
    stat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stat_r = stat.add_run("MVPs 1-7 fechados · 732 testes de regressão · Zero dívida técnica")
    stat_r.font.size = Pt(10)
    stat_r.font.color.rgb = EMERALD
    stat_r.italic = True

    page_break(doc)

    # ═══ 1. Resumo executivo ═══
    h1(doc, "1. Resumo executivo")
    para(doc,
         "O GCA (Gestão de Codificação Assistida) é uma plataforma instalável por cliente "
         "para governança de projetos de desenvolvimento de software assistida por "
         "Inteligência Artificial. O produto cobre o ciclo completo de um projeto, da "
         "solicitação externa até a entrega do Release Bundle, passando por questionário "
         "de requisitos, OCG (Objeto Canônico de Governança), Gatekeeper de sete pilares, "
         "ingestão de artefatos, Arguidor, geração de código com scaffolders determinísticos, "
         "QA Readiness e Documentação Viva.")

    h2(doc, "1.1. Ciclo do projeto coberto pelo GCA")
    bullet(doc, "Entrada — solicitação externa + aprovação administrativa + cadastro do GP.")
    bullet(doc, "Descoberta — questionário de 49 perguntas e geração do OCG via pipeline de 8 agentes de IA.")
    bullet(doc, "Validação — Gatekeeper avalia 7 pilares (Business, Architecture, Stack, Testing, Compliance, Risk, Deliverables).")
    bullet(doc, "Aprofundamento — ingestão de documentos, Arguidor, backlog, roadmap.")
    bullet(doc, "Execução — CodeGen assistida com scaffolders em 7 linguagens (Python, Java, Kotlin, Go, C#, PHP, Node.js).")
    bullet(doc, "Qualidade — QA Readiness, revisão de testes, Release Bundle.")
    bullet(doc, "Operação — backup diário por projeto, tickets de incidente, releases versionadas preservando dados, métricas de IA.")

    h2(doc, "1.2. Diferenciais competitivos")
    bullet(doc, "Uma instância por cliente — sem SaaS multi-tenant. Dados nunca saem do ambiente do cliente.")
    bullet(doc, "Compartimentalização dura por projeto em banco, IA, storage, auditoria e backup.")
    bullet(doc, "IA configurável por cliente e por projeto — Anthropic Claude, OpenAI GPT, Google Gemini, DeepSeek, Ollama local.")
    bullet(doc, "Governança auditável — todas as decisões críticas geram entrada em audit_log_global com hash chain.")
    bullet(doc, "Proteção de código — sete camadas anti-engenharia-reversa documentadas.")
    bullet(doc, "Instalação assistida em Windows e Ubuntu com 10 etapas guiadas.")

    page_break(doc)

    # ═══ 2. O que o GCA resolve ═══
    h1(doc, "2. O que o GCA resolve")
    para(doc,
         "Projetos de desenvolvimento de software em organizações de médio e grande porte "
         "enfrentam, tipicamente, sete problemas estruturais. O GCA endereça cada um deles "
         "com mecanismo específico e auditável.")

    problems = [
        ("2.1. Perda de contexto entre fases",
         "Sem um objeto canônico único, cada fase do projeto (levantamento, arquitetura, codificação, "
         "testes, entrega) é conduzida com premissas próprias. O que foi decidido no levantamento se "
         "perde até chegar no CodeGen.",
         "Resposta: o OCG é a fonte única de verdade. Todas as fases consultam e evoluem o mesmo "
         "objeto. Mudanças geram delta versionado em ocg_delta_log."),
        ("2.2. Governança ad-hoc e difícil de auditar",
         "Aprovações informais, decisões em reuniões sem registro, deliberações por e-mail. Na hora "
         "do post-mortem, ninguém encontra quem aprovou o quê.",
         "Resposta: cada decisão crítica gera evento em audit_log_global com hash chain — qualquer "
         "adulteração invalida a cadeia. Governança e compliance ganham base factual."),
        ("2.3. Escolha de IA acoplada ao produto",
         "Muitas plataformas casam com um único provedor de IA, forçando o cliente a aceitar a "
         "escolha do fornecedor. Isso quebra quando o cliente tem política de dados, custo ou "
         "compliance diferente.",
         "Resposta: política explícita de roteamento híbrido por criticidade e configuração por "
         "instância e por projeto. O cliente decide qual provedor vai processar cada projeto dele."),
        ("2.4. Dependência de conhecimento tácito do GP",
         "Gerente de Projeto é gargalo — ele sabe como fazer mas não existe sistema que codifique "
         "esse saber. Quando sai, o projeto trava.",
         "Resposta: o pipeline GCA transforma o conhecimento do GP em artefatos auditáveis (OCG, "
         "backlog, roadmap, Documentação Viva). Mesmo com troca de GP, o projeto continua rastreável."),
        ("2.5. Tickets de suporte perdidos em canais paralelos",
         "Bug reportado por e-mail, urgência no Slack, feature solicitada em reunião — nada fica "
         "centralizado e o ciclo ticket → correção → release não é rastreável.",
         "Resposta: módulo de Tickets de Incidente com roteamento automático por papel. Cada ticket "
         "tem categoria, prioridade, seção onde ocorreu, fluxo executado e anexos. O MVP 7 amarra "
         "tickets resolvidos às releases seguintes."),
        ("2.6. Atualizações que destroem dados do cliente",
         "Release nova aplica migration destrutiva e os dados preenchidos pelo cliente se perdem ou "
         "ficam inconsistentes.",
         "Resposta: MVP 7 — Entrega versionada preservando dados do usuário. Releases não-destrutivas "
         "aplicadas automaticamente; destrutivas exigem confirmação do Admin e geram snapshot "
         "automático por projeto antes de aplicar."),
        ("2.7. Ausência de backup automático e recuperação",
         "Cliente esquece de fazer backup, ou faz sem consistência, ou perde dados por falha de disco.",
         "Resposta: backup automático diário às 12:00 de cada projeto ativo. Retenção de 10 backups "
         "por projeto. Catch-up no startup se faltou algum. Restore com validação SHA-256 e "
         "preservação de audit_log_global."),
    ]
    for title, problem, answer in problems:
        h2(doc, title)
        para(doc, problem)
        para(doc, answer, bold=True, color=EMERALD)

    page_break(doc)

    # ═══ 3. Compartimentalização ═══
    h1(doc, "3. Compartimentalização por projeto")
    para(doc,
         "O GCA adota compartimentalização dura como princípio arquitetural. Cada projeto "
         "é um compartimento isolado em seis dimensões. Falha em um projeto (ou comprometimento "
         "de chave de IA) não contamina os outros — mesmo dentro da mesma instância.")

    table(doc,
          ["Dimensão", "Como é compartimentalizado"],
          [
              ["Banco de dados",
               "Todas as tabelas com dado de projeto têm project_id no predicado (30+ tabelas). Nenhuma query cross-projeto fora de endpoints administrativos."],
              ["Provedores de IA",
               "Provedor e chave de API configuráveis por projeto; chaves criptografadas com Fernet, isoladas no vault do projeto."],
              ["Storage",
               "Uploads de ingestão, anexos de tickets e bundles prefixados por /projects/{id}/ ou /incidents/{tid}/ em gca-uploads-storage."],
              ["Backups",
               "Backup por projeto: 29 tabelas filtradas por project_id, empacotadas em zip com manifest + sha256. Retenção independente por projeto."],
              ["Auditoria",
               "audit_log_global com hash chain; logs de pipeline e operação com project_id. Filtragem por resource_id=project_id."],
              ["E-mail / SMTP",
               "SMTP compartimentalizado por projeto (servidor, credenciais, FROM). Fallback global se projeto não configurar."],
          ],
          widths=[4, 13])

    para(doc,
         "Um cliente com 50 projetos tem 50 ilhas operacionais dentro da mesma instância. "
         "O único canal cross-projeto é o endpoint administrativo /admin/*, acessível apenas "
         "a Admin e Equipe Sustentação para tickets escalados e visão agregada.",
         bold=True)

    page_break(doc)

    # ═══ 4. RBAC ═══
    h1(doc, "4. Papéis canônicos (RBAC)")
    para(doc,
         "O GCA implementa cinco papéis canônicos + papel virtual, organizados em duas "
         "camadas soberanas.")

    h2(doc, "4.1. Camada da instância")
    table(doc,
          ["Papel", "Flag", "Função"],
          [
              ["Admin", "is_admin=true",
               "Configura a instância, aprova projetos, gerencia usuários, aplica releases destrutivas. Não atua operacionalmente em projetos."],
              ["Sustentação", "is_support=true",
               "Recebe tickets escalados a Admin. Pode atuar no atendimento sem ganhar poderes administrativos. Admin herda Sustentação automaticamente."],
          ],
          widths=[3, 3, 11])

    h2(doc, "4.2. Camada do projeto")
    table(doc,
          ["Papel", "Escopo", "Função"],
          [
              ["GP", "Projeto (soberano)",
               "Pós-emenda 2026-04-19: soberano do projeto. Tem união das ações de Dev, Tester e QA. Conduz o projeto, aprova módulos e OCG, convida time, opera todos os fluxos quando necessário."],
              ["Dev", "Projeto",
               "Implementa código, opera ingestão, Arguidor, CodeGen e commits. Não aprova módulo no Gatekeeper."],
              ["Tester", "Projeto",
               "Cria, edita e executa testes. Registra evidências."],
              ["QA", "Projeto",
               "Revisa e aprova resultados. Valida qualidade final, segurança e compliance. Não edita conteúdo de teste."],
              ["admin_viewer", "Virtual",
               "Atribuído a Admin que acessa projeto sem membership. Vê projeto; não atua."],
          ],
          widths=[3, 3.5, 10.5])

    h2(doc, "4.3. Hierarquia visual")
    image(doc, DIAGRAM_DIR / "rbac_papeis.png",
          "Figura 1 — Hierarquia dos papéis canônicos pós-emenda 2026-04-19.",
          width_in=6.3)

    quote(doc, "GP : projeto :: Admin : instância")
    para(doc,
         "Admin herda Sustentação. GP nunca herda Admin automaticamente — promoção a Admin "
         "é caminho separado (/admin/users → Promover).")

    page_break(doc)

    # ═══ 5. Área Administrativa ═══
    h1(doc, "5. Área Administrativa")
    para(doc,
         "A área administrativa é acessada por usuários com is_admin=true (ou is_support=true "
         "para subset de telas). Ao logar via /login, a sidebar renderiza a seção "
         "ADMINISTRAÇÃO com nove entradas: Dashboard Global, Projetos, Usuários, Auditoria "
         "Global, Métricas, Backups, Incidentes, Equipe Sustentação e Releases.")

    admin_screens = [
        ("5.1. Login Administrador", "01_publica_login_admin.png",
         "Login via /login com e-mail e senha. Retorna JWT. Senhas armazenadas com bcrypt. "
         "Sem refresh token nesta versão.",
         "Figura 2 — Tela de login do Administrador."),
        ("5.2. Dashboard Global", "10_admin_dashboard_global.png",
         "Primeira tela após login. Agrega contadores: projetos ativos por status, total de "
         "tickets abertos, releases recentes, tendência de uso de IA últimas 24 horas.",
         "Figura 3 — Dashboard Global com KPIs cross-projeto."),
        ("5.3. Gestão de Projetos", "11_admin_gestao_projetos.png",
         "Lista todos os projetos da instância com badges de lifecycle (Ativo, Pausado, "
         "Desativado, Excluído-órfão). Ações por linha: Pausar, Desativar, Reativar, Limpar "
         "órfão, Substituir GP, Mensagem ao solicitante.",
         "Figura 4 — Gestão de Projetos com badges de lifecycle."),
        ("5.4. Visão Admin do Projeto", "12_admin_projeto_visao_admin.png",
         "Admin pode entrar em qualquer projeto (mesmo sem membership) com papel virtual "
         "admin_viewer. Vê o projeto sem poder atuar operacionalmente. Útil para fiscalizar "
         "ou diagnosticar.",
         "Figura 5 — Visão Admin entrando em projeto sem ser membro operacional."),
        ("5.5. Gestão de Usuários", "13_admin_gestao_usuarios.png",
         "Lista todos os usuários da instância. Ações: Promover/Rebaixar Admin (bloqueia auto-"
         "rebaixa do último admin ativo), Excluir, Ativar/Desativar, Convidar Administrador.",
         "Figura 6 — Gestão de Usuários com promoção de Admin."),
        ("5.6. Modal Convidar Administrador", "14_admin_convidar_admin_modal.png",
         "Modal acionado em /admin/users. Se e-mail já existe, promove a Admin sem mexer na "
         "senha. Se não existe, cria usuário novo com senha temporária aleatória e tenta enviar "
         "e-mail. Se SMTP falhar, a senha é exibida inline com aviso.",
         "Figura 7 — Modal de convite de Administrador com fallback de senha inline."),
        ("5.7. Auditoria Global", "15_admin_auditoria_global.png",
         "Todos os eventos críticos da instância (login admin, criação/aprovação de projeto, "
         "promoção/rebaixamento, aplicação de release, restore de backup) registrados em "
         "audit_log_global com hash chain.",
         "Figura 8 — Auditoria Global com hash chain de integridade."),
        ("5.8. Métricas Operacionais", "16_admin_metricas.png",
         "Dashboard de uso de IA em toda a instância com totais consolidados, uso por "
         "provider × operation, breakdown compartimentalizado POR PROJETO (nome, status, "
         "chamadas, tokens, custo ordenado por custo decrescente), eventos de audit top 20, "
         "projetos por status e usuários ativos/admin.",
         "Figura 9 — Métricas com breakdown compartimentalizado por projeto."),
        ("5.9. Backups (visão agregada)", "17_admin_backups.png",
         "Lista todos os backups de todos os projetos. Filtros: projeto, status, trigger "
         "source. Quick action: disparar backup imediato para projeto a pedido do GP.",
         "Figura 10 — Backups agregados com filtros e quick action."),
        ("5.10. Incidentes (tickets escalados)", "18_admin_incidents.png",
         "Lista agregada cross-projeto de tickets com target_scope='admin' (abertos por GPs "
         "ou por outros Admins). Filtros por status e projeto. Admin e Sustentação veem a "
         "mesma lista.",
         "Figura 11 — Tickets escalados a Admin e Sustentação."),
        ("5.11. Equipe Sustentação", "19_admin_equipe_sustentacao.png",
         "Administra a flag is_support. Lista membros atuais (Admin + Support). Promove "
         "usuário comum a Support. A UI não promove Support a Admin — esse fluxo fica na "
         "gestão canônica de usuários.",
         "Figura 12 — Equipe Sustentação com herança automática de Admin."),
        ("5.12. Releases", "20_admin_releases.png",
         "Lista de releases aplicadas e pendentes. Releases destrutivas destacadas em âmbar. "
         "Cada release mostra tag semântica (v0.8.0), título, status, número de items, data "
         "de aplicação, YAML de origem.",
         "Figura 13 — Releases com separação pendentes/aplicadas."),
        ("5.13. Detalhe de Release", "21_admin_release_detail.png",
         "Items do changelog (MVP, emenda, ticket, feature, fix, schema change), log de "
         "aplicação (applied, snapshot_taken, rolled_back), botão Aplicar com snapshot em "
         "releases destrutivas pending.",
         "Figura 14 — Detalhe da release com items e log."),
        ("5.14. Changelog (user-facing)", "25_global_changelog_user_admin.png",
         "Acessível via ícone ✨ no topbar. Lista releases aplicadas com changelog segmentado "
         "por papel. Admin vê tudo.",
         "Figura 15 — Changelog segmentado por papel (visto como Admin)."),
    ]

    for title, filename, desc, caption in admin_screens:
        h2(doc, title)
        para(doc, desc)
        image(doc, SCREENSHOT_DIR / filename, caption, width_in=5.8)

    page_break(doc)

    # ═══ 6. Área de Projeto ═══
    h1(doc, "6. Área de Projeto")
    para(doc,
         "Usuários com is_admin=false (Dev, Tester, QA, GP) entram na instância via "
         "/p/{short_slug} — ProjectLoginPage. Após autenticar, são levados direto ao "
         "/projects/{id}. A sidebar renderiza a seção MEUS PROJETOS com o projeto ativo "
         "expandido exibindo 18 sub-itens. A seção ADMINISTRAÇÃO não aparece.")

    h2(doc, "6.1. Login via projeto (ProjectLoginPage)")
    para(doc,
         "Acessível em /p/{short_slug} — o short_slug é gerado automaticamente a partir do "
         "nome do projeto. O endpoint público /projects/by-slug/{short_slug} retorna resumo "
         "do projeto (nome, status); o form de login usa o contexto do projeto para "
         "direcionar ao dashboard correto.")
    image(doc, SCREENSHOT_DIR / "02_publica_login_projeto.png",
          "Figura 16 — Tela de login via projeto, contextualizada pelo slug.",
          width_in=5.8)

    project_screens = [
        ("6.2. Lista de projetos (visão do GP)", "30_projeto_gp_lista_projetos.png",
         "Segmentada pela project_memberships do usuário: lista apenas os projetos onde ele é "
         "membro aceito. GP pode ser membro de múltiplos projetos simultaneamente.",
         "Figura 17 — Lista de projetos visível ao GP (sem menu administrativo)."),
        ("6.3. Dashboard do projeto", "31_projeto_gp_dashboard.png",
         "Entrada principal do projeto. KPIs: saúde do OCG, readiness por pilar, últimos "
         "deltas, backlog priorizado, pendências pós-release, consumo de IA do projeto. "
         "Observe na sidebar os 18 sub-itens expandidos do projeto.",
         "Figura 18 — Dashboard do projeto com sidebar MEUS PROJETOS expandida."),
        ("6.4. Equipe", "32_projeto_gp_team.png",
         "Gestão de membros do projeto. GP convida Dev, Tester, QA por e-mail. Sistema "
         "permite multi-papel (um usuário pode ser Dev e Tester simultaneamente no mesmo "
         "projeto).",
         "Figura 19 — Equipe do projeto com convites e multi-papel."),
        ("6.5. OCG — Objeto Canônico de Governança", "33_projeto_gp_ocg.png",
         "Fonte única de verdade do projeto. Gerado pelo pipeline de 8 agentes de IA a partir "
         "do questionário aprovado. Estrutura: PROJECT_PROFILE, STACK_RECOMMENDATION, "
         "ARCHITECTURE_OVERVIEW, PILLAR_SCORES (7 pilares), COMPOSITE_SCORE, TESTING, "
         "COMPLIANCE, DELIVERABLES, RISK, APPROVAL_STATUS. Toda mudança versionada em "
         "ocg_delta_log.",
         "Figura 20 — OCG com os 7 pilares avaliados."),
        ("6.6. Repositórios externos", "34_projeto_gp_external_repos.png",
         "Lista de repositórios Git vinculados ao projeto (GitHub/GitLab/Bitbucket) via PAT "
         "criptografada com Fernet. Cada repo pode disparar o Arguidor quando houver commits "
         "novos.",
         "Figura 21 — Repositórios externos com integração Git."),
        ("6.7. Ingestão", "35_projeto_gp_ingestion.png",
         "Upload de documentos (PDF, DOCX, TXT, MD) para análise. Detector de PII aplica "
         "quarentena quando detecta dados pessoais. GP pode liberar manualmente falsos-"
         "positivos. Documentos ingeridos alimentam o Arguidor.",
         "Figura 22 — Ingestão de documentos com quarentena PII."),
        ("6.8. Gatekeeper", "36_projeto_gp_gatekeeper.png",
         "Avaliação por sete pilares. Cada pilar tem thresholds parametrizáveis pelo Admin. "
         "Bloqueia avanço de fase quando pilar crítico não atinge o mínimo. Findings geram "
         "items de backlog.",
         "Figura 23 — Gatekeeper avaliando 7 pilares com thresholds."),
        ("6.9. Arguidor", "37_projeto_gp_arguider.png",
         "Módulo de análise profunda de documentos ingeridos via LLM. Gera findings "
         "classificados por severidade (BLOCKER, CRITICAL, WARNING, INFO). Usa o provedor de "
         "IA configurado no projeto.",
         "Figura 24 — Arguidor com análise LLM de documentos."),
        ("6.10. Geração de Código (CodeGen)", "38_projeto_gp_codegen.png",
         "Gera estrutura inicial do projeto com base no OCG. Suporta 7 linguagens via "
         "scaffolders determinísticos: Java/Spring, Java/Quarkus, Kotlin/Spring, Go, C#/.NET 8, "
         "PHP/Laravel 11, Node.js/NestJS. Python via LLM. Após emenda RBAC 2026-04-19, GP "
         "pode operar CodeGen.",
         "Figura 25 — Geração de Código com scaffolders para 7 linguagens."),
        ("6.11. QA Readiness", "39_projeto_gp_qa_readiness.png",
         "Estado de prontidão de QA com cobertura por pilar. Derivado automaticamente do OCG, "
         "testes executados e evidências.",
         "Figura 26 — QA Readiness com cobertura de testes por pilar."),
        ("6.12. Revisão de Testes", "40_projeto_gp_tester_review.png",
         "Workflow de aprovação: Tester registra teste executado com evidência; QA revisa e "
         "aprova ou rejeita. Toda decisão auditada.",
         "Figura 27 — Revisão de Testes (fluxo Tester → QA)."),
        ("6.13. Backlog", "41_projeto_gp_backlog.png",
         "Items derivados do OCG + findings do Arguidor, priorizados. Cada item tem tipo, "
         "descrição, pilar relacionado, status.",
         "Figura 28 — Backlog priorizado do projeto."),
        ("6.14. Roadmap", "42_projeto_gp_roadmap.png",
         "Visão temporal das entregas. Consolidação do backlog em timeline.",
         "Figura 29 — Roadmap do projeto."),
        ("6.15. Documentação Viva", "43_projeto_gp_docs.png",
         "Documentação gerada automaticamente a partir do OCG e atualizada a cada mudança. "
         "Inclui: Arquitetura, Stack, Testing Strategy, Compliance, Risk Assessment, "
         "Deliverables.",
         "Figura 30 — Documentação Viva sincronizada com OCG."),
        ("6.16. Definition of Done (Readiness)", "44_projeto_gp_readiness.png",
         "Checklist de critérios de entrega final. Quando todos os critérios atendidos, "
         "libera geração do Release Bundle (pacote consolidado de artefatos).",
         "Figura 31 — Definition of Done, critérios de Release Bundle."),
        ("6.17. Configurações do projeto", "45_projeto_gp_settings.png",
         "Três tabs: Questionário (upload/re-upload do PDF), Repositório (Git + PAT), "
         "Provedor de IA (configuração do LLM do projeto, override do provedor global; "
         "aceita Ollama com base_url).",
         "Figura 32 — Configurações do projeto com 3 tabs."),
        ("6.18. Pipeline Audit", "46_projeto_gp_audit.png",
         "Logs de auditoria específicos do projeto. Filtros por módulo (OCG, Ingestão, "
         "Arguidor, CodeGen, QA, Release) e por ator.",
         "Figura 33 — Pipeline Audit com filtros por módulo."),
        ("6.19. Backups do projeto", "47_projeto_gp_backups.png",
         "Lista de backups do projeto. Botão Backup agora (Admin ou GP). Botão Reverter em "
         "cada backup (restore com confirmação). Download .zip. Polling a cada 3 segundos "
         "quando existe backup em running.",
         "Figura 34 — Backups do projeto com rollback e download."),
        ("6.20. Incidentes do projeto", "48_projeto_gp_incidents.png",
         "Tickets abertos pelo time do projeto. Filtros por status e prioridade. Listagem "
         "segmentada: Dev, Tester e QA veem apenas os próprios tickets; GP vê todos os "
         "tickets target_scope='gp' do projeto.",
         "Figura 35 — Incidentes do projeto com segmentação por papel."),
        ("6.21. Modal de abertura de ticket", "49_projeto_gp_abrir_ticket_modal.png",
         "Formulário com título, categoria, prioridade, descrição, seção onde o erro ocorreu "
         "(autopreenchida pela rota atual, editável), fluxo executado obrigatório (modal "
         "recusa vazio), anexos (5 arquivos, 10 MB, 9 tipos). Roteamento automático por papel.",
         "Figura 36 — Modal Abrir Ticket com seção/fluxo obrigatórios e anexos."),
        ("6.22. Métricas do Projeto", "50_projeto_gp_metrics.png",
         "Dashboard operacional do projeto: uso de IA (chamadas, tokens, custo), eventos de "
         "audit do projeto. Acessível a Admin, Sustentação e qualquer membro aceito. "
         "Seletor de janela 24h / 7d / 30d.",
         "Figura 37 — Métricas compartimentalizadas do projeto."),
        ("6.23. Changelog visto pelo GP", "55_global_gp_changelog_user.png",
         "Acessível pelo ícone ✨ no topbar. O mesmo endpoint /releases segmenta o changelog "
         "pelo papel do usuário. GP vê items relevantes ao projeto dele.",
         "Figura 38 — Changelog visto pelo GP com segmentação por papel."),
    ]

    for title, filename, desc, caption in project_screens:
        h2(doc, title)
        para(doc, desc)
        image(doc, SCREENSHOT_DIR / filename, caption, width_in=5.8)

    page_break(doc)

    # ═══ 7. Diagramas de sequência ═══
    h1(doc, "7. Diagramas de sequência")
    para(doc,
         "Os seis diagramas abaixo formalizam os fluxos críticos do GCA. Todos foram "
         "renderizados a partir de código Mermaid, que permanece disponível em "
         "docs/diagrams/*.mmd para manutenção.")

    sequences = [
        ("7.1. Login Administrador",
         "Fluxo padrão de autenticação administrativa. JWT com expiração configurável é o "
         "único token emitido nesta versão — não há refresh token.",
         "seq_login_admin.png", "Figura 39 — Login administrativo."),
        ("7.2. Login via projeto (ProjectLoginPage)",
         "Entrada contextualizada em /p/{short_slug}: o frontend busca resumo do projeto e "
         "apresenta form de login com o nome do projeto no topo; após autenticar, valida "
         "membership e redireciona direto ao dashboard do projeto.",
         "seq_login_projeto.png", "Figura 40 — Login via slug de projeto."),
        ("7.3. Criação e aprovação de projeto",
         "Da solicitação pública até a aprovação pelo Admin e promoção automática do "
         "solicitante a GP do projeto recém-criado.",
         "seq_criar_projeto.png", "Figura 41 — Criação e aprovação de projeto."),
        ("7.4. Geração do OCG (pipeline de 8 agentes)",
         "Pipeline em cascata com 7 pilares em paralelo: analyzer → P1..P7 (paralelo) → "
         "consolidator. Escolha de provedor por criticidade (alta → premium). Registro em "
         "ai_usage_log.",
         "seq_ocg_generation.png", "Figura 42 — Geração do OCG em 8 agentes."),
        ("7.5. Rastreabilidade Ticket → Release",
         "Conforme MVP 6 + MVP 7. A release seguinte à resolução do ticket amarra-se via "
         "ref_id=TICKET-{id}. Snapshot automático antes de aplicar quando release é "
         "destrutiva.",
         "seq_ticket_release.png", "Figura 43 — Rastreabilidade Ticket → Release."),
        ("7.6. Backup diário e restore",
         "Scheduler às 12:00 (tz America/Sao_Paulo), catch-up no startup se última execução "
         "> 24h, restore com validação SHA-256 preservando audit_log_global.",
         "seq_backup_restore.png", "Figura 44 — Ciclo de backup e restore (DT-063)."),
    ]

    for title, desc, filename, caption in sequences:
        h2(doc, title)
        para(doc, desc)
        image(doc, DIAGRAM_DIR / filename, caption, width_in=6.3)

    page_break(doc)

    # ═══ 8. Diagramas de fluxo e estado ═══
    h1(doc, "8. Diagramas de fluxo e estado")

    flows = [
        ("8.1. Ciclo de vida do projeto",
         "Projeto passa por quatro estágios principais: solicitado (externo), pendente "
         "(aguardando Admin), aprovado (linha em projects). A partir de aprovado, o Admin "
         "pode transitar entre active, paused e inactive sem nunca apagar dados.",
         "flow_projeto_lifecycle.png", "Figura 45 — Estados do projeto."),
        ("8.2. Ciclo de vida do ticket",
         "Aberto → em andamento → resolvido → fechado. Tickets resolvidos podem ser reabertos. "
         "Tickets fechados entram no histórico permanente.",
         "flow_ticket_lifecycle.png", "Figura 46 — Estados do ticket de incidente."),
        ("8.3. Ciclo de vida da release",
         "Declarada (YAML) → pending (sync no startup) → applied (auto ou manual). Releases "
         "destrutivas exigem confirmação e snapshot antes de aplicar. Rollback é por-projeto.",
         "flow_release_lifecycle.png", "Figura 47 — Estados da release."),
    ]

    for title, desc, filename, caption in flows:
        h2(doc, title)
        para(doc, desc)
        image(doc, DIAGRAM_DIR / filename, caption, width_in=6.3)

    page_break(doc)

    # ═══ 9. Segurança e proteção de código ═══
    h1(doc, "9. Segurança e proteção de código")

    h2(doc, "9.1. Princípio honesto")
    para(doc,
         "Não existe proteção absoluta contra engenharia reversa em código que roda na "
         "máquina do cliente. O GCA implementa sete camadas que elevam o custo de reversão "
         "ao ponto de torná-la economicamente inviável para cliente médio. Quem tem tempo, "
         "ferramentas profissionais e motivação, sempre pode reverter — só proteção em "
         "hardware (HSM, secure enclave) atingiria garantia forte, fora do escopo desta versão.")

    h2(doc, "9.2. As sete camadas")
    table(doc,
          ["#", "Camada", "O que faz", "Ganho"],
          [
              ["1", "Cython compile",
               "Python .py → .c → binário .so (Linux) / .pyd (Windows).",
               "Decompile via uncompyle6 não funciona; só Ghidra/IDA, que produz pseudo-C ilegível."],
              ["2", "PyArmor BCC wrapper",
               "Wrapper adicional em módulos sensíveis (auth, vault, LLM key resolver).",
               "Runtime check; módulo só carrega no contexto esperado."],
              ["3", "Imagens Docker multi-stage",
               "Stage builder separado do runtime; sem toolchain na imagem final.",
               "Atacante com shell no container não encontra gcc/g++ para recompilar."],
              ["4", "Registry privado autenticado",
               "Imagens em registry.gca-produto.com com tokens rotacionáveis.",
               "Quem não é autorizado não baixa."],
              ["5", "JavaScript-obfuscator frontend",
               "Control-flow flattening + string array base64 + dead code injection.",
               "Bundle JS vira ilegível (_0x5e4f, switch-case com estado, strings em array)."],
              ["6", "Integrity check no startup",
               "Backend calcula SHA-256 dos .so em runtime e compara com manifest assinado.",
               "Modificação em qualquer binário quebra o SHA e aborta startup."],
              ["7", "Licença JWT com expiração",
               "JWT assinado com chave privada, validado no startup.",
               "Sem JWT válido, aplicação não sobe; expiração atingida bloqueia até renovar."],
          ],
          widths=[0.8, 3.5, 5.5, 6])

    h2(doc, "9.3. O que não é feito (transparência)")
    bullet(doc, "Sem atestação remota (remote attestation).")
    bullet(doc, "Sem secure enclave (Intel SGX, ARM TrustZone).")
    bullet(doc, "Sem criptografia do .so em repouso — só o código-fonte deixa de ser Python legível.")
    bullet(doc, "JWT da licença pode ser capturado por atacante com acesso root ao host (defesa: responsabilidade do cliente, EULA item 3.1).")

    page_break(doc)

    # ═══ 10. O que faz / o que não faz ═══
    h1(doc, "10. O que o GCA faz / o que ele ainda não faz")

    h2(doc, "10.1. O que o GCA faz (V1)")
    para(doc, "Lista completa das funcionalidades entregues nos MVPs 1-7 mais emenda:", bold=True)

    h3(doc, "Gestão de projeto")
    bullet(doc, "Cadastro e aprovação de projetos via wizard externo + aprovação admin.")
    bullet(doc, "Multi-projeto por instância, compartimentalizado.")
    bullet(doc, "Lifecycle (ativo, pausado, desativado) sem perda de dados.")
    bullet(doc, "Limpeza de solicitações órfãs.")

    h3(doc, "RBAC")
    bullet(doc, "Cinco papéis canônicos (Admin, GP, Dev, Tester, QA) + Sustentação via flag.")
    bullet(doc, "Multi-papel por projeto.")
    bullet(doc, "GP soberano do projeto (emenda 2026-04-19).")
    bullet(doc, "Anti-órfão no último Admin.")

    h3(doc, "Inteligência Artificial")
    bullet(doc, "Provedor configurável por instância e por projeto.")
    bullet(doc, "Suporte: Anthropic Claude, OpenAI GPT, Google Gemini, DeepSeek, Ollama local.")
    bullet(doc, "Roteamento por criticidade (alta → premium; baixa → local/barato).")
    bullet(doc, "Registro de uso em ai_usage_log (tokens, custo, operação) por projeto.")

    h3(doc, "Pipeline de projeto")
    bullet(doc, "Questionário de 49 perguntas (PDF AcroForm).")
    bullet(doc, "Pipeline de 8 agentes de IA gerando OCG.")
    bullet(doc, "Gatekeeper com 7 pilares e thresholds parametrizáveis.")
    bullet(doc, "Ingestão de documentos com quarentena PII.")
    bullet(doc, "Arguidor com análise LLM.")
    bullet(doc, "Backlog derivado do OCG e findings.")
    bullet(doc, "Roadmap.")
    bullet(doc, "CodeGen com scaffolders em 7 linguagens.")
    bullet(doc, "QA Readiness + revisão Tester/QA.")
    bullet(doc, "Documentação Viva sincronizada com OCG.")
    bullet(doc, "Release Bundle ao fim do projeto.")

    h3(doc, "Operação")
    bullet(doc, "Backup automático diário às 12:00 por projeto.")
    bullet(doc, "Retenção de 10 backups por projeto.")
    bullet(doc, "Catch-up no startup.")
    bullet(doc, "Restore com validação SHA-256 e preservação de auditoria.")
    bullet(doc, "Upgrade idempotente com 9 etapas.")

    h3(doc, "Tickets de incidente")
    bullet(doc, "Roteamento automático por papel (Dev/Tester/QA → GP; GP → Admin + Sustentação).")
    bullet(doc, "Seção autopreenchida pela rota.")
    bullet(doc, "Fluxo executado obrigatório.")
    bullet(doc, "Anexos (5 arquivos, 10 MB, 9 tipos).")
    bullet(doc, "Comentários, mudança de status, notificação in-app.")
    bullet(doc, "Auditoria compartimentalizada.")

    h3(doc, "Releases")
    bullet(doc, "Declarativas via YAML shipado com o código.")
    bullet(doc, "Aplicação automática de não-destrutivas no startup.")
    bullet(doc, "Destrutivas exigem confirmação + snapshot automático por projeto.")
    bullet(doc, "Rollback por-projeto preservando auditoria.")
    bullet(doc, "Changelog segmentado por papel.")
    bullet(doc, "Rastreabilidade ticket → release.")

    h3(doc, "Observabilidade")
    bullet(doc, "/metrics/health público para load balancer.")
    bullet(doc, "/metrics/dashboard agregado global.")
    bullet(doc, "/metrics/per-project breakdown por projeto.")
    bullet(doc, "/projects/{id}/metrics/dashboard por projeto.")
    bullet(doc, "/metrics/prometheus texto para scrape externo.")
    bullet(doc, "Audit log global com hash chain.")

    h3(doc, "Segurança")
    bullet(doc, "Sete camadas anti-engenharia-reversa documentadas.")
    bullet(doc, "Credenciais criptografadas com Fernet (PATs, LLM keys).")
    bullet(doc, "Senhas em bcrypt com salt único.")
    bullet(doc, "JWT HS256 com chave por instância.")
    bullet(doc, "Auditoria com integridade verificável.")

    h3(doc, "Distribuição")
    bullet(doc, "Instalador Windows via Inno Setup (10 telas + conclusão).")
    bullet(doc, "Instalador Ubuntu via install.sh interativo ou pacote .deb.")
    bullet(doc, "Script de build para imagens de produção.")

    h2(doc, "10.2. O que o GCA ainda não faz (roadmap)")
    para(doc,
         "Recursos mapeados como candidatos a MVPs futuros. Não há data nem compromisso "
         "contratual; dependem de autorização formal do stakeholder do produto.", italic=True)

    table(doc,
          ["Item", "Por quê ficou fora do V1"],
          [
              ["Atestação remota (HSM / TPM)",
               "Fora do escopo. Clientes que exigem classe governo/defesa contratam extensão."],
              ["SSO corporativo (SAML, OIDC, LDAP)",
               "Planejado para MVP 8 se o stakeholder autorizar. Login hoje é local com bcrypt."],
              ["Refresh token JWT",
               "JWT atual tem expiração fixa; re-login manual quando expira."],
              ["SLA com escalonamento automático",
               "Escalonamento é manual (GP → Admin via abrir novo ticket). Timeline SLA visível mas sem gatilho."],
              ["Anexos com scan de PII",
               "Anexos permitidos; responsabilidade de conteúdo sensível é do autor (declarado no EULA)."],
              ["Integração Jira/Linear/Zendesk",
               "Tickets hoje ficam internos. Bridge para ferramentas externas exige projeto dedicado."],
              ["E-mail bidirecional",
               "Notificações saem via SMTP; respostas por e-mail não voltam para o sistema."],
              ["Downgrade do container",
               "upgrade.sh só vai pra frente. Downgrade requer restore.sh + imagem anterior manualmente."],
              ["Compartilhamento de correção entre instâncias",
               "Cada cliente recebe release pelo fluxo de instalação. Sem sincronização central."],
              ["Marketplace de plugins",
               "Escopo fora da V1."],
              ["Multi-língua do frontend",
               "UI hoje é 100% em português brasileiro. Internacionalização requer ajuste dedicado."],
              ["Assinatura digital X.509 no Release Bundle",
               "Hash SHA-256 presente; PKI formal é projeto próprio."],
              ["Auto-upgrade totalmente autônomo",
               "Upgrade é acionado manualmente; sem daemon autorrolando."],
              ["Dashboards customizáveis por usuário",
               "Dashboards hoje são fixos."],
              ["Exportação de relatórios em PDF/Excel",
               "Consulta via UI; geração PDF/XLSX é projeto futuro."],
              ["Billing / cobrança por uso",
               "Uso de IA é registrado mas não há módulo de cobrança. Cliente lê custo em /admin/metrics."],
              ["Multi-instância federada",
               "Cada instância é soberana. Federação não está no V1."],
          ],
          widths=[6, 11])

    h2(doc, "10.3. Decisões arquiteturais conscientes (não são dívida)")
    bullet(doc, "Python não tem scaffolder determinístico — é gerado via LLM pela maturidade do ecossistema.")
    bullet(doc, "Backup usa JSONL via row_to_json (não pg_dump) porque pg_dump não filtra por WHERE — e compartimentalização exige filtro.")
    bullet(doc, "Comentários de código em português-BR por decisão de idioma do produto.")
    bullet(doc, "Admin não atua operacionalmente em projetos por separação de responsabilidades.")

    page_break(doc)

    # ═══ 11. Roadmap futuro ═══
    h1(doc, "11. Roadmap futuro")
    para(doc, "Sem data nem compromisso contratual. Candidatos a MVPs futuros sujeitos a "
              "decisão do stakeholder.", italic=True)

    table(doc,
          ["MVP futuro", "Tema"],
          [
              ["MVP 8", "SSO corporativo (SAML 2.0 + OpenID Connect + fallback local)"],
              ["MVP 9", "Relatórios exportáveis (PDF do OCG, Excel do backlog, PDF do Release Bundle)"],
              ["MVP 10", "Integração externa de tickets (bridge Jira, Linear, Zendesk)"],
              ["MVP 11", "Marketplace de plugins (extensões configuráveis por instância)"],
              ["MVP 12", "Federação entre instâncias (benchmarks, trocas autorizadas)"],
              ["MVP 13", "Assinatura digital de artefatos (PKI X.509 para Release Bundle)"],
              ["MVP 14", "Internacionalização (pt-BR + en-US + es no V2)"],
          ],
          widths=[4, 13])

    para(doc,
         "Protocolo de adição de MVP (contrato §7.0): stakeholder-soberano autoriza, time de "
         "produto atualiza contrato + progress em commit atômico, MVP nasce com escopo "
         "declarado e só é executado com autorização explícita de início.")

    page_break(doc)

    # ═══ 12. Glossário ═══
    h1(doc, "12. Glossário")
    para(doc, "Termos em ordem alfabética. Referência cruzada ao contrato canônico quando "
              "aplicável.", size=10, color=SLATE_MEDIUM)

    glossary = [
        ("Admin", "Usuário da instância com is_admin=true. Soberano da instância; configura provedores, aprova projetos, gerencia usuários, aplica releases destrutivas. Não atua operacionalmente em projetos."),
        ("admin_viewer", "Papel virtual atribuído a Admin quando acessa projeto sem membership. Permite ver o projeto mas não editar conteúdo."),
        ("ai_usage_log", "Tabela que registra cada chamada a provedor de IA com provider, operation, tokens_input, tokens_output, cost_usd. Compartimentalizada por project_id."),
        ("Anthropic", "Provedor de IA (Claude). Recomendado para tarefas de alta criticidade (consolidação do OCG, CodeGen)."),
        ("APScheduler", "Biblioteca Python de agendamento assíncrono. Dispara o backup diário às 12:00 e o catch-up no startup."),
        ("Arguidor", "Agente de IA que analisa documentos ingeridos e gera findings classificados (BLOCKER, CRITICAL, WARNING, INFO)."),
        ("audit_log_global", "Tabela de auditoria com hash chain (previous_hash + current_hash). Integridade verificável offline. Preservada mesmo em restore."),
        ("Backup", "Cópia consistente dos dados do projeto. Gerenciado pelo project_backup_service (DT-063). Até 10 retidos por projeto."),
        ("Backlog", "Lista de items derivados do OCG e findings do Arguidor, priorizados."),
        ("BCC", "BytecodeCompiler do PyArmor (variante gratuita) usada no empacotamento do GCA."),
        ("Bcrypt", "Algoritmo de hash de senhas utilizado no GCA com salt único por usuário."),
        ("Caddy", "Servidor web com HTTPS automático via Let's Encrypt. Recomendado como proxy reverso."),
        ("Catch-up", "Rotina do scheduler que dispara backup se o último foi há mais de 24 horas."),
        ("Changelog", "Lista de items de uma release com kind (mvp, feature, fix, ticket) e affected_roles. Segmentado por papel na UI /releases."),
        ("Chave de ativação", "Token GCA-PROD-XXXXX-XXXXX-XXXXX-XXXXX fornecido na contratação. Determina validade e limites."),
        ("CodeGen", "Módulo de geração assistida de código (MVP 3). Usa OCG + provedor de IA + scaffolders determinísticos por linguagem."),
        ("Completion task", "Pendência pós-release por projeto quando release adiciona campo novo obrigatório."),
        ("Compartimentalização", "Princípio do contrato §2.2: todo acesso a dado de projeto inclui project_id no predicado."),
        ("Container", "Unidade de deployment isolada via Docker. O GCA usa seis containers."),
        ("Contexto A", "Uso de IA para desenvolvimento do produto GCA em si. Pode usar IA premium. Custo do time de produto."),
        ("Contexto B", "Uso de IA pelo cliente dentro da instância dele. Configurável por instância e por projeto."),
        ("Cython", "Compilador que transforma Python em C e depois em binário .so/.pyd. Usado na proteção de código."),
        ("DeepSeek", "Provedor de IA de baixo custo. Adequado a tarefas simples; não recomendado para OCG consolidado."),
        ("Dev", "Papel canônico do projeto. Implementa código, opera ingestão, Arguidor, CodeGen e commits."),
        ("Docker", "Runtime de containers usado pelo GCA. Versão mínima 24."),
        ("Docker Compose", "Orquestrador multi-container definido em docker-compose.yml."),
        ("Dogfood", "Prática de usar a própria instância do GCA para conduzir o desenvolvimento do GCA. O projeto Automação Jurídica Assistida é o dogfood ativo."),
        ("DT", "Dívida técnica. Identificada por número (DT-001 a DT-063)."),
        ("Emenda 2026-04-19", "Alteração formal no contrato canônico §4.1 (GP soberano do projeto) e §7 MVP 6 (Sustentação, anexos, seção/fluxo)."),
        ("Equipe Sustentação", "Conjunto de usuários com is_support=true (ou is_admin=true via herança). Recebe tickets com target_scope=admin."),
        ("EULA", "End User License Agreement — contrato de licença aceito no Passo 2 da instalação."),
        ("Fernet", "Algoritmo de criptografia simétrica autenticada. Usado no GCA para PATs e chaves de API de provedores."),
        ("Gatekeeper", "Módulo de avaliação baseado em sete pilares (Business, Architecture, Stack, Testing, Compliance, Risk, Deliverables)."),
        ("GCA", "Gestão de Codificação Assistida / Gerenciador Central de Arquiteturas."),
        ("GCA_MASTER_KEY", "Chave mestra da instância. Gerada durante a instalação; usada pelo Fernet."),
        ("GP", "Gerente de Projeto. Soberano do projeto pós-emenda 2026-04-19. Tem união das ações de Dev + Tester + QA + exclusivas."),
        ("Health check", "Endpoint /api/v1/metrics/health que retorna 200 quando o backend está operacional."),
        ("Hash chain", "Cadeia de hashes em audit_log_global onde cada registro tem previous_hash apontando para o current_hash do anterior."),
        ("incident_tickets", "Tabela de tickets de incidente (MVP 6). Campos target_scope, category, priority, status, title, description, section_reference, flow_description."),
        ("Ingestão", "Módulo que recebe documentos, detecta PII, quarentena quando aplicável, aciona Arguidor."),
        ("Inno Setup", "Ferramenta open-source para gerar instaladores Windows .exe. Usada no GCA."),
        ("install.sh", "Script interativo de instalação em Ubuntu."),
        ("is_admin", "Flag em users. Habilita acesso à camada administrativa da instância."),
        ("is_support", "Flag em users que habilita recebimento de tickets escalados a Admin. Independente de is_admin."),
        ("JWT", "JSON Web Token. Único tipo de token de autenticação nesta versão. Assinado com HS256."),
        ("Lifecycle", "Ciclo de vida. Aplicado a projetos (active/paused/inactive), tickets (open/in_progress/resolved/closed) e releases (pending/applied/rolled_back)."),
        ("LLM", "Large Language Model. Provedor de IA usado pelo GCA."),
        ("manifest.json", "Arquivo dentro de cada backup e cada release YAML que lista conteúdo, contagens e hashes."),
        ("Mermaid", "Linguagem de marcação para diagramas (sequência, fluxo, estado)."),
        ("Migration", "Script SQL aplicado ao banco durante upgrade. Nunca destrutivo por default (MVP 7)."),
        ("MVP", "Minimum Viable Product. No GCA, unidade de escopo canônico (MVP 1 a 7)."),
        ("Multi-stage", "Técnica Docker de separar build e runtime. Usada nas imagens de produção."),
        ("OCG", "Objeto Canônico de Governança. Fonte única de verdade do projeto. Gerado pelo pipeline de 8 agentes."),
        ("Ollama", "Provedor de IA local (self-hosted). Suportado via endpoint OpenAI-compatible."),
        ("Órfão", "Registro em project_requests com status=APPROVED cuja linha em projects não existe mais."),
        ("PAT", "Personal Access Token. Credencial de integração com Git. Criptografada com Fernet."),
        ("Pilares", "Sete categorias do Gatekeeper: P1 Business, P2 Architecture, P3 Stack, P4 Testing, P5 Compliance, P6 Risk, P7 Deliverables."),
        ("PII", "Personally Identifiable Information. O GCA detecta e quarentena documentos com PII na ingestão."),
        ("Playwright", "Framework de automação de navegador usado para captura das 40 telas da aplicação."),
        ("PostgreSQL", "Banco de dados relacional usado pelo GCA (versão 15)."),
        ("project_id", "Chave de compartimentalização. Obrigatória em todo predicado que acessa dado de projeto."),
        ("project_requests", "Tabela de solicitações externas de projeto. Estados: PENDING, APPROVED, REJECTED."),
        ("ProjectLoginPage", "Tela de login contextualizada para usuários não-admin em /p/{short_slug}."),
        ("PyArmor", "Ferramenta de obfuscação de Python. Variante BCC (gratuita) é a usada no GCA."),
        ("QA", "Quality Assurance. Papel canônico do projeto. Revisa e aprova resultados. Não edita conteúdo de teste."),
        ("Questionnaire", "Questionário de 49 perguntas (PDF AcroForm) que inicia o pipeline do projeto."),
        ("RBAC", "Role-Based Access Control. Modelo de autorização do GCA."),
        ("Registry", "Repositório de imagens Docker. O GCA usa registry privado autenticado."),
        ("Release", "Unidade de entrega de software para a instância. Declarada em backend/releases/*.yaml."),
        ("Release Bundle", "Pacote consolidado de artefatos entregues ao fim do MVP 4."),
        ("Rollback", "Restauração de dados de projeto via snapshot prévio. Por-projeto."),
        ("Scaffolder", "Gerador determinístico de estrutura inicial de projeto por linguagem."),
        ("Scheduler", "APScheduler — dispara backups diários às 12:00 (America/Sao_Paulo)."),
        ("SHA-256", "Algoritmo de hash usado para verificar integridade de backups e manifestos de release."),
        ("short_slug", "Versão curta do slug do projeto usada no URL /p/{short_slug}."),
        ("Smoke test", "Verificação rápida pós-upgrade de que o sistema está operacional."),
        ("SMTP", "Protocolo de envio de e-mail. O GCA usa SMTP compartimentalizado por projeto."),
        ("Snapshot", "Backup prévio automático antes de aplicação de release destrutiva."),
        ("Solicitante", "Usuário externo que submete /solicitar-projeto. Quando o projeto é aprovado, torna-se GP automaticamente."),
        ("Sustentação", "Ver Equipe Sustentação."),
        ("target_scope", "Campo em incident_tickets. Valores: gp (ticket vai para GPs) ou admin (vai para Admins + Sustentação)."),
        ("Tester", "Papel canônico do projeto. Cria, edita e executa testes; registra evidências."),
        ("Thresholds", "Valores limites do Gatekeeper por pilar. Parametrizáveis em /admin (Settings)."),
        ("Ticket", "Registro de incidente aberto por usuário do projeto."),
        ("Ubuntu", "Distribuição Linux suportada nativamente (22.04+)."),
        ("Upgrade", "Processo de atualização do GCA via scripts/upgrade.sh."),
        ("Uvicorn", "Servidor ASGI que executa o FastAPI do GCA."),
        ("Vite", "Ferramenta de build para frontend React. Em produção usa vite preview."),
        ("Volume", "Área de armazenamento persistente do Docker. Três volumes: gca-postgres-data, gca-uploads-storage, gca-backups."),
        ("Windows", "Sistema operacional suportado via instalador Inno Setup (.exe)."),
        ("WSL 2", "Windows Subsystem for Linux versão 2. Obrigatório para Docker Desktop em Windows."),
        ("YAML", "Formato de declaração de releases em backend/releases/*.yaml."),
        ("ZIP", "Formato usado para empacotar backups por projeto."),
    ]

    for term, definition in sorted(glossary, key=lambda x: x[0].lower()):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        t_run = p.add_run(f"{term}. ")
        t_run.bold = True
        t_run.font.size = Pt(11)
        t_run.font.color.rgb = VIOLET
        d_run = p.add_run(definition)
        d_run.font.size = Pt(11)

    doc.add_paragraph()
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = final.add_run("— Fim do documento —")
    fr.italic = True
    fr.font.color.rgb = SLATE_MEDIUM
    fr.font.size = Pt(10)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    build(doc)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    size_mb = OUT_PATH.stat().st_size / 1024 / 1024
    print(f"✓ Documento gerado: {OUT_PATH}")
    print(f"  Tamanho: {size_mb:.1f} MB")


if __name__ == "__main__":
    main()
