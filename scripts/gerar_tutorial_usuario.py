#!/usr/bin/env python3
"""
Gera o tutorial de uso do GCA em formato .docx.

Saída: /home/luiz/GCA/GCA_Tutorial_Usuario.docx

Estrutura:
    1. O que é o GCA
    2. Conceitos fundamentais
    3. Tutorial passo a passo (10 jornadas)
    4. O que o GCA FAZ
    5. O que o GCA NÃO FAZ
    6. Perguntas e Respostas (FAQ)
    7. Glossário
    8. Suporte e versão
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

OUTPUT = Path("/home/luiz/GCA/GCA_Tutorial_Usuario.docx")

# ────────────────────────────────────────────────────────────────────
# Helpers de formatação
# ────────────────────────────────────────────────────────────────────

VIOLET = RGBColor(0x7C, 0x3A, 0xED)
SLATE_DARK = RGBColor(0x1E, 0x29, 0x3B)
SLATE_MID = RGBColor(0x47, 0x55, 0x69)
EMERALD = RGBColor(0x05, 0x96, 0x69)
AMBER = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)


def set_default_font(doc: Document, family: str = "Calibri", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)
    style.font.color.rgb = SLATE_DARK


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = VIOLET


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = VIOLET


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = SLATE_DARK


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_callout(doc: Document, label: str, text: str, color: RGBColor = AMBER) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    label_run = p.add_run(f"[{label}] ")
    label_run.bold = True
    label_run.font.color.rgb = color
    body_run = p.add_run(text)
    body_run.font.color.rgb = SLATE_DARK


def add_qa(doc: Document, question: str, answer: str) -> None:
    pq = doc.add_paragraph()
    qrun = pq.add_run(f"P: {question}")
    qrun.bold = True
    qrun.font.color.rgb = VIOLET
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    pa.paragraph_format.space_after = Pt(10)
    arun = pa.add_run(f"R: {answer}")
    arun.font.color.rgb = SLATE_DARK


def add_table_two_cols(doc: Document, header: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid Accent 4"
    hdr = table.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    for i, (k, v) in enumerate(rows, start=1):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = k
        c1.text = v


# ────────────────────────────────────────────────────────────────────
# Conteúdo
# ────────────────────────────────────────────────────────────────────


def build_document() -> Document:
    doc = Document()
    set_default_font(doc)

    # ── Capa ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("GCA")
    t_run.bold = True
    t_run.font.size = Pt(40)
    t_run.font.color.rgb = VIOLET

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("Gerenciador Central de Arquiteturas")
    s_run.italic = True
    s_run.font.size = Pt(16)
    s_run.font.color.rgb = SLATE_MID

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline.add_run("Tutorial de Uso para o Usuário")
    tag_run.font.size = Pt(14)
    tag_run.font.color.rgb = SLATE_DARK

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m_run = meta.add_run(f"Versão do documento: {datetime.now().strftime('%Y-%m-%d')}")
    m_run.font.size = Pt(10)
    m_run.font.color.rgb = SLATE_MID

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m2_run = meta2.add_run("GCA — plataforma de governança e geração assistida de software")
    m2_run.font.size = Pt(10)
    m2_run.font.color.rgb = SLATE_MID

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 1. O que é o GCA
    # ════════════════════════════════════════════════════════════════
    add_h1(doc, "1. O que é o GCA")

    add_p(
        doc,
        "O GCA (Gerenciador Central de Arquiteturas) é uma plataforma de governança "
        "e geração assistida de software com IA. Ele organiza o ciclo de vida completo "
        "de um projeto de TI — da solicitação inicial até a entrega de código, testes e "
        "documentação — garantindo que decisões técnicas, de negócio e de compliance "
        "sejam registradas, avaliadas e auditáveis."
    )

    add_p(
        doc,
        "Em uma frase: o GCA recebe a descrição de um projeto, conduz uma análise "
        "rigorosa em sete dimensões (os Pilares), produz um documento vivo de contexto "
        "(o OCG), gera código sob governança, e mantém a documentação atualizada a "
        "cada mudança."
    )

    add_h2(doc, "Para quem é?")
    add_bullet(doc, "Administradores de TI que precisam orquestrar múltiplos projetos com rastreabilidade.")
    add_bullet(doc, "Gerentes de Projeto (GP) que querem governar entregas sem perder velocidade.")
    add_bullet(doc, "Times de desenvolvimento que aceitam apoio de IA mas com aprovação humana e auditoria.")
    add_bullet(doc, "Times de QA, Compliance e Stakeholders que precisam visibilidade sem editar código.")

    add_h2(doc, "Filosofia central")
    add_p(
        doc,
        "Tudo o que o GCA produz precisa ser explicável, reversível e aprovável por um "
        "humano. A IA acelera, mas nunca decide sozinha. Cada arquivo gerado, cada "
        "decisão de stack, cada item do backlog tem origem rastreável e responsável "
        "designado."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 2. Conceitos fundamentais
    # ════════════════════════════════════════════════════════════════
    add_h1(doc, "2. Conceitos fundamentais")
    add_p(doc, "Antes de usar, vale conhecer cinco ideias que aparecem em quase toda tela.")

    add_h2(doc, "2.1 OCG — Objeto de Contexto Global")
    add_p(
        doc,
        "O OCG é o documento vivo do seu projeto. Ele começa pequeno — só com as "
        "respostas do questionário inicial — e cresce a cada nova ingestão de "
        "documento, repositório externo, requisito ou validação. É a fonte única "
        "de verdade que alimenta todos os outros módulos: o Gatekeeper avalia o OCG, "
        "o Backlog deriva do OCG, o CodeGen consulta o OCG."
    )
    add_callout(
        doc, "Importante",
        "O OCG NÃO é uma planilha que você preenche. É construído automaticamente "
        "pelo pipeline conforme você fornece subsídios. Ingestão ruim contrai o OCG; "
        "ingestão boa expande."
    )

    add_h2(doc, "2.2 Os 7 Pilares")
    add_p(doc, "Toda decisão técnica passa por avaliação em sete dimensões:")
    pilares = [
        ("P1 — Negócio", "Modelo de negócio, ROI, KPIs, monetização, receita esperada."),
        ("P2 — Compliance", "LGPD, regulatório setorial (BACEN, ANS, ANATEL), políticas de dados."),
        ("P3 — Escopo", "Funcional vs não-funcional, MVP, fora de escopo, dependências."),
        ("P4 — NFR (Requisitos Não-Funcionais)", "Performance, escalabilidade, disponibilidade, SLA."),
        ("P5 — Arquitetura", "Stack escolhida, padrões, dependências, infraestrutura."),
        ("P6 — Dados", "Modelagem, ETL, governança de dados, retenção."),
        ("P7 — Segurança", "Autenticação, autorização, criptografia, auditoria, OWASP."),
    ]
    for nome, desc in pilares:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{nome}: ")
        run.bold = True
        p.add_run(desc)
    add_callout(
        doc, "Bloqueante",
        "Score abaixo de 70% nos pilares de Segurança (P7) ou Compliance (P2) "
        "trava o pipeline. O Admin pode parametrizar outros thresholds em "
        "Admin → Configurações."
    )

    add_h2(doc, "2.3 Pipeline de 8 agentes IA")
    add_p(
        doc,
        "Para gerar o OCG, o GCA aciona um sistema de oito agentes especializados "
        "trabalhando em paralelo:"
    )
    add_bullet(doc, "Analyzer — decompõe o problema do projeto em subproblemas.")
    add_bullet(doc, "7 especialistas — um por pilar (P1 a P7), cada um avalia sua dimensão.")
    add_bullet(doc, "Consolidator — junta tudo, resolve conflitos e produz o OCG final.")
    add_p(
        doc,
        "Cada agente declara confiança e gaps encontrados. Se um agente diz 'não tenho "
        "informação suficiente sobre X', o GCA registra a lacuna e cobra ingestão "
        "adicional. Provedores de IA suportados: Anthropic, OpenAI, Google Gemini, "
        "DeepSeek, Grok e Qwen — configuráveis em Admin → Provedores de IA."
    )

    add_h2(doc, "2.4 RBAC — 7 papéis distintos")
    add_p(doc, "Cada usuário tem um ou mais papéis em projetos. Admin é exceção — atua só na camada do sistema.")
    add_table_two_cols(
        doc,
        ("Papel", "O que faz / não faz"),
        [
            ("Admin", "Configura sistema, aprova projetos, gerencia usuários e provedores de IA. NÃO atua dentro de projetos."),
            ("GP (Gerente de Projeto)", "Conduz o projeto, configura repos, convida time, aprova entregas. NÃO escreve código."),
            ("Tech Lead", "Define arquitetura, revisa CodeGen, valida decisões técnicas."),
            ("Dev (Sênior / Pleno)", "Edita código, abre PRs, executa testes. NÃO aprova suas próprias entregas."),
            ("QA", "Revisa planos de teste e execuções, aprova qualidade. NUNCA edita código."),
            ("Compliance", "Audita pilares P2 (Compliance) e P7 (Segurança), aprova governança."),
            ("Stakeholder / Viewer", "Visualiza progresso, comenta. Sem permissão de edição."),
        ],
    )

    add_h2(doc, "2.5 Pipeline do projeto (10 fases)")
    add_p(doc, "Cada projeto passa por uma sequência clara, visível na barra lateral:")
    fases = [
        "1. Questionário inicial — captura escopo, restrições, KPIs.",
        "2. Repositório do Projeto (obrigatório) — todo projeto precisa ter um repo Git próprio.",
        "3. Repos Externos — código legado/referência para análise.",
        "4. Ingestão — documentos (PDF, DOCX, MD) viram contexto para o OCG.",
        "5. Gatekeeper — avalia os 7 pilares e libera ou trava o pipeline.",
        "6. OCG — gerado e atualizado reativamente conforme novos subsídios.",
        "7. Arguidor — ajusta arquitetura, stack, padrões com base no OCG.",
        "8. Backlog → Roadmap → CodeGen — geração de código sob aprovação.",
        "9. Testes (QA Readiness + Tester Review) — planos e execuções.",
        "10. Documentação Viva — atualizada a cada commit.",
    ]
    for f in fases:
        add_bullet(doc, f)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 3. Tutorial passo a passo
    # ════════════════════════════════════════════════════════════════
    add_h1(doc, "3. Tutorial passo a passo")
    add_p(
        doc,
        "Esta seção cobre as jornadas mais comuns. Cada uma está desenhada como "
        "uma sequência de passos que você executa na interface."
    )

    add_h2(doc, "3.1 Solicitar um novo projeto (sem ter conta no GCA)")
    add_numbered(doc, "Acesse a página de login do GCA.")
    add_numbered(doc, "Clique em 'Solicitar novo projeto' (link abaixo do botão Entrar).")
    add_numbered(doc, "Preencha o passo 1 do wizard: seu nome, email, nome do projeto, tipo de entregável e descrição (mínimo 30 caracteres).")
    add_numbered(doc, "Se nenhum tipo padrão couber, escolha 'Outro' e descreva (ex: 'Browser Extension', 'CLI Tool').")
    add_numbered(doc, "No passo 2, responda as perguntas obrigatórias específicas do tipo escolhido. Essas respostas alimentam o seed inicial do OCG.")
    add_numbered(doc, "Envie. O administrador receberá uma notificação in-app e por email.")
    add_numbered(doc, "Você receberá um email de aprovação ou recusa (sempre no email que cadastrou).")
    add_callout(
        doc, "Atenção",
        "Quem solicita um projeto se torna automaticamente o GP (Gerente de Projeto) "
        "quando o admin aprovar. Use seu email principal — ele será o canal oficial de "
        "comunicação."
    )

    add_h2(doc, "3.2 Fazer login (Admin OU membro de projeto)")
    add_p(
        doc,
        "A página de login é única. O comportamento muda conforme o que você seleciona:"
    )
    add_bullet(doc, "Sem projeto selecionado: só admins conseguem entrar (vão para a área administrativa).")
    add_bullet(doc, "Com projeto no combo: o sistema valida que você é membro daquele projeto e aplica suas permissões de papel.")
    add_p(doc, "Passos:")
    add_numbered(doc, "Selecione um projeto no combo (o último acessado fica no topo, lembrado pelo navegador).")
    add_numbered(doc, "Digite email e senha.")
    add_numbered(doc, "Clique em Entrar.")
    add_callout(
        doc, "Mensagens de erro",
        "'Email ou senha inválidos' = credenciais erradas. 'Você não é membro deste "
        "projeto' = o GP precisa adicionar você ao time. 'Selecione seu projeto no "
        "combo acima' = você não é admin e tentou entrar sem projeto."
    )

    add_h2(doc, "3.3 Aprovar uma solicitação de projeto (Admin)")
    add_numbered(doc, "Em Admin → Gestão de Projetos, localize a linha com status 'Pendente'.")
    add_numbered(doc, "Clique no ícone de documento (FileText) para abrir os detalhes da solicitação: descrição completa, tipo, perguntas e respostas do wizard. Avalie a aderência.")
    add_numbered(doc, "Para aprovar: clique no ícone verde de check. O tenant é provisionado, o GP vira membro do projeto e dois emails são enviados (aprovação + convite).")
    add_numbered(doc, "Para rejeitar: clique no ícone âmbar de aviso. Escreva o motivo (mínimo 10 caracteres) — ele será enviado por email ao solicitante.")
    add_numbered(doc, "Se quiser apenas excluir sem notificar (ex: spam): use o ícone vermelho de lixeira.")
    add_callout(
        doc, "Boas práticas",
        "Antes de aprovar, leia as respostas do wizard. Se faltam informações críticas, "
        "use o botão 'Mensagem' (lápis) para pedir esclarecimentos sem rejeitar — o "
        "solicitante recebe um email para complementar."
    )

    add_h2(doc, "3.4 Configurar o repositório do projeto (GP)")
    add_p(
        doc,
        "Sem repositório, o pipeline não avança. O GCA suporta GitHub, GitLab e Bitbucket."
    )
    add_numbered(doc, "Entre no projeto e vá em 'Repositório'.")
    add_numbered(doc, "Escolha o provider e cole a URL do repo.")
    add_numbered(doc, "Cole um Personal Access Token (PAT) com permissão de leitura e escrita.")
    add_numbered(doc, "Clique em 'Verificar conexão'. Se ok, o repo é vinculado e a branch padrão é detectada.")

    add_h2(doc, "3.5 Ingerir documentos para o OCG")
    add_numbered(doc, "Vá em Ingestão.")
    add_numbered(doc, "Faça upload de PDFs, DOCX ou Markdown com requisitos, atas, laudos, regulamentações.")
    add_numbered(doc, "O sistema indexa, calcula SHA-256 (deduplicação) e dispara análise automática.")
    add_numbered(doc, "Acompanhe o status: pending → analyzing → processed.")
    add_numbered(doc, "Documentos processados aparecem na lista e alimentam o OCG automaticamente.")
    add_callout(
        doc, "Repos externos",
        "Em 'Repos Externos' você pode apontar para repositórios já existentes "
        "(legado, similares). O Arguidor analisa o código e gera um documento sintético "
        "que entra na ingestão como contexto."
    )

    add_h2(doc, "3.6 Acompanhar o Gatekeeper")
    add_numbered(doc, "Vá em Gatekeeper.")
    add_numbered(doc, "Veja o score de cada um dos 7 pilares.")
    add_numbered(doc, "Itens com tag BLOCKER ou CRITICAL precisam ser resolvidos antes do pipeline avançar.")
    add_numbered(doc, "Para cada item, é possível: marcar como resolvido, anexar evidência, ou registrar justificativa.")

    add_h2(doc, "3.7 Gerar e revisar código (CodeGen)")
    add_numbered(doc, "Após o OCG estar verde no Gatekeeper, vá em Backlog → Roadmap.")
    add_numbered(doc, "Selecione um item do roadmap e dispare CodeGen.")
    add_numbered(doc, "Aguarde a geração (uma barra mostra o progresso).")
    add_numbered(doc, "Revise os arquivos no editor Monaco — diagnostics destacam erros do validador.")
    add_numbered(doc, "Para regerar um arquivo específico, use 'Regerar arquivo' (não regera o projeto inteiro).")
    add_numbered(doc, "Quando aprovar, o GCA commita os arquivos diretamente no repositório do projeto.")
    add_callout(
        doc, "Docstrings obrigatórias",
        "Política do GCA: todo módulo, classe e função em código gerado precisa ter "
        "docstring/comentário documentado (PEP 257 para Python, JSDoc para JS/TS, "
        "Javadoc para Java, godoc para Go). Arquivos sem docstring são rebaixados a "
        "status 'todo' automaticamente."
    )

    add_h2(doc, "3.8 Planejar e executar testes (QA Readiness + Tester Review)")
    add_numbered(doc, "Em QA Readiness, o GCA gera planos de teste a partir do OCG.")
    add_numbered(doc, "O QA revisa, aprova ou edita o plano.")
    add_numbered(doc, "O Tester executa os testes. O GCA registra evidências (logs, artefatos).")
    add_numbered(doc, "QA aprova ou rejeita a execução. Aprovações destravam o item no roadmap.")

    add_h2(doc, "3.9 Verificar Documentação Viva")
    add_numbered(doc, "Vá em Documentação Viva.")
    add_numbered(doc, "A árvore mostra todos os arquivos .md em /docs do repositório.")
    add_numbered(doc, "Clique para visualizar (read-only — para editar use seu IDE e commit).")
    add_numbered(doc, "A cada commit, a documentação é regenerada automaticamente.")

    add_h2(doc, "3.10 Empacotar a entrega (Release Bundle)")
    add_numbered(doc, "Em Readiness, quando todos os deliverables estiverem 'verified', clique 'Gerar Release Bundle'.")
    add_numbered(doc, "O GCA compacta tudo: código, OCG, docs, ADRs, SBOM, plano de testes, métricas.")
    add_numbered(doc, "Você recebe um .zip versionado com MANIFEST.json + RELEASE_NOTES.md + checksum SHA-256.")
    add_numbered(doc, "Esse bundle é o entregável oficial para o cliente / time de produção.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 4. O que o GCA FAZ
    # ════════════════════════════════════════════════════════════════
    add_h1(doc, "4. O que o GCA FAZ")
    add_p(doc, "Lista clara e objetiva do que está disponível hoje na plataforma.")

    add_h2(doc, "Governança de projeto")
    add_bullet(doc, "Recebe e aprova solicitações de novos projetos com wizard tipado e perguntas obrigatórias.")
    add_bullet(doc, "Provisiona um schema PostgreSQL isolado por projeto (multi-tenancy real).")
    add_bullet(doc, "Mantém RBAC granular com 7 papéis e auditoria completa de toda ação.")
    add_bullet(doc, "Permite substituição de GP, transferência de papéis e bloqueio de exclusão de usuários ativos.")

    add_h2(doc, "Análise e contexto (OCG)")
    add_bullet(doc, "Constrói e atualiza o OCG reativamente conforme você adiciona subsídios.")
    add_bullet(doc, "Aplica deltas (mudanças incrementais) ao invés de reescrever — rápido e auditável.")
    add_bullet(doc, "Mantém histórico completo de versões com rollback.")
    add_bullet(doc, "Avalia 7 pilares e bloqueia o pipeline quando Segurança/Compliance < 70%.")

    add_h2(doc, "Geração assistida")
    add_bullet(doc, "Gera código (scaffold completo ou arquivo único) com IA, validado e com docstrings obrigatórias.")
    add_bullet(doc, "Commita arquivos aprovados direto no repositório Git do projeto.")
    add_bullet(doc, "Gera 9+ artefatos automáticos: ADR, OpenAPI, Dockerfile, CI pipeline, observability dashboard, SBOM, plano de testes, diagrama de arquitetura, compliance doc.")
    add_bullet(doc, "Detecta gaps e cobra novos subsídios quando os agentes têm baixa confiança.")

    add_h2(doc, "Integração e infraestrutura")
    add_bullet(doc, "Integra com GitHub, GitLab e Bitbucket via PAT (com plano de migrar para Fernet/criptografia).")
    add_bullet(doc, "Suporta seis provedores de IA: Anthropic, OpenAI, Gemini, DeepSeek, Grok e Qwen.")
    add_bullet(doc, "Permite parametrizar pesos de pilares, thresholds e modelos de IA por instância.")
    add_bullet(doc, "Notificações in-app + email para todos os eventos relevantes.")

    add_h2(doc, "Qualidade e entrega")
    add_bullet(doc, "Pipeline de QA: planos de teste gerados, revisão humana, execução com evidências.")
    add_bullet(doc, "Validador de código tier-1 (Pyflakes para Python, esprima para JS/TS, stdlib para JSON/YAML/TOML).")
    add_bullet(doc, "Editor Monaco integrado com diagnostics e markers em tempo real.")
    add_bullet(doc, "Release Bundle: empacotamento .zip versionado com manifest e changelog.")
    add_bullet(doc, "Documentação viva regenerada a cada commit.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 5. O que o GCA NÃO FAZ
    # ════════════════════════════════════════════════════════════════
    add_h1(doc, "5. O que o GCA NÃO FAZ")
    add_p(
        doc,
        "Tão importante quanto saber o que está incluso é saber os limites. "
        "Estas restrições são intencionais (princípio de design) ou conhecidas como gaps "
        "do roadmap atual."
    )

    add_h2(doc, "Limites por design")
    add_bullet(doc, "Não substitui a decisão humana. A IA propõe, o humano aprova. Toda mudança crítica passa por papel responsável.")
    add_bullet(doc, "Não escreve código sem repositório Git configurado. É bloqueio explícito — sem repo, nada avança.")
    add_bullet(doc, "Não compartilha dados entre projetos: cada projeto tem schema PostgreSQL isolado, com RLS.")
    add_bullet(doc, "Não permite que Admin atue dentro de projetos. Admin é camada de sistema; quem opera é GP/Tech Lead/Dev.")
    add_bullet(doc, "Não permite que QA edite código. Não permite que Dev aprove sua própria entrega.")
    add_bullet(doc, "Não exclui usuários que são GP de projetos ativos — exige transferência prévia.")
    add_bullet(doc, "Não persiste credenciais sensíveis em texto plano nos logs (mas o PAT do Git ainda está em texto plano no DB — gap a resolver).")
    add_bullet(doc, "Não aceita aprovação de projetos com score Segurança/Compliance < 70% sem justificativa expressa.")

    add_h2(doc, "Não é um IDE")
    add_bullet(doc, "Edição de código no GCA é leve (Monaco read-mostly). Para refatorar grandes blocos, use seu IDE preferido e faça commit.")
    add_bullet(doc, "O GCA não roda debugger nem testes automatizados na infra dele — apenas registra que foram executados.")
    add_bullet(doc, "Não substitui pipelines de CI/CD existentes — gera o YAML mas a execução é externa.")

    add_h2(doc, "Não é um SaaS multi-cliente")
    add_bullet(doc, "GCA é produto instalável (Docker Compose hoje). Cada cliente tem sua própria instância.")
    add_bullet(doc, "Não há tenant compartilhado. Não há billing centralizado. Não há marketplace de templates ainda.")

    add_h2(doc, "Gaps conhecidos (no roadmap)")
    add_bullet(doc, "Instalador único (binário ou installer GUI) — hoje só Docker Compose.")
    add_bullet(doc, "Auto-upgrade com preview de changes — hoje upgrade manual.")
    add_bullet(doc, "Backup e restore com export/import de configuração — em planejamento.")
    add_bullet(doc, "Hardening de produção (rate limit, WAF, certificados gerenciados) — em planejamento.")
    add_bullet(doc, "Versionamento de schema com migrations idempotentes — parcial, será completo na v1.0.")
    add_bullet(doc, "PAT do Git criptografado com Fernet — pendente.")
    add_bullet(doc, "Marketplace de templates de projeto (ex: 'Sistema de e-commerce', 'API CRUD') — futuro.")

    add_h2(doc, "Não é mágico — quando dará trabalho")
    add_bullet(doc, "Se a descrição do projeto for vaga, o OCG vai ser raso e o Gatekeeper vai cobrar mais ingestão.")
    add_bullet(doc, "Se você não fornecer documentos, regulamentações, contratos — os pilares de Negócio e Compliance ficarão amarelos.")
    add_bullet(doc, "Se aprovar geração de código sem revisar, vai pagar depois com retrabalho.")
    add_bullet(doc, "Se ignorar bloqueios do Gatekeeper, o pipeline vai quebrar mais adiante (em testes ou Release).")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 6. Perguntas e Respostas (FAQ)
    # ════════════════════════════════════════════════════════════════
    add_h1(doc, "6. Perguntas e Respostas (FAQ)")

    add_h2(doc, "Sobre o produto")
    add_qa(
        doc,
        "Por que GCA e não outro framework de IA?",
        "O GCA não é só um wrapper de LLM. Ele orquestra papéis humanos, governa "
        "decisões com 7 pilares e mantém auditoria completa. A IA é um meio; o fim "
        "é entrega rastreável."
    )
    add_qa(
        doc,
        "O GCA precisa de internet?",
        "Sim, para chamar os provedores de IA (Anthropic, OpenAI, etc) e para "
        "comunicar com seu repositório Git remoto. Nada do código do seu projeto "
        "sai da sua instância — o que sai são prompts construídos com o contexto "
        "necessário (você controla o provedor)."
    )
    add_qa(
        doc,
        "Posso usar o GCA com modelos locais (Ollama, LM Studio)?",
        "Hoje suportamos os provedores listados (Anthropic, OpenAI, Gemini, DeepSeek, "
        "Grok, Qwen). Modelos locais via API compatível com OpenAI funcionam — basta "
        "configurar o endpoint customizado em Admin → Provedores. Não há suporte "
        "oficial ainda."
    )
    add_qa(
        doc,
        "É open source?",
        "A instância atual é dogfood do criador. A política de licenciamento da v1.0 "
        "será definida antes do release público — consulte o canal oficial."
    )

    add_h2(doc, "Sobre conta e acesso")
    add_qa(
        doc,
        "Esqueci minha senha. O que fazer?",
        "Na tela de login, clique 'Esqueci minha senha'. Você receberá um email com "
        "link de redefinição válido por 24 horas."
    )
    add_qa(
        doc,
        "Posso ter múltiplos papéis em um mesmo projeto?",
        "Sim. Por exemplo, alguém pode ser Tech Lead e Dev Sênior simultaneamente. "
        "As permissões somam, mas as restrições mais fortes prevalecem (ex: ainda assim "
        "não aprova suas próprias entregas)."
    )
    add_qa(
        doc,
        "Como adiciono pessoas ao meu projeto?",
        "Como GP, vá em Equipe (dentro do projeto). Convide por email com o papel "
        "desejado. A pessoa recebe um link de aceite válido por 5 dias úteis."
    )

    add_h2(doc, "Sobre o pipeline")
    add_qa(
        doc,
        "Quanto tempo leva para o OCG ficar pronto?",
        "Depende do volume de ingestão e do provedor de IA escolhido. Com Anthropic "
        "Claude Sonnet 4.6 e ingestão moderada (5–10 documentos), tipicamente 30 "
        "segundos a 3 minutos. Atualizações reativas (delta) levam 5 a 15 segundos."
    )
    add_qa(
        doc,
        "O GCA pode regenerar o OCG do zero?",
        "Sim, em casos excepcionais. Em Admin → Configurações você pode forçar "
        "regeneração completa. Cuidado: descarta deltas manuais. Use só após mudanças "
        "estruturais (novo escopo, nova regulamentação)."
    )
    add_qa(
        doc,
        "O que acontece se eu rejeitar uma sugestão da IA?",
        "Sua decisão fica registrada com motivo. O GCA aprende do contexto do projeto "
        "— sugestões futuras consideram seus rejeições passados (via OCG)."
    )
    add_qa(
        doc,
        "O CodeGen sobrescreve meus arquivos editados manualmente?",
        "Depende do fluxo: 'Regerar projeto' (scaffold) sobrescreve tudo — use só na "
        "primeira geração ou reset. 'Regerar arquivo' regera um arquivo único — "
        "preserva o resto. Sempre revise o diff no editor Monaco antes de commitar."
    )

    add_h2(doc, "Sobre dados e segurança")
    add_qa(
        doc,
        "Onde ficam meus documentos ingeridos?",
        "No volume de storage da sua instância (pasta backend/storage/ingested/<project_id>/). "
        "Não saem do seu servidor. Apenas o conteúdo necessário para o prompt da IA é "
        "enviado ao provedor escolhido (você controla qual)."
    )
    add_qa(
        doc,
        "O GCA armazena minha senha em texto plano?",
        "Não. Senhas são hasheadas com bcrypt antes de gravar no banco. Tokens de "
        "redefinição expiram em 24h. JWT são assinados com RS256."
    )
    add_qa(
        doc,
        "E os Personal Access Tokens (PAT) do Git?",
        "Hoje, o PAT é armazenado em texto plano no campo 'pat_encrypted' (sim, o "
        "nome é mentiroso — herança histórica). Está no roadmap migrar para Fernet "
        "(criptografia simétrica). Use PATs com escopo mínimo até lá."
    )
    add_qa(
        doc,
        "O GCA tem auditoria?",
        "Sim. Toda ação relevante (login, aprovação, geração de código, edição de OCG, "
        "exclusão) gera entrada em audit_log_global ou no audit do tenant. Visível em "
        "Admin → Auditoria Global e em cada projeto na aba Auditoria."
    )

    add_h2(doc, "Sobre desenvolvimento e código gerado")
    add_qa(
        doc,
        "Em quais linguagens o CodeGen funciona?",
        "Suportadas: Python, JavaScript/TypeScript, Java, Go. Validador tier-1 cobre "
        "Python (Pyflakes), JS/TS (esprima), JSON/YAML/TOML (stdlib). Outras linguagens "
        "são geradas mas sem validação automática."
    )
    add_qa(
        doc,
        "Posso usar meu próprio template/skeleton?",
        "Hoje o scaffold é determinado pelo OCG e pelo Arguidor. Templates customizados "
        "estão no roadmap (parte do marketplace futuro). Como workaround: faça o scaffold "
        "primeiro e depois adapte; o GCA respeita edições nos arquivos individuais."
    )
    add_qa(
        doc,
        "Por que docstrings são obrigatórias?",
        "Por política de governança do GCA: código sem documentação inline é dívida. "
        "O validador rebaixa arquivos sem docstring para status 'todo', o que impede "
        "marcação como 'verified' no Readiness. Isenções: arquivos de config, "
        "__init__.py triviais, e markdown."
    )

    add_h2(doc, "Sobre a entrega final")
    add_qa(
        doc,
        "O que está dentro de um Release Bundle?",
        "Um arquivo .zip versionado contendo: código-fonte do projeto (snapshot do "
        "repo), OCG completo, documentação viva (markdown), ADRs, SBOM CycloneDX 1.5, "
        "plano de testes, diagrama de arquitetura, MANIFEST.json com lista de "
        "deliverables verificados, RELEASE_NOTES.md e checksum SHA-256 para integridade."
    )
    add_qa(
        doc,
        "Posso entregar o Release Bundle direto pro cliente?",
        "Sim — esse é o propósito. O bundle é auto-contido e auditável. O cliente "
        "recebe tudo que precisa para entender, executar, manter e auditar a entrega."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 7. Glossário
    # ════════════════════════════════════════════════════════════════
    add_h1(doc, "7. Glossário")
    add_table_two_cols(
        doc,
        ("Termo", "Definição"),
        [
            ("OCG", "Objeto de Contexto Global. Documento vivo único por projeto, alimentado pela ingestão e usado por todos os módulos."),
            ("Pilares", "As 7 dimensões avaliadas pelo Gatekeeper: Negócio, Compliance, Escopo, NFR, Arquitetura, Dados, Segurança."),
            ("Gatekeeper", "Módulo que avalia e bloqueia ou libera o pipeline com base no score dos pilares."),
            ("Arguidor", "Módulo que ajusta arquitetura, stack e padrões com base no OCG e em repos externos analisados."),
            ("CodeGen", "Geração de código assistida por IA, com validação automática e commit no Git."),
            ("RBAC", "Role-Based Access Control. Sistema de papéis e permissões."),
            ("Tenant", "Schema PostgreSQL isolado por projeto. Garante separação de dados."),
            ("PAT", "Personal Access Token. Token usado para autenticar com GitHub/GitLab/Bitbucket."),
            ("SBOM", "Software Bill of Materials. Lista de todas as dependências e versões usadas pelo projeto. Padrão CycloneDX 1.5."),
            ("ADR", "Architecture Decision Record. Documento curto registrando uma decisão arquitetural e seu motivo. Padrão MADR."),
            ("Definition of Done", "Conjunto de deliverables que precisam estar 'verified' para o projeto ser considerado pronto para entrega."),
            ("Release Bundle", "Empacotamento final em .zip com código + OCG + docs + ADRs + SBOM + manifest + changelog."),
            ("Delta OCG", "Mudança incremental aplicada ao OCG (ao invés de reescrever todo o documento). Mais rápido e auditável."),
            ("Dogfood", "Uso do próprio produto pela equipe que o desenvolve para validá-lo na prática."),
        ],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 8. Suporte e versão
    # ════════════════════════════════════════════════════════════════
    add_h1(doc, "8. Suporte e versão")

    add_h2(doc, "Onde pedir ajuda")
    add_bullet(doc, "Dentro do GCA: cada tela tem tooltips de ajuda (passe o mouse sobre os ícones de '?').")
    add_bullet(doc, "Documentação viva: cada projeto mantém em /docs material atualizado.")
    add_bullet(doc, "Suporte ao admin da sua instância: ele pode acessar logs, auditoria e parametrizações.")

    add_h2(doc, "Reportar problema")
    add_p(
        doc,
        "Ao reportar um erro, sempre inclua: (1) tela onde ocorreu, (2) ação executada, "
        "(3) mensagem de erro completa, (4) horário aproximado. Isso permite ao admin "
        "localizar o evento na auditoria rapidamente."
    )

    add_h2(doc, "Versão do GCA")
    add_p(
        doc,
        "A versão da sua instância aparece no rodapé da tela de login (ex: 'v0.8.0'). "
        "Use-a sempre que reportar bugs ou comparar comportamentos. Em release futuro, "
        "haverá tela 'Sobre' com detecção de update disponível."
    )

    add_p(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = p.add_run("— Fim do tutorial —")
    final_run.italic = True
    final_run.font.color.rgb = SLATE_MID

    return doc


def main() -> None:
    doc = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"OK: {OUTPUT}")


if __name__ == "__main__":
    main()
