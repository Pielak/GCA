#!/usr/bin/env python3
"""
Gera dois .docx para o pipeline de ingestão do GCA:

  1. docs/GCA_TEMPLATE_INGESTAO.docx
     Template universal preenchível para qualquer projeto. Mostra
     exatamente o formato que o Arguidor lê (parágrafos, sem tabelas),
     com seções que disparam criação de roadmap.

  2. docs/Automacao_Juridica_Assistida-v2.0.docx
     Reescrita do documento original do projeto (v1.0) no formato do
     template, consolidando:
     - Documento original de requisitos (RF-01..RF-60)
     - Questionário preenchido (respostas)
     - Doc técnico DataJud (integração CNJ)
     - Endpoints de tribunais (aliases)

Problema que resolve: o Arguidor do GCA parseia .docx lendo APENAS
`Document(path).paragraphs[].text`. Tabelas não são lidas. O documento
original v1.0 tinha RF-01..RF-60 dentro de tabelas → Arguidor
literalmente não os viu → backlog ficou vazio → roadmap não foi gerado.

Solução: todo requisito, entrega, módulo e fase em parágrafo de texto
numerado com prefixos padronizados que o prompt do Arguidor reconhece.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("/home/luiz/GCA/docs")
TEMPLATE_PATH = OUT_DIR / "GCA_TEMPLATE_INGESTAO.docx"
V2_PATH = OUT_DIR / "Automacao_Juridica_Assistida-v2.0.docx"

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
SLATE_DARK = RGBColor(0x1E, 0x29, 0x3B)
SLATE_MEDIUM = RGBColor(0x64, 0x74, 0x8B)


def h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = VIOLET
        r.font.size = Pt(18)


def h2(doc, text: str):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = SLATE_DARK
        r.font.size = Pt(14)


def h3(doc, text: str):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = SLATE_MEDIUM
        r.font.size = Pt(12)


def para(doc, text: str, *, size: int = 11, bold: bool = False, italic: bool = False,
         color: RGBColor = None, justify: bool = True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


def rf_para(doc, rf_id: str, priority: str, pillar: str, text: str):
    """Requisito funcional em parágrafo (não-tabela) legível pelo
    Arguidor. Formato: `RF-XX [PRIORIDADE] [PILAR]: texto.`"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)

    r1 = p.add_run(f"{rf_id} ")
    r1.bold = True
    r1.font.color.rgb = VIOLET
    r1.font.size = Pt(11)

    color = RGBColor(0xC2, 0x41, 0x0C) if priority == "MANDATÓRIO" else RGBColor(0x05, 0x96, 0x69)
    r2 = p.add_run(f"[{priority}] ")
    r2.bold = True
    r2.font.color.rgb = color
    r2.font.size = Pt(10)

    r3 = p.add_run(f"[{pillar}] ")
    r3.font.color.rgb = SLATE_MEDIUM
    r3.font.size = Pt(10)
    r3.italic = True

    r4 = p.add_run(text)
    r4.font.size = Pt(11)


def bullet(doc, text: str, *, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


def numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


def delivery(doc, nome: str, desc: str, fase: str, prioridade: str):
    """Entregável em parágrafo estruturado com prefixos que o Arguidor lê."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(4)

    r1 = p.add_run(f"ENTREGÁVEL: {nome}. ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = VIOLET

    r2 = p.add_run(f"Descrição: {desc} ")
    r2.font.size = Pt(11)

    r3 = p.add_run(f"Fase: {fase}. ")
    r3.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = SLATE_MEDIUM

    r4 = p.add_run(f"Prioridade: {prioridade}.")
    r4.italic = True
    r4.font.size = Pt(10)
    r4.font.color.rgb = SLATE_MEDIUM


def module_para(doc, numero: str, nome: str, descricao: str, pilares: str,
                prioridade: str, dependencias: str, ready: bool):
    """Módulo identificado pelo Arguidor como module_candidate.
    Formato que permite extrair: name, description, priority, pillar_impact,
    dependencies, ready_for_codegen."""
    h3(doc, f"{numero} Módulo: {nome}")

    para(doc,
         f"MÓDULO: {nome}. Prioridade: {prioridade}. Pilares: {pilares}. "
         f"Dependências: {dependencias}. Pronto para CodeGen: {'SIM' if ready else 'NÃO'}.",
         bold=True, color=VIOLET, size=10)
    para(doc, f"Descrição completa: {descricao}")


def page_break(doc):
    doc.add_page_break()


def titlepage(doc, title: str, subtitle: str, version: str, author: str, date: str):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        doc.add_paragraph()
    tr = t.add_run(title)
    tr.font.size = Pt(40)
    tr.bold = True
    tr.font.color.rgb = VIOLET

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.font.size = Pt(18)
    sr.font.color.rgb = SLATE_DARK

    for _ in range(4):
        doc.add_paragraph()

    v = doc.add_paragraph()
    v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = v.add_run(version)
    vr.font.size = Pt(14)
    vr.italic = True

    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = a.add_run(author)
    ar.font.size = Pt(12)
    ar.bold = True

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = d.add_run(date)
    dr.font.size = Pt(11)
    dr.font.color.rgb = SLATE_MEDIUM

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# TEMPLATE UNIVERSAL
# ═══════════════════════════════════════════════════════════════════

def build_template(doc: Document):
    doc.core_properties.title = "GCA — Template de Documento de Ingestão"
    doc.core_properties.author = "Luiz Carlos Pielak"
    doc.core_properties.subject = "Template padrão para docs ingeridos no pipeline do GCA"

    titlepage(doc,
              "GCA", "Template de Documento de Ingestão",
              "Versão 1.0 — formato ideal para o pipeline do Arguidor",
              "Autor: Luiz Carlos Pielak",
              datetime.now().strftime("%d de %B de %Y"))

    # ─── Introdução metodológica ──────────────────────────────────
    h1(doc, "Como usar este template")
    para(doc,
         "Este documento é o formato ideal para qualquer material que o cliente "
         "queira ingerir no GCA (menu Ingestão do projeto) para alimentar o OCG, o "
         "Gatekeeper e o Roadmap. O pipeline do GCA parseia arquivos .docx lendo "
         "apenas parágrafos de texto — tabelas, imagens, SmartArt e elementos "
         "visuais não são interpretados.")
    para(doc,
         "Por isso, todo requisito funcional (RF), entregável, módulo, fase e "
         "prioridade deve estar em parágrafo numerado com prefixos padronizados. "
         "As seções deste template são o mínimo que o Arguidor precisa para gerar "
         "module_candidates, DELIVERABLES e fases de roadmap automaticamente.")

    h2(doc, "Regras duras de formatação")
    bullet(doc, "Use apenas parágrafos de texto. Nunca coloque requisitos em tabelas.")
    bullet(doc, "Cada RF começa com identificador único (RF-01, RF-02, ...) em parágrafo próprio.")
    bullet(doc, "Prefixos obrigatórios por RF: [MANDATÓRIO] ou [RECOMENDADO] + [P1-Business..P7-Deliverables].")
    bullet(doc, "Cada entregável em parágrafo próprio começando com ENTREGÁVEL:.")
    bullet(doc, "Cada módulo em seção própria começando com MÓDULO:.")
    bullet(doc, "Fases de roadmap numeradas (Fase 1, Fase 2...) com lista de entregas.")
    bullet(doc, "Glossário no final. Termos em ordem alfabética.")

    h2(doc, "Seções esperadas pelo pipeline")
    numbered(doc, "Identificação do projeto (nome, versão, autor, data).")
    numbered(doc, "Objetivo e escopo.")
    numbered(doc, "Glossário de termos específicos do domínio.")
    numbered(doc, "Requisitos funcionais (RF-XX) em parágrafos.")
    numbered(doc, "Requisitos não-funcionais (RNF-XX) em parágrafos.")
    numbered(doc, "Módulos ou componentes (cada um com MÓDULO: nome + descrição + pilares + prioridade).")
    numbered(doc, "Entregáveis (cada um com ENTREGÁVEL: nome + fase + prioridade).")
    numbered(doc, "Roadmap por fases (Fase 1, 2, 3, ...) listando entregas.")
    numbered(doc, "Riscos conhecidos.")
    numbered(doc, "Critérios de aceite (checklist).")
    numbered(doc, "Integrações externas (APIs, bases, fontes de dados).")

    page_break(doc)

    # ─── Template propriamente dito ────────────────────────────────
    h1(doc, "1. Identificação do projeto")
    para(doc, "Nome: [preencher]", bold=True)
    para(doc, "Slug: [preencher, ex: automacao-juridica]", bold=True)
    para(doc, "Versão: 1.0")
    para(doc, "Data: [preencher]")
    para(doc, "Autor/Elaboração: [preencher]")
    para(doc, "Classificação: [pública / interna / confidencial / restrita]")

    h1(doc, "2. Objetivo")
    para(doc,
         "Descreva em 3-5 parágrafos o problema que o projeto resolve, o ganho "
         "esperado e o público-alvo. Seja específico — o Arguidor usa este texto "
         "para avaliar o pilar P1 (Business).")

    h2(doc, "2.1 Objetivos específicos")
    bullet(doc, "[Objetivo mensurável 1]")
    bullet(doc, "[Objetivo mensurável 2]")
    bullet(doc, "[Objetivo mensurável 3]")

    h1(doc, "3. Escopo")
    h2(doc, "3.1 Em escopo")
    bullet(doc, "[Capacidade central 1 — descrição completa]")
    bullet(doc, "[Capacidade central 2]")

    h2(doc, "3.2 Fora de escopo")
    bullet(doc, "[O que NÃO será entregue nesta versão — evite ambiguidade]")

    h1(doc, "4. Glossário do projeto")
    para(doc,
         "Liste termos específicos do domínio em ordem alfabética. Cada termo "
         "em um parágrafo começando com o termo em negrito.", italic=True, size=10)
    para(doc, "Termo. Definição clara e autocontida.")

    h1(doc, "5. Requisitos funcionais")
    para(doc,
         "Cada requisito funcional (RF) em parágrafo próprio, numerado. Prefixos: "
         "[MANDATÓRIO] ou [RECOMENDADO] + [P1-Business / P2-Architecture / "
         "P3-Stack / P4-Testing / P5-Compliance / P6-Risk / P7-Deliverables]. "
         "Exemplos abaixo — substitua pelos RFs reais do projeto.", italic=True, size=10)

    h2(doc, "5.1 Módulo exemplo")
    rf_para(doc, "RF-01", "MANDATÓRIO", "P1-Business",
            "Permitir cadastro manual de [entidade] com campos obrigatórios A, B, C e campos opcionais X, Y.")
    rf_para(doc, "RF-02", "MANDATÓRIO", "P2-Architecture",
            "Manter histórico completo de alterações em [entidade] com timestamp, usuário e diff.")
    rf_para(doc, "RF-03", "RECOMENDADO", "P4-Testing",
            "Permitir exportação em formato [CSV/PDF/Excel] respeitando filtros ativos.")

    h1(doc, "6. Requisitos não-funcionais")
    rf_para(doc, "RNF-01", "MANDATÓRIO", "P2-Architecture",
            "Tempo de resposta para operações síncronas de listagem ≤ 2 segundos para base com [volume].")
    rf_para(doc, "RNF-02", "MANDATÓRIO", "P5-Compliance",
            "Conformidade com LGPD: [detalhar controles — consentimento, minimização, retenção].")
    rf_para(doc, "RNF-03", "MANDATÓRIO", "P6-Risk",
            "Disponibilidade operacional alvo: 99,5% mensal.")

    h1(doc, "7. Módulos do sistema")
    para(doc,
         "Cada módulo em seção própria. Esta seção alimenta module_candidates do "
         "Arguidor, que por sua vez geram items de roadmap priorizados.",
         italic=True, size=10)

    module_para(doc, "7.1", "Nome do Módulo 1",
                "Descrição completa do que o módulo faz, inputs, outputs, regras de negócio principais, "
                "casos de uso cobertos. Detalhe suficiente para o Arguidor marcar ready_for_codegen=true.",
                "P1-Business, P2-Architecture",
                "MANDATÓRIO",
                "Nenhuma",
                True)
    module_para(doc, "7.2", "Nome do Módulo 2",
                "Descrição completa...",
                "P2-Architecture, P3-Stack",
                "MANDATÓRIO",
                "Depende de Módulo 1",
                True)

    h1(doc, "8. Entregáveis")
    para(doc,
         "Cada entregável em parágrafo próprio. Prefixo ENTREGÁVEL: disparador "
         "do pipeline que popula o campo DELIVERABLES do OCG.",
         italic=True, size=10)

    delivery(doc, "Setup inicial do projeto",
             "Repositório Git inicializado com README, LICENSE, .gitignore, CI e estrutura de pastas padrão.",
             "Fase 1", "MANDATÓRIO")
    delivery(doc, "Schema do banco",
             "Migrations SQL iniciais criando entidades principais com chaves e índices.",
             "Fase 1", "MANDATÓRIO")
    delivery(doc, "API REST inicial",
             "Endpoints de CRUD das entidades principais com autenticação JWT.",
             "Fase 1", "MANDATÓRIO")
    delivery(doc, "Frontend SPA inicial",
             "Telas de login, dashboard, listagens principais com roteamento protegido.",
             "Fase 1", "MANDATÓRIO")
    delivery(doc, "Documentação técnica",
             "README de setup, doc de API gerada (OpenAPI), ADRs das decisões arquiteturais.",
             "Fase 1", "RECOMENDADO")

    h1(doc, "9. Roadmap")
    para(doc,
         "Fases numeradas (Fase 1, Fase 2, ...) com entregas nomeadas por fase. "
         "Esta seção orienta o roadmap_service a agrupar module_candidates em "
         "fases de execução.", italic=True, size=10)

    h2(doc, "Fase 1 — Fundação")
    bullet(doc, "Setup inicial do projeto (repo, CI, .gitignore)")
    bullet(doc, "Schema inicial do banco")
    bullet(doc, "Autenticação + autorização")
    bullet(doc, "CRUD principal das entidades centrais")
    bullet(doc, "Telas básicas do frontend")

    h2(doc, "Fase 2 — Funcionalidades centrais")
    bullet(doc, "[Módulos do domínio específico]")
    bullet(doc, "[Integrações de primeira linha]")
    bullet(doc, "[Relatórios básicos]")

    h2(doc, "Fase 3 — Maturidade")
    bullet(doc, "[Recursos avançados / analytics]")
    bullet(doc, "[Integrações complexas]")
    bullet(doc, "[Otimização de performance]")

    h2(doc, "Fase 4 — Expansões futuras (opcional)")
    bullet(doc, "[Features nice-to-have]")

    h1(doc, "10. Riscos")
    para(doc, "RISCO-01: [descrição do risco + probabilidade (alta/média/baixa) + impacto + mitigação].")
    para(doc, "RISCO-02: ...")

    h1(doc, "11. Critérios de aceite")
    numbered(doc, "Todos os RF-[MANDATÓRIO] implementados e testados.")
    numbered(doc, "Cobertura de testes unitários ≥ [X]% nos módulos críticos.")
    numbered(doc, "Documentação técnica atualizada.")
    numbered(doc, "Aprovação do Gatekeeper em todos os 7 pilares.")

    h1(doc, "12. Integrações externas")
    para(doc,
         "Cada integração em seção própria com: nome da fonte, natureza (pública/"
         "privada/paga), método de acesso (REST/SOAP/webhook), autenticação, "
         "limites de taxa, riscos conhecidos.", italic=True, size=10)

    h2(doc, "12.1 [Nome da integração]")
    para(doc, "Fonte: [URL ou descrição].")
    para(doc, "Natureza: [pública / privada / paga com contrato X].")
    para(doc, "Método: [HTTP POST com JSON / GraphQL / webhook / ...].")
    para(doc, "Autenticação: [APIKey no header / OAuth2 / JWT / ...].")
    para(doc, "Limites: [requisições por minuto, janela, quotas].")
    para(doc, "Riscos: [disponibilidade, mudança de contrato, latência].")
    para(doc, "Uso no projeto: [o que alimenta, qual módulo consome].")

    h1(doc, "13. Conformidade e compliance")
    bullet(doc, "LGPD: [controles aplicáveis — base legal, minimização, DPO, portabilidade].")
    bullet(doc, "Segurança da informação: [criptografia em trânsito, em repouso, gestão de chaves].")
    bullet(doc, "Auditoria: [trilhas obrigatórias, retenção, integridade verificável].")

    page_break(doc)
    h1(doc, "Anexo A — Checklist de verificação antes da ingestão")
    numbered(doc, "Todos os RFs estão em parágrafos (não em tabelas)?")
    numbered(doc, "Cada RF tem identificador único (RF-01, RF-02, ...)?")
    numbered(doc, "Cada RF tem marcador de prioridade [MANDATÓRIO] ou [RECOMENDADO]?")
    numbered(doc, "Cada RF tem marcador de pilar [P1-Business..P7-Deliverables]?")
    numbered(doc, "Seção 'Entregáveis' preenchida com prefixos ENTREGÁVEL:?")
    numbered(doc, "Seção 'Módulos' preenchida com prefixos MÓDULO:?")
    numbered(doc, "Seção 'Roadmap' com fases numeradas e lista de entregas?")
    numbered(doc, "Glossário em ordem alfabética ao final?")
    numbered(doc, "Documento salvo como .docx (não .pdf nem .odt)?")
    numbered(doc, "Tamanho inferior a 10 MB?")


# ═══════════════════════════════════════════════════════════════════
# AUTOMAÇÃO JURÍDICA ASSISTIDA v2.0
# Consolidação: doc v1.0 + questionário + DataJud + endpoints
# ═══════════════════════════════════════════════════════════════════

def build_auto_juridica_v2(doc: Document):
    doc.core_properties.title = "Automação Jurídica Assistida — v2.0"
    doc.core_properties.author = "Luiz Carlos Pielak"
    doc.core_properties.subject = "Especificação técnica ingestável no GCA"

    titlepage(doc,
              "Automação Jurídica Assistida",
              "Especificação Técnica de Requisitos",
              "Versão 2.0 — formato ideal para ingestão no GCA",
              "Autor: Luiz Carlos Pielak",
              datetime.now().strftime("Consolidado em %d de %B de %Y"))

    # ─── 1. Identificação ─────────────────────────────────────────
    h1(doc, "1. Identificação do projeto")
    para(doc, "Nome: Automação Jurídica Assistida.", bold=True)
    para(doc, "Slug: automacao-juridica-assistida.", bold=True)
    para(doc, "Versão: 2.0 (consolidação do v1.0 + documento técnico DataJud + endpoints de tribunais + questionário preenchido).")
    para(doc, "Data: consolidado em 2026-04-19.")
    para(doc, "Elaboração: Luiz Carlos Pielak — Equipe Técnica de Especificação de Requisitos.")
    para(doc, "Classificação: Documento Confidencial.")
    para(doc, "Tipo (questionário Q4): Novo sistema + Automação interna.")
    para(doc, "Criticidade (questionário Q5): Alta.")
    para(doc, "Classificação da informação (questionário Q6): Restrita.")

    # ─── 2. Objetivo ──────────────────────────────────────────────
    h1(doc, "2. Objetivo")
    para(doc,
         "Entregar aplicativo desktop instalável para automação assistida do ciclo "
         "cível do escritório de advocacia, com arquitetura local-first e aderência "
         "total à LGPD. A Solução automatiza as tarefas repetitivas do contencioso "
         "cível (cadastro e dossiê do caso, ingestão documental, pesquisa "
         "jurisprudencial, geração de peças, motor de cálculos, análise de risco, "
         "auditoria) sem dependência de serviços em nuvem e sem expor dados "
         "sensíveis a terceiros.")

    h2(doc, "2.1 Objetivos específicos")
    bullet(doc, "Reduzir em 60% o tempo médio de elaboração de peças padronizadas (contestação, petição inicial, recurso) via geração assistida por IA local.")
    bullet(doc, "Centralizar em dossiê único todo o histórico documental do caso, com linha do tempo auditável.")
    bullet(doc, "Integrar consulta automática à API Pública do DataJud para enriquecimento de dados processuais.")
    bullet(doc, "Automatizar cálculos cíveis (juros, correção monetária, honorários, custas) com memória de cálculo exportável.")
    bullet(doc, "Prover painel gerencial com visão de portfólio, riscos e prazos críticos.")
    bullet(doc, "Operar 100% offline por padrão, com sincronização opcional em versões posteriores.")

    # ─── 3. Escopo ────────────────────────────────────────────────
    h1(doc, "3. Escopo")

    h2(doc, "3.1 Em escopo")
    bullet(doc, "Aplicativo desktop instalável (Windows 10/11 e Ubuntu 22.04+).")
    bullet(doc, "Armazenamento local em SQLite ou PostgreSQL embarcado, criptografado em repouso.")
    bullet(doc, "Consulta à API Pública do DataJud para enriquecimento processual (leitura).")
    bullet(doc, "OCR embutido para PDFs digitalizados.")
    bullet(doc, "Modelos de peças parametrizáveis por tipo de ação cível.")
    bullet(doc, "Inteligência artificial local para sugestões — modelo baixado e executado na máquina do advogado.")
    bullet(doc, "Anexo regional: Execução Fiscal do Estado do Paraná.")

    h2(doc, "3.2 Fora de escopo")
    bullet(doc, "Peticionamento eletrônico automatizado em tribunais (envio de peças).")
    bullet(doc, "Assinatura digital ICP-Brasil (fica para Fase 3).")
    bullet(doc, "Sincronização multi-dispositivo (fica para Fase 3).")
    bullet(doc, "Integração contábil completa (fica para Fase 3).")
    bullet(doc, "Contencioso criminal, trabalhista e tributário (fora deste MVP; só cível).")
    bullet(doc, "Marketplace de modelos (fica para Fase 4).")

    h2(doc, "3.3 Premissa crítica de parametrização")
    para(doc,
         "A Solução é instalada no computador do advogado com configuração mínima. "
         "Não há servidor central; não há dependência de conexão permanente com a "
         "Internet. Conexão só é usada quando o usuário solicita consulta ao "
         "DataJud, ou quando baixa atualização de modelo de IA. Todas as peças "
         "geradas e dossiês ficam no disco local até decisão explícita do usuário "
         "de exportar ou compartilhar.")

    h2(doc, "3.4 Princípios arquiteturais norteadores")
    bullet(doc, "Local-first: dados sensíveis ficam apenas no host do cliente.")
    bullet(doc, "Privacidade por design: minimização, consentimento, retenção configurável.")
    bullet(doc, "Auditabilidade: toda operação sensível gera registro imutável.")
    bullet(doc, "Modularidade: cada capacidade encapsulada em módulo independente.")
    bullet(doc, "Extensibilidade via conectores: fontes externas plugáveis via configuração.")

    # ─── 4. Glossário ────────────────────────────────────────────
    h1(doc, "4. Glossário do projeto")
    para(doc, "Termos específicos do domínio jurídico e da Solução em ordem alfabética.",
         italic=True, size=10)

    glossary_juridico = [
        ("CNJ", "Conselho Nacional de Justiça. Órgão que mantém a API Pública do DataJud."),
        ("CNPJ", "Cadastro Nacional da Pessoa Jurídica."),
        ("Contencioso cível", "Conjunto de ações judiciais de natureza civil (danos, contratos, família, sucessões, imobiliário)."),
        ("DataJud", "Base Nacional de Dados do Poder Judiciário; API pública operada pelo CNJ."),
        ("Decadência", "Perda do direito material pelo decurso de prazo legal."),
        ("Dossiê", "Pasta digital centralizada com todos os documentos de um caso."),
        ("Execução Fiscal", "Ação para cobrança de dívida ativa da Fazenda Pública, regida pela Lei 6.830/1980 e por legislação estadual complementar."),
        ("ICP-Brasil", "Infraestrutura de Chaves Públicas Brasileira — sistema oficial de certificação digital."),
        ("LGPD", "Lei Geral de Proteção de Dados Pessoais (Lei 13.709/2018)."),
        ("Local-first", "Arquitetura onde dados primários ficam no dispositivo do usuário, sem servidor obrigatório."),
        ("MoSCoW", "Método de priorização: Must / Should / Could / Won't."),
        ("Motor de cálculos", "Módulo que aplica índices oficiais (Selic, IPCA, TR) e regras de juros/correção/honorários."),
        ("Número CNJ", "Identificador padronizado do processo no formato NNNNNNN-DD.AAAA.J.TR.OOOO."),
        ("OCR", "Optical Character Recognition — reconhecimento óptico de caracteres em imagens/PDFs digitalizados."),
        ("Peça", "Documento processual redigido pelo advogado (petição, contestação, recurso, etc.)."),
        ("Prescrição", "Perda da pretensão de exercer um direito pelo decurso do prazo legal."),
        ("Query DSL", "Domain Specific Language do Elasticsearch para consultas estruturadas ao DataJud."),
        ("RF", "Requisito Funcional. Comportamento observável da Solução."),
        ("RNF", "Requisito Não-Funcional. Atributo de qualidade (performance, segurança, disponibilidade)."),
        ("TJPR", "Tribunal de Justiça do Estado do Paraná."),
    ]
    for termo, defin in glossary_juridico:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{termo}. ")
        r1.bold = True
        r1.font.color.rgb = VIOLET
        r1.font.size = Pt(11)
        r2 = p.add_run(defin)
        r2.font.size = Pt(11)

    page_break(doc)

    # ─── 5. Requisitos funcionais ──────────────────────────────
    h1(doc, "5. Requisitos funcionais")
    para(doc,
         "Cada RF em parágrafo próprio com prefixos [MANDATÓRIO]/[RECOMENDADO] + "
         "pilar impactado. Numeração sequencial e contínua através dos módulos.",
         italic=True, size=10)

    h2(doc, "5.1 Módulo de cadastro e dossiê do caso")
    rf_para(doc, "RF-01", "MANDATÓRIO", "P1-Business",
            "Permitir criação manual de caso com número CNJ opcional, cliente, partes, polo (autor/réu), assunto, classe processual, tribunal e observações estratégicas.")
    rf_para(doc, "RF-02", "MANDATÓRIO", "P1-Business",
            "Permitir anexar documentos, mensagens eletrônicas exportadas (EML/MSG), imagens, decisões, despachos, contratos e provas ao dossiê do caso.")
    rf_para(doc, "RF-03", "MANDATÓRIO", "P2-Architecture",
            "Classificar cada item anexado por tipo documental (petição, decisão, despacho, contrato, documento pessoal, prova), origem, data, emissor e relevância.")
    rf_para(doc, "RF-04", "MANDATÓRIO", "P1-Business",
            "Manter linha do tempo do caso com eventos processuais e documentos correlacionados, ordenados cronologicamente com filtros por tipo.")
    rf_para(doc, "RF-05", "RECOMENDADO", "P5-Compliance",
            "Permitir marcação de sigilo local, destaque de prioridade e etiquetas customizáveis por caso.")

    h2(doc, "5.2 Módulo de ingestão e normalização documental")
    rf_para(doc, "RF-06", "MANDATÓRIO", "P3-Stack",
            "Extrair texto e metadados de arquivos PDF pesquisáveis e suportar OCR opcional em documentos digitalizados, com indicação de qualidade do OCR.")
    rf_para(doc, "RF-07", "MANDATÓRIO", "P2-Architecture",
            "Detectar a entidade emissora do documento ou evento, distinguindo juízo, parte autora, parte ré, Ministério Público, perito, cliente e origem interna.")
    rf_para(doc, "RF-08", "MANDATÓRIO", "P2-Architecture",
            "Identificar automaticamente elementos relevantes como número CNJ, partes, classe, tribunal, prazos, pedidos, fundamentos legais e anexos citados.")
    rf_para(doc, "RF-09", "MANDATÓRIO", "P1-Business",
            "Permitir confirmação e correção manual dos metadados extraídos automaticamente antes da indexação final.")
    rf_para(doc, "RF-10", "MANDATÓRIO", "P6-Risk",
            "Registrar falhas de extração em log auditável e permitir reprocessamento manual do item.")

    h2(doc, "5.3 Módulo pré-processual")
    rf_para(doc, "RF-11", "MANDATÓRIO", "P1-Business",
            "Permitir triagem do cliente com cadastro de dados essenciais, histórico jurídico prévio, conflitos conhecidos e data do primeiro contato.")
    rf_para(doc, "RF-12", "MANDATÓRIO", "P5-Compliance",
            "Executar verificação automática de conflito de interesses contra base local de clientes e partes, com alertas e bloqueio quando aplicável.")
    rf_para(doc, "RF-13", "RECOMENDADO", "P1-Business",
            "Permitir elaboração de notificação extrajudicial com modelos parametrizáveis, contendo apresentação do caso, proposta de resolução e prazos de resposta.")
    rf_para(doc, "RF-14", "RECOMENDADO", "P1-Business",
            "Permitir elaboração de minuta de acordo ou transação com cláusulas padrão e específicas, incluindo sigilo, desistência e custas.")
    rf_para(doc, "RF-15", "MANDATÓRIO", "P6-Risk",
            "Manter checklist de prescrição e decadência com cálculo automático de prazos legais e alertas de aproximação (30, 60 e 90 dias).")
    rf_para(doc, "RF-16", "RECOMENDADO", "P1-Business",
            "Registrar consulta inicial com parecer jurídico preliminar e análise de viabilidade.")

    h2(doc, "5.4 Pesquisa assistida de legislação e jurisprudência")
    rf_para(doc, "RF-17", "MANDATÓRIO", "P1-Business",
            "Permitir pesquisa por tema, artigo de lei, tese, tribunal, classe processual, palavra-chave e contexto do caso.")
    rf_para(doc, "RF-18", "MANDATÓRIO", "P3-Stack",
            "Priorizar fontes oficiais e bases configuradas pelo usuário, com hierarquia de confiança parametrizável.")
    rf_para(doc, "RF-19", "MANDATÓRIO", "P5-Compliance",
            "Exibir junto ao resultado a fonte consultada, data da consulta, resumo, trechos relevantes e referência quando disponível.")
    rf_para(doc, "RF-20", "MANDATÓRIO", "P1-Business",
            "Associar resultados relevantes ao dossiê do caso com tag de uso pretendido (fundamento, contra-argumento, precedente).")
    rf_para(doc, "RF-21", "RECOMENDADO", "P2-Architecture",
            "Permitir filtros de jurisprudência por tribunal superior, precedentes qualificados, recência e aderência ao caso.")
    rf_para(doc, "RF-22", "RECOMENDADO", "P2-Architecture",
            "Suportar integração automática à API Pública do DataJud para enriquecimento de decisões mencionadas no caso.")

    h2(doc, "5.5 Geração assistida de peças e documentos")
    rf_para(doc, "RF-23", "MANDATÓRIO", "P1-Business",
            "Oferecer modelos de peças por tipo de ação cível (contestação, réplica, recurso, memoriais, embargos).")
    rf_para(doc, "RF-24", "MANDATÓRIO", "P2-Architecture",
            "Permitir edição colaborativa local com histórico de versões e comparação lado a lado.")
    rf_para(doc, "RF-25", "MANDATÓRIO", "P3-Stack",
            "Gerar rascunho inicial de peça a partir de contexto do caso usando modelo de IA local, sem envio de dados a serviços externos.")
    rf_para(doc, "RF-26", "MANDATÓRIO", "P4-Testing",
            "Permitir validação automática de referências legais citadas (existência do artigo, vigência, revogação).")
    rf_para(doc, "RF-27", "RECOMENDADO", "P1-Business",
            "Permitir exportação em ODT, DOCX e PDF com metadados do caso embutidos.")

    h2(doc, "5.6 Fluxo de trabalho, tarefas e revisão")
    rf_para(doc, "RF-28", "MANDATÓRIO", "P1-Business",
            "Permitir criação de tarefas vinculadas ao caso com responsável, prazo, prioridade e estado.")
    rf_para(doc, "RF-29", "MANDATÓRIO", "P1-Business",
            "Manter fluxo de revisão peer-review antes de exportar peça final.")
    rf_para(doc, "RF-30", "RECOMENDADO", "P4-Testing",
            "Gerar checklist de revisão específico por tipo de peça (fatos, fundamentos, pedidos, valor da causa).")

    h2(doc, "5.7 Motor de cálculos cíveis")
    rf_para(doc, "RF-31", "MANDATÓRIO", "P2-Architecture",
            "Aplicar juros simples/compostos com taxa legal ou convencional configurada.")
    rf_para(doc, "RF-32", "MANDATÓRIO", "P2-Architecture",
            "Aplicar correção monetária por IPCA, INPC, IGP-M, TR ou índice customizável, com tabela oficial atualizada localmente.")
    rf_para(doc, "RF-33", "MANDATÓRIO", "P2-Architecture",
            "Calcular honorários advocatícios por percentual sobre o valor da condenação ou por tabela OAB local.")
    rf_para(doc, "RF-34", "MANDATÓRIO", "P2-Architecture",
            "Calcular custas processuais por tabela do tribunal aplicável.")
    rf_para(doc, "RF-35", "MANDATÓRIO", "P5-Compliance",
            "Gerar memória de cálculo detalhada, exportável em PDF, com fórmulas explícitas e referências normativas.")

    h2(doc, "5.8 Análise de risco e estratégia")
    rf_para(doc, "RF-36", "RECOMENDADO", "P6-Risk",
            "Calcular probabilidade de êxito por tipo de pedido usando modelo local treinado com jurisprudência histórica.")
    rf_para(doc, "RF-37", "RECOMENDADO", "P6-Risk",
            "Sugerir estratégias processuais alternativas com base em precedentes consolidados.")
    rf_para(doc, "RF-38", "RECOMENDADO", "P1-Business",
            "Simular cenários de desfecho (procedência total, parcial, improcedência) com valores financeiros estimados.")

    h2(doc, "5.9 Fluxos específicos do contencioso cível")
    rf_para(doc, "RF-39", "MANDATÓRIO", "P1-Business",
            "Cobrir fluxo de execução: cálculo de dívida, indicação de bens, impugnação, embargos à execução, penhora.")
    rf_para(doc, "RF-40", "MANDATÓRIO", "P1-Business",
            "Cobrir fluxo de conhecimento: inicial, contestação, réplica, audiência, instrução, sentença, recurso.")
    rf_para(doc, "RF-41", "RECOMENDADO", "P1-Business",
            "Cobrir fluxo de tutela provisória: urgência, evidência, contracautelas, efeito suspensivo.")

    h2(doc, "5.10 Gestão de portfólio e painel gerencial")
    rf_para(doc, "RF-42", "MANDATÓRIO", "P1-Business",
            "Exibir painel de carteira com contagem de casos por status, tribunal, cliente e prazo crítico.")
    rf_para(doc, "RF-43", "MANDATÓRIO", "P1-Business",
            "Exibir alertas de prazo processual com D-7, D-3 e D-0 configuráveis.")
    rf_para(doc, "RF-44", "RECOMENDADO", "P1-Business",
            "Gerar relatório mensal de atividade do escritório com horas por caso, peças geradas, pareceres emitidos.")

    h2(doc, "5.11 Integração com cliente e gestão financeira")
    rf_para(doc, "RF-45", "RECOMENDADO", "P1-Business",
            "Registrar contrato de honorários por caso com modalidade (êxito, fixo, hora, misto) e cronograma de pagamento.")
    rf_para(doc, "RF-46", "RECOMENDADO", "P5-Compliance",
            "Emitir recibo e controle de inadimplência com alerta ao responsável.")

    h2(doc, "5.12 Configurações, credenciais e conectores")
    rf_para(doc, "RF-47", "MANDATÓRIO", "P5-Compliance",
            "Gerenciar credenciais de conectores (API keys do DataJud, credenciais de bases privadas opcionais) com criptografia em repouso via Fernet ou equivalente.")
    rf_para(doc, "RF-48", "MANDATÓRIO", "P2-Architecture",
            "Configurar fontes jurisprudenciais ativas, com prioridade e fallback, via UI.")

    h2(doc, "5.13 Auditoria e rastreabilidade")
    rf_para(doc, "RF-49", "MANDATÓRIO", "P5-Compliance",
            "Registrar em trilha imutável toda operação sensível: acesso a documento sigiloso, alteração de peça, exportação, consulta a DataJud, operação financeira.")
    rf_para(doc, "RF-50", "MANDATÓRIO", "P5-Compliance",
            "Garantir integridade da trilha via hash chain (SHA-256) entre registros consecutivos.")

    h2(doc, "5.14 Integração com DataJud (consolidado do documento técnico)")
    rf_para(doc, "RF-51", "MANDATÓRIO", "P2-Architecture",
            "Implementar conector HTTP POST para a API Pública do DataJud do CNJ, com URL base https://api-publica.datajud.cnj.jus.br/ e sufixo /api_publica_{alias}/_search.")
    rf_para(doc, "RF-52", "MANDATÓRIO", "P5-Compliance",
            "Usar autenticação via cabeçalho Authorization: APIKey {chave}, com chave pública publicada pelo CNJ na wiki do DataJud.")
    rf_para(doc, "RF-53", "MANDATÓRIO", "P2-Architecture",
            "Suportar consultas em Query DSL compatível com Elasticsearch (match, bool, term, range, aggregations) conforme glossário de dados do DataJud.")
    rf_para(doc, "RF-54", "MANDATÓRIO", "P2-Architecture",
            "Implementar paginação com size + sort + search_after para grandes volumes, respeitando limite máximo de 10.000 resultados por janela.")
    rf_para(doc, "RF-55", "MANDATÓRIO", "P6-Risk",
            "Implementar circuit breaker e back-off exponencial para falhas da API pública (Beta — instabilidade documentada).")
    rf_para(doc, "RF-56", "MANDATÓRIO", "P5-Compliance",
            "Aderir ao Termo de Uso oficial da API Pública do DataJud: uso institucional, não comercial, sem scraping agressivo.")

    h2(doc, "5.15 Aliases de tribunais suportados")
    para(doc, "A API Pública do DataJud expõe consultas por alias do tribunal. A Solução suporta todos os aliases abaixo (fonte: documento de endpoints consolidado).")
    para(doc, "Tribunais superiores: api_publica_tst, api_publica_tse, api_publica_stj, api_publica_stm.")
    para(doc, "Justiça Federal: api_publica_trf1..api_publica_trf6.")
    para(doc, "Justiça Estadual: api_publica_tjac, tjal, tjam, tjap, tjba, tjce, tjdft, tjes, tjgo, tjma, tjmg, tjms, tjmt, tjpa, tjpb, tjpe, tjpi, tjpr, tjrj, tjrn, tjro, tjrr, tjrs, tjsc, tjse, tjsp, tjto (27 tribunais).")
    para(doc, "Justiça do Trabalho: api_publica_trt1..api_publica_trt24.")
    rf_para(doc, "RF-57", "MANDATÓRIO", "P2-Architecture",
            "Expor seletor de tribunal no frontend alimentado por tabela local de aliases (sem hardcode no código — tabela versionada).")

    h2(doc, "5.16 Anexo regional — Execução Fiscal do Estado do Paraná")
    rf_para(doc, "RF-58", "RECOMENDADO", "P5-Compliance",
            "Aplicar Lei Estadual do Paraná 15.552/2007 (dívida ativa estadual) como módulo plugável ativado por configuração regional.")
    rf_para(doc, "RF-59", "RECOMENDADO", "P2-Architecture",
            "Implementar conector de consulta ao TJPR (alias api_publica_tjpr) para enriquecimento de processos de execução fiscal estadual.")
    rf_para(doc, "RF-60", "RECOMENDADO", "P6-Risk",
            "Auditar separadamente as operações do módulo de execução fiscal por exigência de compliance fiscal estadual (transparência pública).")

    page_break(doc)

    # ─── 6. Requisitos não-funcionais ──────────────────────────
    h1(doc, "6. Requisitos não-funcionais")
    rf_para(doc, "RNF-01", "MANDATÓRIO", "P2-Architecture",
            "Tempo de resposta para abertura de dossiê ≤ 1 segundo até 5.000 documentos; até 3 segundos até 50.000 documentos.")
    rf_para(doc, "RNF-02", "MANDATÓRIO", "P3-Stack",
            "Aplicativo disponível para Windows 10/11 e Ubuntu 22.04+, como binário instalável.")
    rf_para(doc, "RNF-03", "MANDATÓRIO", "P5-Compliance",
            "Criptografia em repouso AES-256 do banco local e dos anexos do dossiê.")
    rf_para(doc, "RNF-04", "MANDATÓRIO", "P5-Compliance",
            "Aderência à LGPD: consentimento explícito na primeira abertura, política de minimização, portabilidade via exportação completa, retenção configurável por caso.")
    rf_para(doc, "RNF-05", "MANDATÓRIO", "P6-Risk",
            "Disponibilidade operacional alvo: 99,5% — o advogado não pode perder o caso por indisponibilidade da ferramenta local.")
    rf_para(doc, "RNF-06", "MANDATÓRIO", "P4-Testing",
            "Cobertura de testes unitários ≥ 80% nos módulos de cadastro, cálculo e auditoria; testes de integração para todos os conectores.")
    rf_para(doc, "RNF-07", "MANDATÓRIO", "P5-Compliance",
            "Trilhas de auditoria imutáveis, com hash chain entre registros consecutivos e ferramenta de verificação offline.")
    rf_para(doc, "RNF-08", "RECOMENDADO", "P3-Stack",
            "Instalação não-privilegiada quando possível (instalar em ~/Applications/ no Linux, %LOCALAPPDATA% no Windows).")
    rf_para(doc, "RNF-09", "MANDATÓRIO", "P6-Risk",
            "Tolerância a falhas: ao perder conexão no meio de uma consulta DataJud, permitir retomada sem duplicação de operações.")

    # ─── 7. Módulos do sistema ─────────────────────────────────
    h1(doc, "7. Módulos do sistema")
    para(doc, "Cada módulo lógico da Solução, com prioridade, pilares impactados, "
              "dependências e flag ready_for_codegen.", italic=True, size=10)

    module_para(doc, "7.1", "Dossiê e cadastro do caso",
                "Cadastro, triagem, anexos, linha do tempo, conflito de interesses, prescrição/decadência. "
                "Cobre RF-01 a RF-05 e RF-11 a RF-16. Módulo central do sistema.",
                "P1-Business, P2-Architecture", "MANDATÓRIO", "Nenhuma", True)

    module_para(doc, "7.2", "Ingestão e normalização documental",
                "Upload, OCR, extração de metadados, classificação, correção manual, reprocessamento. "
                "Cobre RF-06 a RF-10. Alimenta dossiê e pesquisa.",
                "P2-Architecture, P3-Stack", "MANDATÓRIO", "Depende de 7.1", True)

    module_para(doc, "7.3", "Pesquisa de legislação e jurisprudência",
                "Busca por tema/tribunal/tese, priorização de fontes, integração DataJud (leitura), "
                "filtros por precedente. Cobre RF-17 a RF-22.",
                "P1-Business, P2-Architecture", "MANDATÓRIO", "Depende de 7.1 e 7.10", True)

    module_para(doc, "7.4", "Geração assistida de peças",
                "Modelos parametrizáveis por tipo de ação, IA local, edição com versões, validação de "
                "referências legais. Cobre RF-23 a RF-27.",
                "P1-Business, P2-Architecture, P3-Stack", "MANDATÓRIO", "Depende de 7.1 e 7.3", True)

    module_para(doc, "7.5", "Fluxo de trabalho e revisão",
                "Tarefas vinculadas a caso, peer-review, checklist por tipo de peça. Cobre RF-28 a RF-30.",
                "P1-Business, P4-Testing", "MANDATÓRIO", "Depende de 7.1", True)

    module_para(doc, "7.6", "Motor de cálculos cíveis",
                "Juros, correção monetária, honorários, custas, memória de cálculo exportável. Cobre RF-31 a RF-35.",
                "P2-Architecture, P5-Compliance", "MANDATÓRIO", "Nenhuma (módulo independente)", True)

    module_para(doc, "7.7", "Análise de risco e estratégia",
                "Probabilidade de êxito, sugestão de estratégias, simulação de cenários. Cobre RF-36 a RF-38.",
                "P6-Risk", "RECOMENDADO", "Depende de 7.3 (histórico jurisprudencial)", False)

    module_para(doc, "7.8", "Fluxos do contencioso cível",
                "Execução, conhecimento e tutela provisória. Cobre RF-39 a RF-41.",
                "P1-Business", "MANDATÓRIO", "Depende de 7.1 e 7.4", True)

    module_para(doc, "7.9", "Painel gerencial",
                "Portfólio, alertas de prazo, relatórios. Cobre RF-42 a RF-44.",
                "P1-Business", "MANDATÓRIO", "Depende de 7.1", True)

    module_para(doc, "7.10", "Conector DataJud",
                "Integração com API Pública CNJ, autenticação APIKey, Query DSL, paginação com "
                "search_after, circuit breaker, tabela de aliases. Cobre RF-22 e RF-51 a RF-57.",
                "P2-Architecture, P5-Compliance, P6-Risk", "MANDATÓRIO",
                "Nenhuma (módulo independente consumido por 7.3)", True)

    module_para(doc, "7.11", "Auditoria e rastreabilidade",
                "Trilha imutável de operações sensíveis com hash chain. Cobre RF-49 a RF-50 e RNF-07.",
                "P5-Compliance", "MANDATÓRIO", "Transversal — todos os módulos gravam eventos aqui", True)

    module_para(doc, "7.12", "Configurações e conectores",
                "Gestão de credenciais criptografadas, configuração de fontes, ajustes de UI. Cobre RF-47 a RF-48.",
                "P5-Compliance", "MANDATÓRIO", "Nenhuma", True)

    module_para(doc, "7.13", "Anexo Paraná — Execução Fiscal Estadual",
                "Lei 15.552/2007, conector TJPR específico, auditoria separada. Cobre RF-58 a RF-60.",
                "P5-Compliance, P6-Risk", "RECOMENDADO",
                "Depende de 7.10 (conector DataJud) e 7.11 (auditoria)", False)

    page_break(doc)

    # ─── 8. Entregáveis ────────────────────────────────────────
    h1(doc, "8. Entregáveis")
    para(doc,
         "Lista de entregáveis concretos por fase. Cada entregável em parágrafo "
         "próprio com prefixo ENTREGÁVEL:, que o Arguidor usa para popular o "
         "campo DELIVERABLES do OCG e gerar items de backlog.",
         italic=True, size=10)

    h2(doc, "8.1 Fase 1 — Versão inicial homologável")
    delivery(doc, "Instalador Windows", "Instalador .exe via Inno Setup com wizard de 8 passos incluindo aceite de termos, pasta de instalação e criação do primeiro usuário.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Instalador Ubuntu", "Pacote .deb + script install.sh interativo com as mesmas 8 etapas.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Schema local criptografado", "Migrations aplicadas em SQLite/PostgreSQL embarcado com AES-256 em repouso.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Módulo de cadastro e dossiê", "CRUD de casos, anexos, linha do tempo, alerta de prescrição/decadência.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Módulo de ingestão documental básico", "Upload, OCR, extração de metadados com correção manual.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Conector DataJud (leitura)", "Implementação completa do conector HTTP POST com autenticação APIKey e paginação.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Motor de cálculos essencial", "Juros, correção, honorários, custas com memória de cálculo exportável.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Geração de peças assistida (modelos básicos)", "Contestação, réplica, recurso com IA local.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Painel gerencial básico", "Carteira, alertas de prazo D-7/D-3/D-0.", "Fase 1", "MANDATÓRIO")
    delivery(doc, "Trilhas de auditoria com hash chain", "Registro imutável de operações sensíveis + ferramenta de verificação offline.", "Fase 1", "MANDATÓRIO")

    h2(doc, "8.2 Fase 2 — Evolução funcional")
    delivery(doc, "Análise de risco e simulação", "Probabilidade de êxito por tipo de pedido usando modelo local treinado.", "Fase 2", "RECOMENDADO")
    delivery(doc, "Gestão de portfólio avançada", "Relatórios mensais, KPIs, métricas de atividade do escritório.", "Fase 2", "RECOMENDADO")
    delivery(doc, "Modelos de peças expandidos", "Cobertura de todos os fluxos cíveis (conhecimento, execução, tutela provisória).", "Fase 2", "MANDATÓRIO")
    delivery(doc, "Gestão financeira", "Honorários, recibos, controle de inadimplência.", "Fase 2", "RECOMENDADO")
    delivery(doc, "Integração com bases jurídicas privadas", "Conectores plugáveis para bases pagas opcionais.", "Fase 2", "RECOMENDADO")

    h2(doc, "8.3 Fase 3 — Maturidade institucional")
    delivery(doc, "Sincronização multi-dispositivo", "Sync opcional local-to-local ou via servidor do cliente.", "Fase 3", "RECOMENDADO")
    delivery(doc, "Assinatura digital ICP-Brasil", "Integração com certificado A1/A3 para assinar peças geradas.", "Fase 3", "RECOMENDADO")
    delivery(doc, "Integração contábil completa", "Exportação para ERPs contábeis (Omie, ContaAzul, Conta Simples).", "Fase 3", "RECOMENDADO")
    delivery(doc, "Analítica do escritório", "Dashboards consolidados com comparativos temporais e benchmarks.", "Fase 3", "RECOMENDADO")

    h2(doc, "8.4 Fase 4 — Expansões futuras")
    delivery(doc, "Peticionamento eletrônico automatizado", "Envio de peças para tribunais que aceitem API — não em V1 pelo risco.", "Fase 4", "RECOMENDADO")
    delivery(doc, "Marketplace de modelos de peças", "Loja interna para advogados compartilharem/venderem modelos.", "Fase 4", "RECOMENDADO")
    delivery(doc, "Especialização por domínios jurídicos", "Módulos adicionais (trabalhista, tributário, criminal) como plugins.", "Fase 4", "RECOMENDADO")

    # ─── 9. Roadmap ───────────────────────────────────────────
    h1(doc, "9. Roadmap")
    para(doc,
         "Sequenciamento técnico de entregas. O Arguidor usa esta seção para "
         "organizar module_candidates em fases priorizadas no Roadmap do projeto.",
         italic=True, size=10)

    h2(doc, "Fase 1 — Versão inicial homologável")
    para(doc, "Duração estimada: 6 meses. Objetivo: entregar versão mínima viável que cobre o fluxo "
              "completo de um caso cível (cadastro → ingestão → pesquisa → peça → cálculo → auditoria).")
    bullet(doc, "Instalador Windows e Ubuntu")
    bullet(doc, "Schema local criptografado")
    bullet(doc, "Módulo de cadastro e dossiê (módulo 7.1)")
    bullet(doc, "Ingestão documental básica com OCR (módulo 7.2)")
    bullet(doc, "Pesquisa jurisprudencial (módulo 7.3)")
    bullet(doc, "Geração de peças com modelos essenciais (módulo 7.4)")
    bullet(doc, "Fluxo de trabalho (módulo 7.5)")
    bullet(doc, "Motor de cálculos essencial (módulo 7.6)")
    bullet(doc, "Fluxos cíveis cobertos (módulo 7.8)")
    bullet(doc, "Painel gerencial básico (módulo 7.9)")
    bullet(doc, "Conector DataJud completo (módulo 7.10)")
    bullet(doc, "Auditoria com hash chain (módulo 7.11)")
    bullet(doc, "Configurações de conectores (módulo 7.12)")

    h2(doc, "Fase 2 — Evolução funcional")
    para(doc, "Duração estimada: 4 meses após Fase 1. Objetivo: adicionar inteligência de risco e "
              "gestão financeira.")
    bullet(doc, "Análise de risco e estratégia (módulo 7.7)")
    bullet(doc, "Gestão de portfólio avançada (extensão do módulo 7.9)")
    bullet(doc, "Gestão financeira (módulo 7.11 ampliado — integração com honorários)")
    bullet(doc, "Integração com bases jurídicas privadas (extensão do módulo 7.3)")
    bullet(doc, "Anexo Paraná — Execução Fiscal Estadual (módulo 7.13)")

    h2(doc, "Fase 3 — Maturidade institucional")
    para(doc, "Duração estimada: 6 meses após Fase 2. Objetivo: tornar a Solução apta a escritórios "
              "de porte médio.")
    bullet(doc, "Sincronização multi-dispositivo opcional")
    bullet(doc, "Colaboração multiusuário local")
    bullet(doc, "Assinatura digital ICP-Brasil")
    bullet(doc, "Integração contábil completa")
    bullet(doc, "Analítica do escritório")

    h2(doc, "Fase 4 — Expansões futuras")
    para(doc, "Sem cronograma fixo. Depende de demanda de mercado e contratos específicos.")
    bullet(doc, "Peticionamento eletrônico automatizado")
    bullet(doc, "Especialização por domínios jurídicos (trabalhista, tributário, criminal)")
    bullet(doc, "Marketplace de modelos")

    # ─── 10. Riscos ───────────────────────────────────────────
    h1(doc, "10. Riscos conhecidos")
    para(doc, "RISCO-01: API Pública do DataJud em fase Beta — pode mudar contrato ou ficar instável. "
              "Probabilidade: média. Impacto: alto (quebra pesquisa). "
              "Mitigação: circuit breaker, cache local agressivo, monitoramento de status.")
    para(doc, "RISCO-02: IA local requer hardware mínimo não atendido por advogado com notebook antigo. "
              "Probabilidade: média. Impacto: médio (degrada UX mas não quebra). "
              "Mitigação: degradar graciosamente para modelos menores; oferecer fallback sem IA.")
    para(doc, "RISCO-03: Dados do cliente expostos se o notebook do advogado for comprometido. "
              "Probabilidade: baixa. Impacto: crítico (LGPD). "
              "Mitigação: AES-256 em repouso, senha obrigatória na primeira abertura, opção de wipe remoto.")
    para(doc, "RISCO-04: Mudança regulatória (LGPD, OAB) que invalide capacidade atual. "
              "Probabilidade: baixa. Impacto: médio. Mitigação: modularidade, monitoramento jurídico.")
    para(doc, "RISCO-05: Divergência na modelagem de Query DSL do DataJud entre tribunais. "
              "Probabilidade: alta. Impacto: médio. "
              "Mitigação: testar cada alias individualmente; tabela de quirks por tribunal.")

    # ─── 11. Critérios de aceite ──────────────────────────────
    h1(doc, "11. Critérios de aceite")
    numbered(doc, "Todos os RFs [MANDATÓRIO] implementados e validados em testes automatizados.")
    numbered(doc, "Cobertura de testes unitários ≥ 80% nos módulos de cadastro, cálculo e auditoria.")
    numbered(doc, "Teste de integração com ao menos 5 tribunais reais via DataJud.")
    numbered(doc, "Instalador funciona em Windows 10, Windows 11 e Ubuntu 22.04+.")
    numbered(doc, "Banco local criptografado verificável por ferramenta externa de inspeção.")
    numbered(doc, "Trilha de auditoria validável offline via ferramenta do próprio produto.")
    numbered(doc, "Documentação técnica completa: README de setup, guia do usuário, ADRs.")
    numbered(doc, "Aprovação do Gatekeeper do GCA em todos os 7 pilares com score ≥ 75.")
    numbered(doc, "Homologação com 3 escritórios-piloto por 30 dias antes de GA.")

    # ─── 12. Integrações externas ─────────────────────────────
    h1(doc, "12. Integrações externas")

    h2(doc, "12.1 API Pública do DataJud (CNJ)")
    para(doc, "Fonte: Base Nacional de Dados do Poder Judiciário, operada pelo CNJ.")
    para(doc, "URL base: https://api-publica.datajud.cnj.jus.br/")
    para(doc, "Natureza: pública, Beta, não-comercial.")
    para(doc, "Método: HTTP POST para /api_publica_{alias}/_search com corpo JSON em Query DSL compatível com Elasticsearch.")
    para(doc, "Autenticação: cabeçalho Authorization: APIKey {chave} — chave pública publicada pelo CNJ na wiki do DataJud.")
    para(doc, "Paginação: size + sort + search_after, limite de 10.000 resultados por janela.")
    para(doc, "Limites: taxa informal recomendada de 100 req/min por alias (não documentada oficialmente).")
    para(doc, "Riscos: Beta — contrato pode mudar, quotas podem ser impostas sem aviso prévio, aliases podem ser descontinuados.")
    para(doc, "Uso no projeto: enriquecer pesquisa jurisprudencial (módulo 7.3) e análise de risco (módulo 7.7) com dados processuais públicos.")

    h2(doc, "12.2 Tabelas oficiais de índices econômicos")
    para(doc, "Fonte: Banco Central do Brasil (Selic), IBGE (IPCA, INPC), FGV (IGP-M, TR).")
    para(doc, "Método: download periódico de tabelas CSV/XLS oficiais para atualização local.")
    para(doc, "Frequência: diária para Selic; mensal para demais.")
    para(doc, "Uso no projeto: motor de cálculos cíveis (módulo 7.6).")

    h2(doc, "12.3 Tabelas de custas dos tribunais")
    para(doc, "Fonte: sites oficiais dos tribunais (TJPR, TJSP, STJ, etc.).")
    para(doc, "Método: download manual (não há API) + revisão trimestral de atualização.")
    para(doc, "Uso no projeto: motor de cálculos (RF-34).")

    # ─── 13. Conformidade ─────────────────────────────────────
    h1(doc, "13. Conformidade, segurança e privacidade")
    bullet(doc, "LGPD (Lei 13.709/2018): consentimento explícito no primeiro uso, minimização de coleta, política de retenção configurável por caso, portabilidade via exportação completa, DPO de referência do escritório.")
    bullet(doc, "Segurança: AES-256 em repouso, TLS 1.3 em trânsito, credenciais criptografadas com Fernet, senha obrigatória com bcrypt + salt.")
    bullet(doc, "Auditoria: trilhas imutáveis com hash chain, retenção mínima de 5 anos conforme recomendação OAB.")
    bullet(doc, "Código de Ética da OAB: sigilo profissional, vedação de compartilhamento de informações sem consentimento do cliente.")
    bullet(doc, "Aderência ao Termo de Uso da API Pública do DataJud.")

    page_break(doc)
    h1(doc, "Anexo A — Mudanças da v1.0 para v2.0")
    bullet(doc, "Todos os RFs saíram de tabelas (ilegíveis para o Arguidor do GCA) e viraram parágrafos numerados com prefixos [MANDATÓRIO/RECOMENDADO] + [Pilar].")
    bullet(doc, "Seção 5 expandida de 50 para 60 RFs com incorporação dos 6 RFs de integração DataJud (RF-51 a RF-57) e anexo Paraná (RF-58 a RF-60).")
    bullet(doc, "Seção 7 (Módulos) nova — cada módulo com prefixo MÓDULO: legível pelo Arguidor.")
    bullet(doc, "Seção 8 (Entregáveis) nova — prefixo ENTREGÁVEL: para popular DELIVERABLES do OCG.")
    bullet(doc, "Seção 9 (Roadmap) reformulada com fases numeradas e bullets (não tabela).")
    bullet(doc, "Seção 12 (Integrações) nova — consolida DataJud e tabelas econômicas.")
    bullet(doc, "Anexo de endpoints de tribunais incorporado em 5.15 como tabela de aliases suportados.")
    bullet(doc, "Glossário movido para início (seção 4) e expandido com termos do domínio jurídico.")

    h1(doc, "Anexo B — Como testar se a ingestão funcionou")
    numbered(doc, "Após ingerir este documento em /projects/{id}/ingestion, aguardar 2-3 minutos.")
    numbered(doc, "Abrir /projects/{id}/ocg e verificar se o campo DELIVERABLES foi populado.")
    numbered(doc, "Abrir /projects/{id}/backlog e verificar items de source='ocg' (entregáveis) e source='arguider' (módulos candidatos).")
    numbered(doc, "Abrir /projects/{id}/roadmap e verificar fases 1 a 4 com entregas por fase.")
    numbered(doc, "Se alguma etapa não gerou itens: verificar log do Arguidor em /projects/{id}/audit — possível problema de prompt ou de provider IA.")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    # Template universal
    doc1 = Document()
    s = doc1.sections[0]
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2)
    build_template(doc1)
    doc1.save(TEMPLATE_PATH)
    size1 = TEMPLATE_PATH.stat().st_size / 1024
    print(f"✓ Template gerado: {TEMPLATE_PATH.name} ({size1:.1f} KB)")

    # Automação Jurídica v2.0
    doc2 = Document()
    s = doc2.sections[0]
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2)
    build_auto_juridica_v2(doc2)
    doc2.save(V2_PATH)
    size2 = V2_PATH.stat().st_size / 1024
    print(f"✓ Automação Jurídica v2.0: {V2_PATH.name} ({size2:.1f} KB)")


if __name__ == "__main__":
    main()
