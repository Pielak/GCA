#!/usr/bin/env python3
"""
Gera GCA_Requisitos_v1.docx — documento de requisitos completo do GCA.

Estrutura:
  1. Apresentação
  2. Visão geral do produto
  3. Regras de negócio
  4. Arquitetura técnica
  5. RBAC e papéis canônicos
  6. MVPs entregues (1-7)
  7. Modelo de dados (schema)
  8. API — endpoints principais
  9. Diagramas de sequência
 10. Diagramas de fluxo e estado
 11. Requisitos não-funcionais
 12. Operação e manutenção
 13. Segurança e proteção de código
 14. Glossário (ordem alfabética)

Autor: Luiz Carlos Pielak
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

DIAG_DIR = Path("/home/luiz/GCA/docs/diagrams")
OUT_PATH = Path("/home/luiz/GCA/docs/GCA_Requisitos_v1.docx")

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
SLATE_DARK = RGBColor(0x1E, 0x29, 0x3B)
SLATE_MEDIUM = RGBColor(0x64, 0x74, 0x8B)
EMERALD = RGBColor(0x05, 0x96, 0x69)
AMBER = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)


# ─── Helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = VIOLET
        run.font.size = Pt(20)


def h2(doc, text: str):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = SLATE_DARK
        run.font.size = Pt(15)


def h3(doc, text: str):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = SLATE_MEDIUM
        run.font.size = Pt(12)


def para(doc, text: str, *, bold: bool = False, size: int = 11,
         color: RGBColor = None, justify: bool = True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text: str, *, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


def numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)


def code(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = SLATE_DARK


def table(doc, headers: list[str], rows: list[list[str]],
          widths: list[float] = None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 4"
    t.autofit = False

    # Header
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "6D28D9")
        if widths and i < len(widths):
            cell.width = Cm(widths[i])

    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if widths and ci < len(widths):
                cell.width = Cm(widths[ci])
    return t


def page_break(doc):
    doc.add_page_break()


def insert_diagram(doc, filename: str, caption: str, width_in: float = 6.3):
    path = DIAG_DIR / filename
    if not path.exists():
        para(doc, f"[diagrama ausente: {filename}]", color=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figura — {caption}")
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = SLATE_MEDIUM


# ─── Documento ────────────────────────────────────────────────────────────

def build(doc: Document):
    # Propriedades
    core = doc.core_properties
    core.title = "GCA — Documento de Requisitos"
    core.author = "Luiz Carlos Pielak"
    core.subject = "Especificação técnica e regras de negócio do GCA"
    core.keywords = "GCA, requisitos, arquitetura, RBAC, OCG, FastAPI, React"
    core.last_modified_by = "Luiz Carlos Pielak"
    core.created = datetime.now()

    # Capa
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("GCA")
    tr.font.size = Pt(60)
    tr.bold = True
    tr.font.color.rgb = VIOLET

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Gestão de Codificação Assistida")
    sr.font.size = Pt(20)
    sr.font.color.rgb = SLATE_DARK

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = desc.add_run("Documento de Requisitos")
    dr.font.size = Pt(22)
    dr.italic = True
    dr.font.color.rgb = SLATE_MEDIUM

    for _ in range(8):
        doc.add_paragraph()

    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = auth.add_run("Autor: Luiz Carlos Pielak")
    ar.font.size = Pt(13)
    ar.bold = True

    dateline = doc.add_paragraph()
    dateline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt = dateline.add_run(datetime.now().strftime("Versão 1.0 — %d de %B de %Y").replace(
        "January", "janeiro").replace("February", "fevereiro").replace("March", "março")
        .replace("April", "abril").replace("May", "maio").replace("June", "junho")
        .replace("July", "julho").replace("August", "agosto").replace("September", "setembro")
        .replace("October", "outubro").replace("November", "novembro").replace("December", "dezembro"))
    dt.font.size = Pt(11)
    dt.font.color.rgb = SLATE_MEDIUM

    page_break(doc)

    # ─── 1. Apresentação ──────────────────────────────────────────────
    h1(doc, "1. Apresentação")
    para(doc,
         "Este documento consolida os requisitos funcionais, as regras de negócio, a arquitetura "
         "técnica, o modelo de dados, o desenho de papéis e os diagramas comportamentais da "
         "plataforma GCA (Gestão de Codificação Assistida / Gerenciador Central de Arquiteturas), "
         "um sistema instalável por cliente para governança de projetos de TI assistida por Inteligência Artificial.")
    para(doc,
         "O GCA foi desenvolvido em ciclos sequenciais (MVPs 1 a 7) com emendas formais sempre que "
         "o escopo canônico precisou ser ampliado por decisão do stakeholder-soberano. Esta versão "
         "1.0 do documento é produzida em 2026-04-19, logo após o fechamento do MVP 7 (entrega "
         "versionada preservando dados do usuário), sua emenda RBAC (GP soberano do projeto) e a "
         "consolidação do sistema de métricas admin com breakdown compartimentalizado por projeto.")
    para(doc,
         "O texto deste documento é binário em relação aos requisitos: quando diz \"deve\", o "
         "comportamento é obrigatório; quando diz \"não deve\", é proibido. Ausência de menção não "
         "equivale a permissão — o que não está previsto aqui ou no contrato canônico não "
         "integra a fase atual do produto.")

    h2(doc, "1.1. Público-alvo do documento")
    bullet(doc, "Arquitetos e engenheiros que vão evoluir o GCA.")
    bullet(doc, "Time de operações responsável pela instalação em clientes.")
    bullet(doc, "Auditores internos e compliance que precisam entender o modelo de dados e o fluxo de auditoria.")
    bullet(doc, "Gerentes de Projeto (GPs) e administradores operacionais que queiram entender limites e garantias do produto.")

    h2(doc, "1.2. Documentos relacionados")
    bullet(doc, "GCA_CANONICAL_CONTRACT.md — contrato canônico soberano; em caso de conflito, prevalece sobre este documento.")
    bullet(doc, "GCA_MVP_PROGRESS.md — rastreabilidade dos MVPs entregues e dívidas técnicas quitadas.")
    bullet(doc, "GCA_Tutorial_Instalacao_v1.docx — tutorial passo a passo de instalação em Windows e Ubuntu, gerado em paralelo a este.")
    bullet(doc, "docs/ANTI_REVERSE_ENGINEERING.md — nota técnica sobre proteção do código binário.")

    page_break(doc)

    # ─── 2. Visão geral do produto ────────────────────────────────────
    h1(doc, "2. Visão geral do produto")

    h2(doc, "2.1. Definição")
    para(doc,
         "O GCA é uma plataforma instalável por cliente, auto-hospedada (on-premises ou em "
         "infraestrutura privada do cliente), que automatiza a governança de projetos de "
         "desenvolvimento de software. O produto cobre o ciclo do projeto desde a solicitação inicial "
         "até a entrega do código, passando por questionário de requisitos, geração de OCG (Objeto "
         "Canônico de Governança), avaliação por Gatekeeper com sete pilares, ingestão de artefatos, "
         "Arguidor, geração de código assistida, Quality Assurance Readiness, Documentação Viva e "
         "Release Bundle.")

    h2(doc, "2.2. Modelo de deployment")
    bullet(doc, "Uma instância por cliente — sem multi-tenant compartilhado.")
    bullet(doc, "Isolamento principal por projeto dentro da instância.")
    bullet(doc, "Cada cliente usa seus próprios provedores de IA, chaves e integrações.")
    bullet(doc, "Sem marketplace central nem operação SaaS compartilhada nesta versão.")

    h2(doc, "2.3. Princípio binário")
    para(doc,
         "Toda regra deste produto é binária: sim ou não, deve ou não deve, tem ou não tem. "
         "Qualquer ambiguidade em prompt de IA, em documentação, em RFC ou em decisão de gate é "
         "tratada como defeito do documento, não como zona cinzenta legítima.")

    h2(doc, "2.4. Roteamento híbrido de IA por criticidade")
    bullet(doc, "Baixa criticidade (classificação, extração, sumarização curta): modelo local (Ollama) ou modelo barato.")
    bullet(doc, "Média criticidade (propostas iniciais, pré-análise): local ou remoto com validação.")
    bullet(doc, "Alta criticidade (consolidação do OCG, decisão arquitetural, compliance crítico, release): modelo premium de raciocínio obrigatório.")

    page_break(doc)

    # ─── 3. Regras de negócio ─────────────────────────────────────────
    h1(doc, "3. Regras de negócio")

    h2(doc, "3.1. Criação e aprovação de projetos")
    numbered(doc, "Solicitante externo submete formulário público em /solicitar-projeto com wizard de 2 passos (dados de identificação + questionário breve).")
    numbered(doc, "Solicitação entra em project_requests com status PENDING e dispara notificação para administradores ativos.")
    numbered(doc, "Admin aprova ou rejeita; aprovação cria linha em projects com status=active e promove o solicitante a GP.")
    numbered(doc, "Rejeição registra rejection_reason e notifica solicitante por e-mail.")

    h2(doc, "3.2. Ciclo de vida do projeto")
    para(doc,
         "Todo projeto tem obrigatoriamente um dos três estados de lifecycle:")
    bullet(doc, "active — operacional; scheduler de backup automático ligado.")
    bullet(doc, "paused — suspenso temporariamente; dados preservados; scheduler não roda.")
    bullet(doc, "inactive — encerramento formal sem deleção; dados preservados para consulta/auditoria.")
    para(doc,
         "Apenas Admin pode alterar o estado via PATCH /admin/projects/{id}/status. "
         "Nenhum dos três estados apaga OCG, questionário, backlog, documentos, backups ou tickets — "
         "a área administrativa é soberana em persistir os itens do projeto em qualquer estado.")

    h2(doc, "3.3. Compartimentalização por projeto (contrato §2.2)")
    bullet(doc, "Toda leitura ou escrita envolvendo dado de projeto deve incluir project_id no predicado.")
    bullet(doc, "Nenhum canal lateral (vault, storage, cache, logs, notificações, Git, n8n, SMTP, IA) pode cruzar projetos sem autorização explícita do contrato.")
    bullet(doc, "Admin é o único com visão cross-projeto, restrita a endpoints marcados como globais (/admin/*).")

    h2(doc, "3.4. OCG como fonte de verdade")
    numbered(doc, "O OCG inicia a partir do questionário externo aprovado, através do pipeline de 8 agentes de IA (analyzer + 7 pilares + consolidator).")
    numbered(doc, "O OCG é um objeto de estado evolutivo: expande com boa ingestão e contrai com ingestão ruim, conflitante ou incompleta.")
    numbered(doc, "Nenhuma decisão arquitetural, funcional, de testes ou de geração de código pode ignorar o OCG.")
    numbered(doc, "Toda mudança relevante no OCG é versionada em ocg_delta_log e auditada em audit_log_global.")

    h2(doc, "3.5. Tickets de incidente (MVP 6 + emenda)")
    para(doc,
         "Usuários do projeto podem abrir tickets de incidente com roteamento automático por papel:")
    bullet(doc, "Dev, Tester e QA — abrem ticket direcionado aos GPs do projeto (target_scope=gp).")
    bullet(doc, "GP — abre ticket direcionado a Admins e Equipe Sustentação (target_scope=admin).")
    bullet(doc, "Admin — abre ticket direcionado a demais Admins e Sustentação.")
    para(doc,
         "Todo ticket carrega obrigatoriamente título, descrição, prioridade (baixa/média/alta/crítica), "
         "categoria (bug/dúvida/pedido_feature/incidente_pipeline), seção onde ocorreu (autopreenchida pela rota) "
         "e fluxo executado (obrigatório — modal recusa vazio). Anexos: até 5 arquivos por ticket, "
         "10 MB cada, tipos png, jpg, jpeg, webp, gif, txt, log, json, pdf.")

    h2(doc, "3.6. Releases e preservação de dados")
    para(doc,
         "Toda entrega de software aplicada à instância vira uma release com tag semântica e "
         "changelog visível ao usuário. Releases são declaradas em backend/releases/*.yaml shipado "
         "com o código. Releases não-destrutivas são aplicadas automaticamente no startup. "
         "Releases destrutivas (is_destructive=true) ficam em status pending até o Admin confirmar "
         "em /admin/releases e disparar snapshot automático (DT-063) de cada projeto ativo antes de "
         "efetivar a aplicação.")
    bullet(doc, "Default não-destrutivo: migrações novas preservam dados (coluna nullable ou com default).")
    bullet(doc, "Quando destrutivo: snapshot prévio por projeto + botão de restaurar por projeto.")
    bullet(doc, "Rollback é por-projeto; status global da release permanece \"applied\" quando um projeto reverte.")

    h2(doc, "3.7. Backups compartimentalizados (DT-063)")
    bullet(doc, "Scheduler APScheduler diário às 12:00 (America/Sao_Paulo) com cron 0 12 * * *.")
    bullet(doc, "Catch-up automático no startup quando o último backup > 24h.")
    bullet(doc, "Até 10 backups retidos por projeto; excedentes limpos após cada execução.")
    bullet(doc, "Admin pode disparar backup manual a pedido do GP via quick action em /admin/backups.")
    bullet(doc, "GP pode disparar backup manual em /projects/{id}/backups.")
    bullet(doc, "Apenas Admin ou GP do projeto podem executar restore; restore preserva audit_log_global.")

    h2(doc, "3.8. Equipe Sustentação")
    para(doc,
         "A flag users.is_support define usuários que recebem tickets escalados a Admin. Admin "
         "herda automaticamente os privilégios de Sustentação (pode ver /admin/incidents). Support "
         "puro vê tickets escalados mas não ganha poderes administrativos. Promoção de user comum a "
         "Support é feita pela UI /admin/support; promoção a Admin continua pela gestão canônica "
         "de usuários em /admin/users. Último Admin ativo não pode se auto-rebaixar (anti-órfão).")

    h2(doc, "3.9. Política de IA do cliente final")
    bullet(doc, "Provedor e modelo são configuráveis por instância e por projeto.")
    bullet(doc, "Ollama é suportado via endpoint OpenAI-compatible (base_url obrigatório, api_key opcional).")
    bullet(doc, "Nenhum modelo de baixa criticidade decide sozinho arquitetura, compliance, segurança ou liberação de pipeline.")
    bullet(doc, "Compatibilidade com endpoint estilo OpenAI não é equivalência funcional entre modelos.")
    bullet(doc, "Cada tarefa relevante de IA registra provedor, modelo, criticidade, custo observado em ai_usage_log.")

    page_break(doc)

    # ─── 4. Arquitetura técnica ───────────────────────────────────────
    h1(doc, "4. Arquitetura técnica")

    h2(doc, "4.1. Stack principal")
    table(doc,
          ["Camada", "Tecnologia", "Observações"],
          [
              ["Backend", "Python 3.11 + FastAPI + SQLAlchemy async", "Uvicorn com 4 workers em produção."],
              ["Frontend", "React 18 + TypeScript + Vite + Tailwind CSS", "Vite preview em produção (build estático)."],
              ["Banco", "PostgreSQL 15", "asyncpg driver. DB gca (prod) e gca_test (pytest isolado)."],
              ["Cache", "Redis (opcional)", "Usado em pipelines de ingestão quando configurado."],
              ["Agendamento", "APScheduler AsyncIOScheduler", "Scheduler de backup diário; startup catch-up."],
              ["IA", "Anthropic / OpenAI / Google / DeepSeek / Ollama local", "Configurável por instância e projeto."],
              ["Email", "SMTP compartimentalizado por projeto (DT-016)", "Fallback global quando projeto não configura."],
              ["Storage", "Volumes Docker nomeados", "gca-postgres-data, gca-uploads-storage, gca-backups."],
              ["Auth", "JWT HS256 via python-jose", "Token em memória do frontend; sem refresh token nesta versão."],
          ],
          widths=[3, 6, 6])

    h2(doc, "4.2. Topologia de deployment")
    para(doc,
         "A instância do GCA é composta por seis serviços em Docker Compose:")
    bullet(doc, "gca-backend — FastAPI com volume ./backend:/app em dev, imagem imutável em prod.")
    bullet(doc, "gca-frontend — Node.js servindo build estático via vite preview na porta 5173.")
    bullet(doc, "gca-postgres — PostgreSQL 15 com volume persistente e configuração de locale pt-BR.")
    bullet(doc, "gca-redis — cache opcional.")
    bullet(doc, "gca-ollama — provedor IA local opcional, exposto internamente na porta 11434.")
    bullet(doc, "gca-n8n — orquestrador de workflows opcional (histórico; uso em declínio pós MVP 4).")

    h2(doc, "4.3. Convenções de código")
    bullet(doc, "Backend: tipagem explícita; operações de IO sempre async; nenhum endpoint existente tem contrato alterado sem justificativa e migração compatível.")
    bullet(doc, "Frontend: TypeScript estrito; zustand para estado global; TanStack Query para data fetching; Tailwind obrigatório (sem estilos inline de cor).")
    bullet(doc, "Banco: nenhuma tabela ou coluna existente é removida sem fase de deprecação; toda mudança estrutural exige migração.")
    bullet(doc, "Comentários em código somente quando descrevem o porquê (constraint, invariante, workaround); nunca o quê.")

    h2(doc, "4.4. Topologia da IA")
    para(doc,
         "O GCA separa de forma dura dois contextos de uso de IA:")
    bullet(doc, "Contexto A — desenvolvimento do produto GCA: pode usar IA premium, é custo do time de produto.")
    bullet(doc, "Contexto B — operação da instância do cliente: IA é configurável pelo cliente; GCA não impõe provedor.")
    para(doc,
         "Nenhuma decisão de desenvolvimento pode transformar conveniência interna em obrigação do cliente. Claude, ao evoluir o GCA, "
         "valida sempre \"esta decisão é apenas do meu ambiente ou está virando indevidamente uma obrigação da operação do cliente?\".")

    page_break(doc)

    # ─── 5. RBAC ──────────────────────────────────────────────────────
    h1(doc, "5. RBAC e papéis canônicos")

    h2(doc, "5.1. Papéis")
    table(doc,
          ["Papel", "Escopo", "Função"],
          [
              ["Admin", "Instância", "Opera a instância, configura provedores, aprova projetos, promove Sustentação. Não atua operacionalmente em projetos (papel virtual admin_viewer quando acessa projeto sem membership)."],
              ["Sustentação", "Instância (auxiliar)", "Flag is_support. Recebe tickets escalados a Admin. Não ganha privilégios administrativos."],
              ["GP", "Projeto (soberano)", "Emenda 2026-04-19: soberano do projeto. Tem união das ações de Dev, Tester e QA. Pode escrever código, executar pipeline, aprovar módulos e OCG."],
              ["Dev", "Projeto", "Implementa código, opera ingestão, Arguidor, CodeGen e commits."],
              ["Tester", "Projeto", "Cria, edita e executa testes. Registra evidências."],
              ["QA", "Projeto", "Revisa e aprova resultados. Valida qualidade final. Não edita conteúdo de teste."],
          ],
          widths=[2.5, 3.2, 9])

    h2(doc, "5.2. Assimetria soberana")
    para(doc,
         "Existem duas soberanias paralelas no GCA:")
    bullet(doc, "Admin é soberano da instância (visão cross-projeto, configura provedores, aprova projetos).")
    bullet(doc, "GP é soberano do projeto (dentro do projeto, tem acesso a todas as funcionalidades que Dev, Tester e QA teriam separadamente).")
    para(doc,
         "Admin herda Sustentação (flag implícita). GP nunca herda Admin automaticamente — "
         "promoção a Admin é caminho separado em /admin/users.")

    insert_diagram(doc, "rbac_papeis.png",
                   "Hierarquia de papéis canônicos pós-emenda 2026-04-19.")

    h2(doc, "5.3. Matriz de ações")
    table(doc,
          ["Action", "Admin", "GP", "Dev", "Tester", "QA"],
          [
              ["project:view",        "✓", "✓", "✓", "✓", "✓"],
              ["project:edit",        "✓", "✓", "—", "—", "—"],
              ["project:manage_team", "✓", "✓", "—", "—", "—"],
              ["code:write",          "—", "✓", "✓", "—", "—"],
              ["code:review",         "—", "✓", "✓", "—", "—"],
              ["git:commit",          "—", "✓", "✓", "—", "—"],
              ["pipeline:review",     "—", "✓", "—", "—", "—"],
              ["pipeline:execute",    "—", "✓", "✓", "✓", "—"],
              ["qa:approve",          "—", "✓", "—", "—", "✓"],
              ["security:review",     "—", "✓", "—", "—", "✓"],
              ["compliance:validate", "—", "✓", "—", "—", "✓"],
              ["backlog:manage",      "—", "✓", "—", "—", "—"],
              ["audit:view",          "✓", "✓", "✓", "✓", "✓"],
              ["audit:export",        "✓", "✓", "—", "✓", "—"],
              ["docs:edit",           "—", "✓", "✓", "—", "—"],
          ],
          widths=[4, 1.6, 1.6, 1.6, 1.6, 1.6])

    page_break(doc)

    # ─── 6. MVPs entregues ────────────────────────────────────────────
    h1(doc, "6. MVPs entregues")

    mvps = [
        ("MVP 1", "Base operacional e saneamento do núcleo",
         "Autenticação, RBAC canônico, bootstrap do primeiro Admin, cadastro e aprovação de projetos, "
         "questionário externo e interno, OCG persistido básico, Gatekeeper básico, auditoria mínima, "
         "configuração básica de provedor de IA."),
        ("MVP 2", "Contexto vivo e governança de conteúdo",
         "Ingestão de documentos, quarentena de PII, OCG versionado com deltas, backlog derivado do OCG, "
         "Arguidor, reavaliação do Gatekeeper após ingestão."),
        ("MVP 3", "Geração assistida controlada",
         "CodeGen controlado, preview, geração cirúrgica por arquivo, integração Git, commits rastreáveis, "
         "validação pós-geração, docstrings obrigatórias, análise de adequação de provedor de IA ao CodeGen."),
        ("MVP 4", "Qualidade, documentação e entrega",
         "QA Readiness, execução e revisão de testes, Documentação Viva, Roadmap coerente, Release Bundle, "
         "evidências e relatórios."),
        ("MVP 5", "Hardening operacional",
         "Criptografia de segredos e PATs, hardening de produção, observabilidade complementar "
         "(/metrics/health e /dashboard), rotinas de backup/restore maduras (DT-061/062/063), "
         "melhorias de deploy/upgrade por cliente."),
        ("MVP 6", "Validação assistida em campo",
         "Tickets de incidente com roteamento automático por papel, área administrativa agregada, "
         "comentários e mudança de status, notificação in-app, auditoria compartimentalizada."),
        ("MVP 6 Emenda 2026-04-19", "Sustentação + anexos + contexto",
         "Flag is_support com Admin herdando. Anexos ao ticket (5 arquivos, 10 MB, 9 extensões). "
         "Campo section_reference autopreenchido e flow_description obrigatório."),
        ("MVP 7", "Entrega versionada preservando dados",
         "Registry YAML de releases shipado com o código, aplicação automática de não-destrutivas no startup, "
         "snapshots DT-063 antes de aplicar destrutivas, rollback por projeto, changelog segmentado por papel, "
         "completion tasks pós-release para preenchimento de dados novos."),
    ]
    for codex, nome, desc in mvps:
        h3(doc, f"{codex} — {nome}")
        para(doc, desc)

    para(doc,
         "Todos os sete MVPs estão fechados em 2026-04-19 com 732 testes de regressão passando. "
         "As 34 dívidas técnicas mapeadas no GCA_MVP_PROGRESS.md foram todas quitadas.")

    page_break(doc)

    # ─── 7. Modelo de dados ───────────────────────────────────────────
    h1(doc, "7. Modelo de dados")

    h2(doc, "7.1. Tabelas principais")
    table(doc,
          ["Tabela", "Chave", "Função"],
          [
              ["users", "id", "Usuários da instância. Flags is_admin, is_support, is_active."],
              ["organizations", "id", "Organizações (pré-MVP 3, agrupamento leve)."],
              ["projects", "id", "Projetos com status (active/paused/inactive/archived)."],
              ["project_requests", "id", "Solicitações externas de projeto (PENDING/APPROVED/REJECTED)."],
              ["project_members", "id", "Membership com role (gp/dev/tester/qa) + accepted_at."],
              ["project_member_roles", "id", "Multi-papel por membership (permite usuário ter múltiplos papéis no projeto)."],
              ["project_invites", "id", "Convites pendentes a usuários."],
              ["questionnaires", "id", "Questionários submetidos (PDF + parse + adherence_score)."],
              ["ocg", "id", "OCG do projeto (versionado via ocg_delta_log)."],
              ["ocg_delta_log", "id", "Histórico de mudanças no OCG (trigger, before, after)."],
              ["ocg_analysis_log", "id", "Logs das execuções dos 8 agentes de IA."],
              ["ingested_documents", "id", "Documentos ingeridos e analisados pelo Arguidor."],
              ["backlog_items", "id", "Itens de backlog derivados do OCG + arguidor."],
              ["qa_tests, qa_test_runs, qa_evidences", "id", "Estrutura de QA Readiness + evidências."],
              ["ai_usage_log", "id", "Registro de chamadas de IA (provider, operation, tokens, cost). Compartimentalizado por project_id."],
              ["audit_log_global", "id", "Auditoria com hash chain (previous_hash + current_hash)."],
              ["project_backups", "id", "Backups por projeto (até 10 retidos)."],
              ["incident_tickets", "id", "Tickets de incidente com target_scope (gp/admin)."],
              ["incident_ticket_comments", "id", "Comentários em tickets."],
              ["incident_ticket_attachments", "id", "Anexos de tickets (imagens, logs, texto, PDF)."],
              ["releases", "id", "Releases aplicadas ou pending (is_destructive)."],
              ["release_items", "id", "Items do changelog de uma release (kind, ref_id, affected_roles)."],
              ["release_application_log", "id", "Log de aplicação de release (applied, snapshot_taken, rolled_back)."],
              ["release_completion_tasks", "id", "Tarefas pós-release por projeto."],
              ["user_notifications", "id", "Notificações in-app entregues."],
              ["project_settings", "id", "Configurações do projeto incluindo SMTP compartimentalizado."],
              ["llm_providers", "id", "Provedores de IA configurados (global e por projeto)."],
          ],
          widths=[5, 2, 8])

    h2(doc, "7.2. Integridade auditável")
    para(doc,
         "A tabela audit_log_global implementa chain integrity: cada evento novo tem previous_hash "
         "apontando para o current_hash do evento anterior. Qualquer mutação posterior invalida a "
         "cadeia e é detectável por varredura offline. O backup/restore (DT-063) preserva "
         "audit_log_global mesmo em rollback — isso é regra dura: nenhuma operação de escopo de "
         "projeto pode apagar linhas de auditoria cross-projeto.")

    page_break(doc)

    # ─── 8. Endpoints principais ──────────────────────────────────────
    h1(doc, "8. API — endpoints principais")

    h2(doc, "8.1. Autenticação e bootstrap")
    table(doc,
          ["Método", "Path", "Descrição"],
          [
              ["POST", "/api/v1/auth/bootstrap", "Cria primeiro Admin na instância (só funciona quando zero admins)."],
              ["POST", "/api/v1/auth/login", "Login com e-mail e senha. Retorna JWT access_token."],
              ["POST", "/api/v1/auth/change-password", "Troca senha do usuário autenticado."],
              ["POST", "/api/v1/auth/reset-password", "Inicia fluxo de reset (envia e-mail com token)."],
          ],
          widths=[2, 5, 9])

    h2(doc, "8.2. Administração")
    table(doc,
          ["Método", "Path", "Descrição"],
          [
              ["GET",   "/api/v1/admin/projects/pending", "Lista solicitações e projetos com project_lifecycle_status."],
              ["POST",  "/api/v1/admin/projects/{id}/approve", "Aprova solicitação de projeto."],
              ["POST",  "/api/v1/admin/projects/{id}/reject", "Rejeita solicitação com motivo."],
              ["PATCH", "/api/v1/admin/projects/{id}/status", "Altera lifecycle para active/paused/inactive."],
              ["POST",  "/api/v1/admin/projects/requests/{id}/cleanup-orphan", "Limpa request aprovada cujo projeto foi deletado."],
              ["GET",   "/api/v1/admin/users", "Lista todos os usuários da instância."],
              ["PATCH", "/api/v1/admin/users/{id}/admin-flag", "Promove ou rebaixa papel de Admin."],
              ["POST",  "/api/v1/admin/invitations/admin", "Cria Admin novo via convite com senha temporária."],
              ["GET",   "/api/v1/admin/support", "Lista Equipe Sustentação (Admin + is_support)."],
              ["PATCH", "/api/v1/admin/support/{id}", "Liga/desliga flag is_support de um usuário."],
          ],
          widths=[2, 6, 8])

    h2(doc, "8.3. Métricas")
    table(doc,
          ["Método", "Path", "Descrição"],
          [
              ["GET", "/api/v1/metrics/health", "Health check público (load balancer / liveness probe)."],
              ["GET", "/api/v1/metrics/dashboard", "Agregação global (Admin + Sustentação). Aceita hours."],
              ["GET", "/api/v1/metrics/per-project", "Breakdown de uso de IA por projeto (Admin-only)."],
              ["GET", "/api/v1/metrics/prometheus", "Formato texto Prometheus para scrape externo."],
              ["GET", "/api/v1/projects/{id}/metrics/dashboard", "Métricas do projeto (Admin/Support/membro aceito)."],
          ],
          widths=[2, 6, 8])

    h2(doc, "8.4. Backups e releases")
    table(doc,
          ["Método", "Path", "Descrição"],
          [
              ["GET",   "/api/v1/projects/{id}/backups", "Lista backups do projeto."],
              ["POST",  "/api/v1/projects/{id}/backups", "Dispara backup manual (Admin ou GP)."],
              ["POST",  "/api/v1/projects/{id}/backups/{bid}/restore?confirm=true", "Restaura projeto a partir do backup."],
              ["GET",   "/api/v1/projects/{id}/backups/{bid}/download", "Download .zip do backup."],
              ["GET",   "/api/v1/admin/backups", "Lista agregada cross-projeto."],
              ["GET",   "/api/v1/backups/active", "Backups em andamento (banner global)."],
              ["GET",   "/api/v1/admin/releases", "Lista releases aplicadas e pending."],
              ["POST",  "/api/v1/admin/releases/{id}/apply", "Aplica release destrutiva com snapshot."],
              ["POST",  "/api/v1/admin/releases/{id}/rollback-project", "Rollback por projeto de release destrutiva."],
              ["GET",   "/api/v1/releases", "Changelog segmentado por papel (user-facing)."],
          ],
          widths=[2, 7, 7])

    h2(doc, "8.5. Tickets")
    table(doc,
          ["Método", "Path", "Descrição"],
          [
              ["GET",   "/api/v1/projects/{id}/incidents", "Lista tickets do projeto."],
              ["POST",  "/api/v1/projects/{id}/incidents", "Abre ticket (roteamento automático por papel)."],
              ["GET",   "/api/v1/incidents/{tid}", "Detalhe do ticket com comentários e anexos."],
              ["PATCH", "/api/v1/incidents/{tid}/status", "Altera status (open/in_progress/resolved/closed)."],
              ["POST",  "/api/v1/incidents/{tid}/comments", "Adiciona comentário."],
              ["POST",  "/api/v1/incidents/{tid}/attachments", "Upload de anexo (multipart, 10 MB, 9 ext)."],
              ["GET",   "/api/v1/incidents/{tid}/attachments/{aid}/download", "Download do anexo."],
              ["DELETE","/api/v1/incidents/{tid}/attachments/{aid}", "Exclui anexo (autor ou Admin)."],
              ["GET",   "/api/v1/admin/incidents", "Agregado cross-projeto (Admin + Sustentação)."],
          ],
          widths=[2, 7, 7])

    page_break(doc)

    # ─── 9. Diagramas de sequência ────────────────────────────────────
    h1(doc, "9. Diagramas de sequência")

    h2(doc, "9.1. Login Admin")
    para(doc, "Fluxo padrão de autenticação administrativa. JWT com expiração configurável é o único "
              "token emitido nesta versão — não há refresh token.")
    insert_diagram(doc, "seq_login_admin.png", "Login administrativo.")

    h2(doc, "9.2. Login via projeto")
    para(doc, "Entrada contextualizada em /p/{slug}: o frontend busca resumo do projeto e apresenta "
              "form de login com o nome do projeto no topo; após autenticar, valida membership e "
              "redireciona direto ao dashboard do projeto.")
    insert_diagram(doc, "seq_login_projeto.png", "Login via slug de projeto.")

    h2(doc, "9.3. Criação de projeto ponta a ponta")
    para(doc, "Da solicitação pública até a aprovação pelo Admin e promoção do solicitante a GP.")
    insert_diagram(doc, "seq_criar_projeto.png", "Criação e aprovação de projeto.")

    h2(doc, "9.4. Geração do OCG")
    para(doc, "Pipeline de 8 agentes (analyzer → 7 pilares em paralelo → consolidator) com "
              "escolha de provedor por criticidade e registro em ai_usage_log.")
    insert_diagram(doc, "seq_ocg_generation.png", "Geração do OCG em 8 agentes.")

    h2(doc, "9.5. Ticket até release")
    para(doc, "Rastreabilidade ticket → release conforme MVP 6 + MVP 7. A release seguinte à "
              "resolução amarra-se ao ticket via ref_id.")
    insert_diagram(doc, "seq_ticket_release.png", "Rastreabilidade ticket → release.")

    h2(doc, "9.6. Backup e restore")
    para(doc, "Scheduler diário, catch-up no startup e restore preservando audit_log_global.")
    insert_diagram(doc, "seq_backup_restore.png", "Ciclo de backup e restore (DT-063).")

    page_break(doc)

    # ─── 10. Fluxo e estado ───────────────────────────────────────────
    h1(doc, "10. Fluxos e estados")

    h2(doc, "10.1. Lifecycle do projeto")
    para(doc,
         "Projeto passa por quatro estágios principais: solicitado (externo), pendente (aguardando "
         "Admin), aprovado (linha em projects). A partir de aprovado, o Admin pode transitar entre "
         "active, paused e inactive sem nunca apagar dados.")
    insert_diagram(doc, "flow_projeto_lifecycle.png", "Estados do projeto (active/paused/inactive).")

    h2(doc, "10.2. Lifecycle do ticket")
    para(doc,
         "Aberto → em andamento → resolvido → fechado. Tickets resolvidos podem ser reabertos; "
         "tickets fechados entram no histórico permanente.")
    insert_diagram(doc, "flow_ticket_lifecycle.png", "Estados do ticket de incidente.")

    h2(doc, "10.3. Lifecycle da release")
    para(doc,
         "Declarada (YAML) → pending (sync no startup) → applied (auto ou manual). Releases "
         "destrutivas exigem confirmação e snapshot antes de aplicar.")
    insert_diagram(doc, "flow_release_lifecycle.png", "Estados da release.")

    page_break(doc)

    # ─── 11. Requisitos não-funcionais ────────────────────────────────
    h1(doc, "11. Requisitos não-funcionais")

    h2(doc, "11.1. Desempenho")
    bullet(doc, "Operações síncronas de listagem (projects, users, tickets, releases) devem retornar em até 2 segundos na instância do cliente com 100 projetos e 500 usuários ativos.")
    bullet(doc, "Geração do OCG completo (8 agentes) é assíncrona; UI mostra progresso e SSE/WebSocket para conclusão.")
    bullet(doc, "Backup por projeto não ultrapassa 60 segundos por projeto típico (até 50 MB de uploads).")

    h2(doc, "11.2. Disponibilidade")
    bullet(doc, "Alvo operacional: 99,5% mensal para ambiente produtivo em cliente de porte médio.")
    bullet(doc, "Health check /metrics/health é público e retorna status 200 + validação trivial de DB.")
    bullet(doc, "Falha de SMTP nunca deve derrubar o produto — logs gravam bounce e guard _is_non_deliverable_email bloqueia domínios de teste.")

    h2(doc, "11.3. Segurança")
    bullet(doc, "PATs e credenciais de provedores IA criptografadas com Fernet (chave derivada via PBKDF2 de GCA_MASTER_KEY).")
    bullet(doc, "Senhas armazenadas com bcrypt (rounds configuráveis via config).")
    bullet(doc, "JWT assinado com HS256 e chave exclusiva por instância.")
    bullet(doc, "audit_log_global com hash chain (previous_hash + current_hash) detecta adulteração.")
    bullet(doc, "Código em produção é protegido por Cython compile + PyArmor BCC + obfuscation de frontend + assinatura de integridade (ver docs/ANTI_REVERSE_ENGINEERING.md).")

    h2(doc, "11.4. Observabilidade")
    bullet(doc, "Logs estruturados em JSON via structlog; correlation_id em audit para ligar eventos.")
    bullet(doc, "Métricas em formato Prometheus expostas em /metrics/prometheus para scrape por Grafana Agent externo.")
    bullet(doc, "Breakdown por projeto em /admin/metrics e em /projects/{id}/metrics.")

    h2(doc, "11.5. Upgrade e operação")
    bullet(doc, "Script scripts/upgrade.sh idempotente com 9 etapas (backup pré → fetch → build → migrations → recreate → health → smoke).")
    bullet(doc, "Script scripts/restore.sh valida SHA256 antes de qualquer DROP e exige dupla confirmação.")
    bullet(doc, "Releases destrutivas geram snapshot automático por projeto ativo antes de aplicar.")

    page_break(doc)

    # ─── 12. Operação e manutenção ────────────────────────────────────
    h1(doc, "12. Operação e manutenção")

    h2(doc, "12.1. Scripts operacionais")
    table(doc,
          ["Script", "Função"],
          [
              ["scripts/upgrade.sh", "Upgrade idempotente em 9 etapas com rollback automático em falha."],
              ["scripts/backup.sh", "Backup completo da instância (postgres dump + volumes) pré-upgrade."],
              ["scripts/restore.sh", "Restore com validação SHA256 e dupla confirmação."],
              ["scripts/health-check.sh", "Health endpoints do stack."],
              ["scripts/capture_screenshots_v2.py", "Captura automatizada das 38 telas da aplicação."],
              ["scripts/gerar_diagramas_mermaid.py", "Gera diagramas Mermaid como PNG."],
              ["scripts/gerar_requisitos.py", "Gera este documento."],
              ["scripts/gerar_tutorial_instalacao.py", "Gera o tutorial de instalação."],
              ["scripts/build_production_images.sh", "Build de imagens com Cython + PyArmor + obfuscator."],
          ],
          widths=[6.5, 10])

    h2(doc, "12.2. Rotina de backup")
    bullet(doc, "12:00 (America/Sao_Paulo) diário — scheduler APScheduler dispara backup de todos projetos active.")
    bullet(doc, "Retenção: 10 últimos backups por projeto. Excedentes limpos após cada execução.")
    bullet(doc, "Catch-up no startup se last_backup_at > 24h.")
    bullet(doc, "Banner global no frontend indica backup em curso via polling /backups/active a cada 5 segundos.")

    h2(doc, "12.3. Monitoramento")
    bullet(doc, "GET /api/v1/metrics/health — público, para load balancer.")
    bullet(doc, "GET /api/v1/metrics/prometheus — texto Prometheus para scrape autenticado.")
    bullet(doc, "GET /api/v1/admin/metrics/dashboard — JSON agregado para UI.")
    bullet(doc, "GET /api/v1/admin/metrics/per-project — breakdown compartimentalizado.")

    page_break(doc)

    # ─── 13. Segurança e proteção de código ───────────────────────────
    h1(doc, "13. Segurança e proteção de código")

    h2(doc, "13.1. Princípio honesto")
    para(doc,
         "Não existe proteção absoluta em código que roda na máquina do cliente. O GCA implementa "
         "sete camadas que elevam o custo de engenharia reversa ao ponto de não ser economicamente "
         "viável para o cliente médio. Quem tem tempo, ferramentas profissionais e motivação "
         "suficiente ainda pode reverter — só proteção em hardware (HSM, enclave, atestação remota) "
         "atingiria garantia forte, e isso está fora do escopo desta versão.")

    h2(doc, "13.2. Sete camadas de proteção")
    numbered(doc, "Backend Python compilado via Cython: cada .py vira .c e depois .so/.pyd binário. Decompile produz pseudo-C ilegível.")
    numbered(doc, "PyArmor BCC wrapper sobre arquivos sensíveis (autenticação, vault, LLM key resolver). Adiciona camada de runtime check.")
    numbered(doc, "Imagens Docker multi-stage com base mínima. Bytecode e binários ficam apenas nas imagens, nunca no host.")
    numbered(doc, "Registry privado autenticado com tokens rotacionáveis. Imagens nunca são públicas.")
    numbered(doc, "Frontend React com minify, tree-shake e javascript-obfuscator (control-flow flattening + string array + dead code injection).")
    numbered(doc, "Verificação de integridade no startup: backend calcula SHA-256 dos próprios arquivos e compara com manifesto assinado. Modificação detectada → recusa subir.")
    numbered(doc, "Licença JWT com expiração, assinada com chave privada do GCA. Sem JWT válido, instância entra em modo somente-leitura ou não sobe.")

    h2(doc, "13.3. Isolamento de dados")
    bullet(doc, "PostgreSQL em container com volume nomeado (gca-postgres-data). Sem exposição de porta ao host por default.")
    bullet(doc, "Uploads em volume nomeado (gca-uploads-storage), acesso apenas pelo backend.")
    bullet(doc, "Backups em volume nomeado (gca-backups) com sha256 por arquivo.")
    bullet(doc, "JWT signing key em variável de ambiente; rotação manual disponível via reinstalação da instância.")

    page_break(doc)

    # ─── 14. Glossário ────────────────────────────────────────────────
    h1(doc, "14. Glossário")
    para(doc, "Termos em ordem alfabética. A referência cruzada ao contrato canônico é indicada quando aplicável.", size=10, color=SLATE_MEDIUM)

    glossary = [
        ("Admin", "Usuário com is_admin=true. Soberano da instância; configura provedores, aprova projetos, gerencia usuários. Não atua operacionalmente em projetos (contrato §4.1)."),
        ("Admin_viewer", "Papel virtual atribuído a Admin quando acessa projeto sem membership. Permite ver o projeto mas não editar conteúdo (só gerenciar GP)."),
        ("APScheduler", "Biblioteca Python de agendamento usada pelo GCA para disparar o backup diário às 12:00 e o catch-up no startup."),
        ("Arguidor", "Agente de IA que analisa documentos ingeridos e gera findings para o Gatekeeper. Criticidade média a alta. Configurável por projeto."),
        ("ai_usage_log", "Tabela que registra cada chamada a provedor de IA com provider, operation, tokens_input, tokens_output, cost_usd. Compartimentalizada por project_id."),
        ("audit_log_global", "Tabela de auditoria com hash chain (previous_hash + current_hash). Integridade verificável offline."),
        ("Backup", "Zip por projeto contendo JSONL de 29 tabelas filtradas por project_id + manifest.json + sha256 global. Gerenciado pelo project_backup_service (DT-063)."),
        ("BCC", "BytecodeCompiler wrapper do PyArmor. Variante gratuita de obfuscação usada no GCA."),
        ("Catch-up", "Rotina executada no startup do backend que dispara backup se o último foi há mais de 24h."),
        ("Changelog", "Lista de items de uma release com kind (mvp, feature, fix, ticket) e affected_roles. Segmentado por papel na UI /releases."),
        ("CodeGen", "Módulo de geração assistida de código (MVP 3). Usa OCG + provedor de IA + scaffolders determinísticos por linguagem."),
        ("Completion task", "Pendência pós-release por projeto quando a release adiciona campo novo obrigatório. Registrada em release_completion_tasks."),
        ("Compartimentalização", "Princípio do contrato §2.2: todo acesso a dado de projeto inclui project_id no predicado. Nenhum canal cruza projetos sem autorização explícita."),
        ("Contexto A", "Uso de IA para desenvolvimento do produto GCA em si. Pode usar IA premium. Custo do time de produto."),
        ("Contexto B", "Uso de IA pelo cliente dentro da instância dele. Configurável por instância e projeto. Custo do cliente."),
        ("Cython", "Compilador que transforma código Python em C e depois em binário .so/.pyd. Usado no empacotamento de produção."),
        ("Dev", "Papel canônico do projeto. Implementa código, opera ingestão, Arguidor, CodeGen e commits. Não aprova módulo no Gatekeeper."),
        ("Dogfood", "Prática de usar a própria instância do GCA para conduzir o desenvolvimento do GCA. O projeto Automação Jurídica Assistida é o dogfood ativo."),
        ("DT", "Dívida técnica. Identificada por número (DT-001 a DT-063). Catalogada em GCA_MVP_PROGRESS.md §3 (abertas) e §4 (quitadas)."),
        ("Emenda 2026-04-19", "Alteração formal no contrato canônico §4.1 (GP soberano do projeto) e §7 MVP 6 (Sustentação, anexos, seção/fluxo)."),
        ("Equipe Sustentação", "Conjunto de usuários com is_support=true (ou is_admin=true via herança). Recebe tickets com target_scope=admin."),
        ("Gatekeeper", "Módulo de avaliação baseado em sete pilares (Business, Architecture, Stack, Testing, Compliance, Risk, Deliverables). Bloqueia avanço se thresholds não são atingidos."),
        ("GP", "Gerente de Projeto. Soberano do projeto. Após emenda 2026-04-19 tem união de ações de Dev + Tester + QA."),
        ("incident_tickets", "Tabela de tickets de incidente (MVP 6). Campos target_scope, category, priority, status, title, description, section_reference, flow_description."),
        ("Ingestão", "Módulo que recebe documentos, detecta PII, quarentena quando aplicável, aciona Arguidor."),
        ("Inno Setup", "Ferramenta para criar instaladores Windows a partir de scripts .iss. Usada no empacotamento do GCA para Windows."),
        ("is_admin", "Flag em users. Habilita acesso à camada administrativa da instância."),
        ("is_support", "Flag em users que habilita recebimento de tickets escalados a Admin (MVP 6 Emenda). Independente de is_admin."),
        ("JWT", "JSON Web Token. Único tipo de token de autenticação nesta versão. Assinado com HS256."),
        ("Lifecycle", "Ciclo de vida. Aplicado a projetos (active/paused/inactive/archived), tickets (open/in_progress/resolved/closed) e releases (pending/applied/rolled_back)."),
        ("LLM", "Large Language Model. Provedor de IA usado pelo GCA (Anthropic, OpenAI, Google, DeepSeek, Ollama)."),
        ("manifest.json", "Arquivo dentro de cada backup e cada release YAML que lista conteúdo, contagens e hashes para verificação de integridade."),
        ("Mermaid", "Linguagem de marcação para diagramas (sequência, fluxo, estado). Renderizada para PNG via mmdc."),
        ("MVP", "Minimum Viable Product. No GCA, unidade de escopo canônico (MVP 1 a 7). Cada MVP tem em-escopo e fora-de-escopo obrigatórios."),
        ("OCG", "Objeto Canônico de Governança. Fonte única de verdade do projeto. Gerado pelo pipeline de 8 agentes a partir do questionário aprovado."),
        ("Ollama", "Provedor de IA local (self-hosted). Suportado pelo GCA via endpoint OpenAI-compatible."),
        ("Órfão", "Registro em project_requests com status=APPROVED cuja linha em projects não existe mais (efeito colateral de deleção hard no passado). Tratado pelo endpoint cleanup-orphan."),
        ("PAT", "Personal Access Token. Credencial de integração com Git (GitHub/GitLab/Bitbucket). Criptografada com Fernet."),
        ("PII", "Personally Identifiable Information. O GCA detecta e quarentena documentos com PII nas ingestões."),
        ("Playwright", "Framework de automação de navegador usado para captura das 38 telas da aplicação."),
        ("project_id", "Chave de compartimentalização. Obrigatória em todo predicado que acessa dado de projeto."),
        ("project_requests", "Tabela de solicitações externas de projeto. Estados: PENDING, APPROVED, REJECTED."),
        ("Pilares", "Sete categorias do Gatekeeper: P1 Business, P2 Architecture, P3 Stack, P4 Testing, P5 Compliance, P6 Risk, P7 Deliverables."),
        ("PyArmor", "Ferramenta de obfuscação de Python. Variante BCC usada no GCA (gratuita). Variante Pro opcional para máxima proteção."),
        ("QA", "Quality Assurance. Papel canônico do projeto. Revisa e aprova resultados, valida qualidade e compliance. Não edita conteúdo de teste."),
        ("RBAC", "Role-Based Access Control. Modelo de autorização do GCA (5 papéis canônicos + soberania cruzada Admin/GP)."),
        ("Release", "Unidade de entrega de software para a instância. Declarada em backend/releases/*.yaml. Pode ser destrutiva ou não-destrutiva."),
        ("Release Bundle", "Pacote consolidado de artefatos entregues ao fim do MVP 4."),
        ("Rollback", "Restauração de dados de projeto via snapshot prévio (DT-063). Por-projeto; não altera status global da release."),
        ("Scaffolder", "Gerador determinístico de estrutura inicial de projeto por linguagem (Java/Spring, Java/Quarkus, Kotlin/Spring, Go, C#, PHP, Node.js/NestJS+Express). Python continua LLM-only por design."),
        ("Schema", "Estrutura do banco de dados. Cerca de 30 tabelas em PostgreSQL. Migrações em backend/migrations/*.sql."),
        ("SLA", "Service Level Agreement. Alvo operacional do GCA: 99,5% mensal."),
        ("SMTP", "Simple Mail Transfer Protocol. Compartimentalizado por projeto (DT-016) com fallback global."),
        ("Snapshot", "Backup pré-release gerado automaticamente antes de aplicação de release destrutiva."),
        ("Solicitante", "Usuário externo que submete /solicitar-projeto. Quando o projeto é aprovado, torna-se GP automaticamente."),
        ("Status global", "Agregação de estado do produto por instância. Acessível em /admin/metrics."),
        ("Sustentação", "Ver \"Equipe Sustentação\"."),
        ("Tester", "Papel canônico do projeto. Cria, edita e executa testes; registra evidências."),
        ("target_scope", "Campo em incident_tickets. Valores: gp (ticket vai para GPs do projeto) ou admin (vai para Admins + Sustentação)."),
        ("Ticket", "Registro de incidente aberto por usuário. Campos obrigatórios incluem título, descrição, prioridade, categoria, seção e fluxo."),
        ("Thresholds", "Valores limites do Gatekeeper por pilar. Parametrizáveis em /admin (Settings)."),
        ("Ubuntu", "Distribuição Linux suportada pelo empacotamento .deb + install.sh."),
        ("Vite", "Ferramenta de build para frontend React + TypeScript. O GCA usa vite preview em produção para servir build estático."),
        ("Windows", "Sistema operacional suportado pelo empacotamento via instalador Inno Setup (.exe)."),
        ("YAML", "Formato usado para declarar releases em backend/releases/*.yaml."),
        ("ZIP", "Formato usado para empacotar backups por projeto."),
    ]

    for term, definition in sorted(glossary, key=lambda x: x[0].lower()):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        t_run = p.add_run(f"{term}. ")
        t_run.bold = True
        t_run.font.size = Pt(11)
        t_run.font.color.rgb = VIOLET
        d_run = p.add_run(definition)
        d_run.font.size = Pt(11)

    # Rodapé final
    doc.add_paragraph()
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = final.add_run("— Fim do documento —")
    fr.italic = True
    fr.font.color.rgb = SLATE_MEDIUM
    fr.font.size = Pt(10)


def main():
    doc = Document()
    # Margens
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    build(doc)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"✓ Documento gerado: {OUT_PATH}")
    print(f"  Tamanho: {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
